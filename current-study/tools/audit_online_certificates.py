from __future__ import annotations

import hashlib
import re
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from PIL import Image, ImageChops
from pypdf import PdfReader

import build_online_certificates as study


ROOT = Path(__file__).resolve().parents[1]
REPOSITORY = ROOT.parent
HUNSINGER_PDF = REPOSITORY / "online-it-credentials-research/source/Hunsinger_Smith_2008_IT_Certification.pdf"
WANG_PDF = REPOSITORY / "online-it-credentials-research/source/Wang_2023_MOOC_TAM_TPB.pdf"
HUNSINGER_PAGE = ROOT / "source/Hunsinger_Smith_2008_page_5.png"
WANG_PAGE = REPOSITORY / "online-it-credentials-research/working/wang-2023-page-05.png"


def fail(message: str) -> None:
    raise AssertionError(message)


def normalized(value: str) -> str:
    value = value.replace("\u2192", "->")
    value = re.sub(r"(?<=\w)-\s+(?=\w)", "", value)
    return re.sub(r"\s+", " ", value).strip()


def document_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def assert_embedded_exactly(docx_path: Path, image_path: Path) -> None:
    expected = image_path.read_bytes()
    with ZipFile(docx_path) as archive:
        media = [archive.read(name) for name in archive.namelist() if name.startswith("word/media/")]
    if not any(blob == expected for blob in media):
        embedded_hashes = ", ".join(sha256(blob)[:12] for blob in media)
        fail(f"{image_path.name} is not embedded byte-for-byte in the manuscript; media hashes: {embedded_hashes}")


def assert_same_pixels(actual: Path, expected: Image.Image, label: str) -> None:
    with Image.open(actual).convert("RGB") as image:
        if image.size != expected.size or ImageChops.difference(image, expected).getbbox() is not None:
            fail(f"{label} is not an unchanged crop of its retained source page")


def source_text(path: Path, page_indexes: list[int]) -> str:
    reader = PdfReader(path)
    return normalized(" ".join((reader.pages[index].extract_text() or "") for index in page_indexes))


def assert_contains(haystack: str, needle: str, label: str) -> None:
    if normalized(needle).casefold() not in normalized(haystack).casefold():
        fail(f"Missing {label}: {needle}")


def audit_sources() -> None:
    for path in (HUNSINGER_PDF, WANG_PDF, HUNSINGER_PAGE, WANG_PAGE):
        if not path.exists():
            fail(f"Required retained source is missing: {path}")

    hunsinger = source_text(HUNSINGER_PDF, [4, 5, 6, 7, 11, 12])
    for phrase in (
        "Attitude toward the behavior is significantly and positively related to intent to pursue IT certification",
        "Subjective Norm is significantly and positively related to intent to pursue IT certification",
        "Perceived Behavioral Control is significantly and positively related to intent to pursue IT certification",
        "I plan to earn",
        "I intend to pursue",
        "To the extent possible, I plan to pursue",
        "Very good / Very bad",
        "IT managers",
        "My professors",
        "Hiring managers",
        "My advisors",
        "My parents",
        "The general public",
        "Learning ability",
        "Knowledge",
        "Skills",
        "Money and resources",
        "All three original TPB constructs are significant",
        "Generally speaking, I do what _____________ think I should do.",
        "For me, having the ___________ to pursue an IT certification would make it [much easier---much more difficult] to earn an IT certification in the next twelve months.",
    ):
        assert_contains(hunsinger, phrase, "Hunsinger and Smith source evidence")

    wang = source_text(WANG_PDF, [4, 14])
    for phrase in (
        "H3: A learner's attitude towards using MOOCs has a positive impact on their behavioral intention",
        "H4: A learner's subjective norm has a positive impact on their behavioral intention",
        "H5: A learner's perceived behavioral control has a positive impact on their behavioral intention",
        "H3: ATT->BI",
        "H4: SN->BI",
        "H5: PBC->BI",
    ):
        # PDF extraction drops arrows and may use curly apostrophes.
        cleaned_phrase = phrase.replace("'", "’") if "learner's" in phrase else phrase
        if normalized(cleaned_phrase).casefold() not in wang.casefold() and normalized(phrase).casefold() not in wang.casefold():
            fail(f"Missing Wang source evidence: {phrase}")

    with Image.open(HUNSINGER_PAGE) as page:
        assert_same_pixels(
            study.HUNSINGER_FRAMEWORK,
            page.convert("RGB").crop((205, 990, 1265, 1645)),
            "Hunsinger and Smith framework figure",
        )
    with Image.open(WANG_PAGE) as page:
        assert_same_pixels(
            study.WANG_FRAMEWORK,
            page.convert("RGB").crop((145, 870, 770, 1455)),
            "Wang framework figure",
        )


def audit_manuscript() -> None:
    if not study.MASTER_OUT.exists():
        fail(f"Manuscript not found: {study.MASTER_OUT}")
    text = document_text(study.MASTER_OUT)
    compact = normalized(text)

    required = [
        study.TITLE,
        study.SUBTITLE,
        "CHAPTER 1: INTRODUCTION",
        "CHAPTER 2: LITERATURE REVIEW",
        "CHAPTER 3: RESEARCH FRAMEWORK",
        "CHAPTER 4: RESEARCH METHODOLOGY",
        "APPENDIX A: QUESTIONNAIRE",
        "Source: Constructed by the author based on Hunsinger and Smith (2008) and Wang (2023).",
        "Questionnaire items are adopted from Hunsinger and Smith (2008, pp. 252-254).",
        "A pilot study with 40 eligible students",
        "at least 384 eligible and complete responses",
        "One multiple linear regression will test H1-H3",
    ]
    required.extend(statement for _, statement in study.HYPOTHESES)
    for phrase in required:
        assert_contains(compact, phrase, "manuscript content")

    codes = [f"BI{i}" for i in range(1, 4)]
    codes += [f"ATT{i}" for i in range(1, 4)]
    codes += [f"SN-NB{i}" for i in range(1, 7)]
    codes += [f"SN-MC{i}" for i in range(1, 7)]
    codes += [f"PBC-CB{i}" for i in range(1, 5)]
    codes += [f"PBC-PF{i}" for i in range(1, 5)]
    if len(codes) != 26:
        fail("The expected questionnaire code list does not contain 26 responses")
    for code in codes:
        if not re.search(rf"(?<![A-Z0-9-]){re.escape(code)}(?![A-Z0-9-])", text):
            fail(f"Questionnaire item code is missing: {code}")

    # Compare every field item with the printed stems and listed labels, not just its code.
    item_rows = {}
    for table in Document(study.MASTER_OUT).tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            if cells and cells[0] in codes:
                if cells[0] in item_rows:
                    fail(f"Questionnaire item is duplicated: {cells[0]}")
                item_rows[cells[0]] = cells[1:]
    target = "a short online IT professional certificate"
    expected = {}
    for index, stem in enumerate(("I plan to earn", "I intend to pursue", "To the extent possible, I plan to pursue"), 1):
        expected[f"BI{index}"] = [f"{stem} {target} in the next twelve months."]
    for index, endpoints in enumerate((("Very good", "Very bad"), ("Very positive", "Very negative"), ("Very helpful", "Very unhelpful")), 1):
        expected[f"ATT{index}"] = list(endpoints)
    referents = ("IT managers", "My professors", "Hiring managers", "My advisors", "My parents", "The general public")
    for index, referent in enumerate(referents, 1):
        verb = "thinks" if referent == "The general public" else "think"
        embedded = referent if referent == "IT managers" else referent[0].lower() + referent[1:]
        expected[f"SN-NB{index}"] = [referent, f"{referent} {verb} I should pursue {target} within the next twelve months.", "+3 to -3"]
        expected[f"SN-MC{index}"] = [referent, f"Generally speaking, I do what {embedded} {verb} I should do.", "1 to 7"]
    for index, factor in enumerate(("learning ability", "knowledge", "skills", "money and resources"), 1):
        expected[f"PBC-CB{index}"] = [factor.title(), f"I have the {factor} to earn {target} within the next twelve months.", "+3 to -3"]
        expected[f"PBC-PF{index}"] = [factor.title(), f"For me, having the {factor} to pursue {target} would make it [much easier---much more difficult] to earn {target} in the next twelve months.", "1 to 7"]
    if item_rows != expected:
        changed = [code for code in codes if item_rows.get(code) != expected[code]]
        fail(f"Questionnaire wording or response cells differ from the source-based specification: {changed}")

    for phrase in (
        "Very good",
        "Very bad",
        "Very positive",
        "Very negative",
        "Very helpful",
        "Very unhelpful",
        "IT managers think I should pursue",
        "The general public thinks I should pursue",
        "I have the learning ability to earn",
        "I have the money and resources to earn",
        "from +3 (strongly agree) to -3 (strongly disagree)",
        "from 1 (very undesirable) to 7 (very desirable)",
        "from 1 (much easier) to 7 (much more difficult)",
    ):
        assert_contains(compact, phrase, "questionnaire wording")

    prohibited = (
        "Prepared for Dr.",
        "Prepared for Professor",
        "Evidence Pack",
        "neurological condition",
        "implausibly short completion time",
        "missing responses",
        "self-developed vignette",
        "First endpoint (+3)",
    )
    for phrase in prohibited:
        if phrase.casefold() in compact.casefold():
            fail(f"Prohibited or obsolete wording remains: {phrase}")

    assert_embedded_exactly(study.MASTER_OUT, study.HUNSINGER_FRAMEWORK)
    assert_embedded_exactly(study.MASTER_OUT, study.WANG_FRAMEWORK)
    assert_embedded_exactly(study.MASTER_OUT, study.AUTHOR_FRAMEWORK)


def main() -> int:
    try:
        audit_sources()
        audit_manuscript()
    except AssertionError as error:
        print(f"ONLINE-CERTIFICATE SOURCE-FIDELITY AUDIT: FAIL\n{error}", file=sys.stderr)
        return 1
    print("ONLINE-CERTIFICATE SOURCE-FIDELITY AUDIT: PASS")
    print("Verified: two direct published framework crops, three published paths, and 26 questionnaire responses.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
