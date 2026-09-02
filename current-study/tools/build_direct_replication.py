from __future__ import annotations

import argparse
from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_documents as base


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = PROJECT_ROOT / "outputs"
SOURCE_DIR = PROJECT_ROOT / "source"

MASTER_OUT = OUTPUTS / "Master_Thesis_Digital_Badges_Thai_IT_Students.docx"
EVIDENCE_OUT = OUTPUTS / "Evidence_Pack_Author_Constructed_Framework_and_Questionnaire.docx"

SOURCE_MODEL_PAGE = SOURCE_DIR / "steenkamp_page_10_model.png"
SOURCE_APPENDIX_28 = SOURCE_DIR / "steenkamp_page_28_appendix.png"
SOURCE_APPENDIX_29 = SOURCE_DIR / "steenkamp_page_29_appendix.png"
SOURCE_TAM3_FIGURE = SOURCE_DIR / "Venkatesh_Bala_2008_TAM3_Figure_2.png"
SOURCE_MIAO_PAGE = SOURCE_DIR / "miao_page_9_framework.png"

MODEL_CROP = OUTPUTS / "Published_Framework_Steenkamp_Figure_1_Direct_Crop.png"
APPENDIX_28_CROP = OUTPUTS / "Published_Questionnaire_Steenkamp_Appendix_p28_Direct_Crop.png"
APPENDIX_29_CROP = OUTPUTS / "Published_Questionnaire_Steenkamp_Appendix_p29_Direct_Crop.png"
TAM3_FIGURE = OUTPUTS / "Published_Framework_Venkatesh_Bala_TAM3_Figure_2.png"
MIAO_FIGURE = OUTPUTS / "Published_Framework_Miao_2024_Figure_1_Direct_Crop.png"
AUTHOR_FRAMEWORK = OUTPUTS / "Conceptual_Framework_Author_Constructed_Digital_Badges.png"

TITLE = (
    "Factors Influencing Thai IT Students' Intentions to Use University-Issued "
    "Digital Badges for IT Micro-Credentials in Job Applications"
)
SUBTITLE = "A Quantitative Study of Usefulness, Job Relevance, Ease of Use, and Social Influence"

BLUE = RGBColor(31, 78, 121)
GRAY = RGBColor(89, 89, 89)


HYPOTHESES = [
    (
        "H1",
        "Perceived ease of use has a positive direct effect on students' perceived usefulness "
        "of university-issued digital badges in job applications.",
    ),
    (
        "H2",
        "Job application relevance has a positive direct effect on students' perceived usefulness "
        "of university-issued digital badges in job applications.",
    ),
    (
        "H3",
        "Perceived usefulness has a positive direct effect on students' behavioural intention "
        "to use university-issued digital badges in job applications.",
    ),
    (
        "H4",
        "Subjective norm has a positive direct effect on students' behavioural intention to use "
        "university-issued digital badges in job applications.",
    ),
]


CONSTRUCTS = [
    (
        "Perceived Usefulness",
        "PU",
        "Figure 1 and hypothesis development (pp. 10-11); Appendix (p. 28)",
        "PU1-PU4",
        "4",
    ),
    (
        "Perceived Ease of Use",
        "PEOU",
        "Figure 1 and hypothesis development (pp. 10-11); Appendix (p. 28)",
        "PEOU1-PEOU4",
        "4",
    ),
    (
        "Computer Self-Efficacy",
        "CSE",
        "Figure 1 and hypothesis development (pp. 12-13); Appendix (p. 28)",
        "CSE1-CSE3",
        "3",
    ),
    (
        "Perceptions of External Control",
        "PEC",
        "Figure 1 and hypothesis development (pp. 12-13); Appendix (p. 28)",
        "PEC1-PEC4",
        "4",
    ),
    (
        "Computer Playfulness",
        "CPLAY",
        "Figure 1 and hypothesis development (pp. 12-13); Appendix (p. 28)",
        "CPLAY1-CPLAY4",
        "4",
    ),
    (
        "Computer Anxiety",
        "CANX",
        "Figure 1 and hypothesis development (pp. 12-13); Appendix (p. 28)",
        "CANX1-CANX4",
        "4",
    ),
    (
        "Subjective Norm",
        "SN",
        "Figure 1 and hypothesis development (pp. 11-12); Appendix (pp. 28-29)",
        "SN1-SN4",
        "4",
    ),
    (
        "Image",
        "IMG",
        "Figure 1 and hypothesis development (pp. 11-12); Appendix (p. 29)",
        "IMG1-IMG3",
        "3",
    ),
    (
        "Job Application Relevance",
        "REL",
        "Figure 1 and hypothesis development (p. 12); Appendix (p. 29)",
        "REL1-REL3",
        "3",
    ),
    (
        "Result Demonstrability",
        "RES",
        "Figure 1 and hypothesis development (p. 12); Appendix (p. 29)",
        "RES1-RES4",
        "4",
    ),
    (
        "Behavioural Intention to Use DB",
        "BI",
        "Figure 1 and hypothesis development (pp. 10-11); Appendix (p. 29)",
        "BI1-BI3",
        "3",
    ),
]


ITEMS = [
    ("PU1", "Perceived Usefulness", "Using university-issued digital badges would increase my employability when I apply for jobs.", "Davis (1989); Venkatesh & Bala (2008)"),
    ("PU2", "Perceived Usefulness", "Using university-issued digital badges would increase the quality of my job applications.", "Davis (1989); Venkatesh & Bala (2008)"),
    ("PU3", "Perceived Usefulness", "Using university-issued digital badges would enhance my effectiveness when I apply for jobs.", "Davis (1989); Venkatesh & Bala (2008)"),
    ("PU4", "Perceived Usefulness", "I would find using university-issued digital badges useful when applying for jobs.", "Davis (1989); Venkatesh & Bala (2008)"),
    ("PEOU1", "Perceived Ease of Use", "If I were to obtain university-issued digital badges, my interaction with them would be clear and understandable.", "Davis (1989); Venkatesh & Bala (2008)"),
    ("PEOU2", "Perceived Ease of Use", "If I were to obtain university-issued digital badges, interacting with them would not require a lot of my mental effort.", "Davis (1989); Venkatesh & Bala (2008)"),
    ("PEOU3", "Perceived Ease of Use", "If I were to obtain university-issued digital badges, I would find them easy to use.", "Davis (1989); Venkatesh & Bala (2008)"),
    ("PEOU4", "Perceived Ease of Use", "If I were to obtain university-issued digital badges, I would find it easy to get them to do what I want them to do.", "Davis (1989); Venkatesh & Bala (2008)"),
    ("CSE1", "Computer Self-Efficacy", "I could complete a job or task using a computer if there was no one around to tell me what to do.", "Sykes et al. (2014)"),
    ("CSE2", "Computer Self-Efficacy", "I could complete a job or task using a computer if I could contact someone if I got stuck.", "Sykes et al. (2014)"),
    ("CSE3", "Computer Self-Efficacy", "I could complete a job or task using a computer if I had a lot of time to complete the job for which the software was provided.", "Sykes et al. (2014)"),
    ("PEC1", "Perceptions of External Control", "I would have control over my university-issued digital badges.", "Venkatesh & Bala (2008)"),
    ("PEC2", "Perceptions of External Control", "I would have the resources necessary to use university-issued digital badges.", "Venkatesh & Bala (2008)"),
    ("PEC3", "Perceptions of External Control", "Given the resources, opportunities and knowledge it takes to use university-issued digital badges, it would be easy for me to use university-issued digital badges.", "Venkatesh & Bala (2008)"),
    ("PEC4", "Perceptions of External Control", "University-issued digital badges are not compatible with other means I use for communicating my skills and competencies to potential employers.", "Venkatesh & Bala (2008)"),
    ("CPLAY1", "Computer Playfulness", "I would characterise myself as spontaneous when using computers.", "Venkatesh & Bala (2008)"),
    ("CPLAY2", "Computer Playfulness", "I would characterise myself as creative when using computers.", "Venkatesh & Bala (2008)"),
    ("CPLAY3", "Computer Playfulness", "I would characterise myself as playful when using computers.", "Venkatesh & Bala (2008)"),
    ("CPLAY4", "Computer Playfulness", "I would characterise myself as unoriginal when using computers.", "Venkatesh & Bala (2008)"),
    ("CANX1", "Computer Anxiety", "Computers do not scare me at all.", "Venkatesh & Bala (2008)"),
    ("CANX2", "Computer Anxiety", "Working with a computer makes me nervous.", "Venkatesh & Bala (2008)"),
    ("CANX3", "Computer Anxiety", "Computers make me feel uncomfortable.", "Venkatesh & Bala (2008)"),
    ("CANX4", "Computer Anxiety", "Computers make me feel uneasy.", "Venkatesh & Bala (2008)"),
    ("SN1", "Subjective Norm", "People who influence my behaviour think that I should use university-issued digital badges when applying for jobs.", "Venkatesh & Bala (2008); Venkatesh et al. (2003)"),
    ("SN2", "Subjective Norm", "People who are important to me think that I should use university-issued digital badges when applying for jobs.", "Venkatesh & Bala (2008); Venkatesh et al. (2003)"),
    ("SN3", "Subjective Norm", "My university would be helpful in using university-issued digital badges for job applications.", "Venkatesh & Bala (2008); Venkatesh et al. (2003)"),
    ("SN4", "Subjective Norm", "In general, my university would support the use of university-issued digital badges for job applications.", "Venkatesh & Bala (2008); Venkatesh et al. (2003)"),
    ("IMG1", "Image", "Students who use university-issued digital badges when applying for jobs are likely to have more prestige than those who do not.", "Venkatesh & Bala (2008)"),
    ("IMG2", "Image", "Students who use university-issued digital badges when applying for jobs will have a high profile.", "Venkatesh & Bala (2008)"),
    ("IMG3", "Image", "Being able to display university-issued digital badges when applying for jobs would be a status symbol for students.", "Venkatesh & Bala (2008)"),
    ("REL1", "Job Application Relevance", "When job searching, usage of university-issued digital badges is important.", "Venkatesh & Bala (2008)"),
    ("REL2", "Job Application Relevance", "When job searching, usage of university-issued digital badges is relevant.", "Venkatesh & Bala (2008)"),
    ("REL3", "Job Application Relevance", "The use of university-issued digital badges is pertinent to my job searching activities.", "Venkatesh & Bala (2008)"),
    ("RES1", "Result Demonstrability", "I would have no difficulty telling others about the results of using university-issued digital badges.", "Venkatesh & Bala (2008)"),
    ("RES2", "Result Demonstrability", "I believe I could communicate to others the consequences for me of using university-issued digital badges.", "Venkatesh & Bala (2008)"),
    ("RES3", "Result Demonstrability", "The results of using university-issued digital badges would be apparent to me.", "Venkatesh & Bala (2008)"),
    ("RES4", "Result Demonstrability", "I would have difficulty explaining why using university-issued digital badges for job searching may or may not be beneficial.", "Venkatesh & Bala (2008)"),
    ("BI1", "Behavioral Intention", "Assuming I had access to university-issued digital badges I intend to use them when I apply for jobs.", "Venkatesh & Bala (2008); Venkatesh & Davis (2000)"),
    ("BI2", "Behavioral Intention", "Given that I had access to university-issued digital badges, I predict that I would use them when I apply for jobs.", "Venkatesh & Bala (2008); Venkatesh & Davis (2000)"),
    ("BI3", "Behavioral Intention", "If I had access to university-issued digital badges, I would plan to use the when I apply for jobs in the next 12 months.", "Venkatesh & Bala (2008); Venkatesh & Davis (2000)"),
]


REFERENCES = [
    "Abdullah, F., and Ward, R. (2016). Developing a general extended technology acceptance model for e-learning (GETAMEL) by analysing commonly used external factors. Computers in Human Behavior, 56, 238-256. https://doi.org/10.1016/j.chb.2015.11.036",
    "Abdullah, F., Ward, R., and Ahmed, E. (2016). Investigating the influence of the most commonly used external variables of TAM on students' perceived ease of use and perceived usefulness of e-portfolios. Computers in Human Behavior, 63, 75-90. https://doi.org/10.1016/j.chb.2016.05.014",
    "Ajzen, I. (1991). The theory of planned behavior. Organizational Behavior and Human Decision Processes, 50(2), 179-211. https://doi.org/10.1016/0749-5978(91)90020-T",
    "Brislin, R. W. (1970). Back-translation for cross-cultural research. Journal of Cross-Cultural Psychology, 1(3), 185-216. https://doi.org/10.1177/135910457000100301",
    "Council of the European Union. (2022). Council Recommendation of 16 June 2022 on a European approach to micro-credentials for lifelong learning and employability (2022/C 243/02). Official Journal of the European Union, C 243, 10-25.",
    "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319-340. https://doi.org/10.2307/249008",
    "Fishbein, M., and Ajzen, I. (1975). Belief, attitude, intention and behavior: An introduction to theory and research. Addison-Wesley.",
    "Granic, A., and Marangunic, N. (2019). Technology acceptance model in educational context: A systematic literature review. British Journal of Educational Technology, 50(5), 2572-2593. https://doi.org/10.1111/bjet.12864",
    "Hair, J. F., Hult, G. T. M., Ringle, C. M., and Sarstedt, M. (2022). A primer on partial least squares structural equation modeling (PLS-SEM) (3rd ed.). Sage.",
    "Hertzog, M. A. (2008). Considerations in determining sample size for pilot studies. Research in Nursing & Health, 31(2), 180-191. https://doi.org/10.1002/nur.20247",
    "Johanson, G. A., and Brooks, G. P. (2010). Initial scale development: Sample size for pilot studies. Educational and Psychological Measurement, 70(3), 394-400. https://doi.org/10.1177/0013164409355692",
    "Kiiskila, P., Kukkonen, A., and Pirkkalainen, H. (2023). Are micro-credentials valuable for students? Perspective on verifiable digital credentials. SN Computer Science, 4, Article 366. https://doi.org/10.1007/s42979-023-01797-y",
    "Krejcie, R. V., and Morgan, D. W. (1970). Determining sample size for research activities. Educational and Psychological Measurement, 30(3), 607-610. https://doi.org/10.1177/001316447003000308",
    "Miao, M., Ahmed, M., Ahsan, N., and Qamar, B. (2024). Intention to use technology for micro-credential programs: Evidence from technology acceptance and self-determination model. International Journal of Educational Management, 38(4), 948-977. https://doi.org/10.1108/IJEM-02-2023-0066",
    "Moore, G. C., and Benbasat, I. (1991). Development of an instrument to measure the perceptions of adopting an information technology innovation. Information Systems Research, 2(3), 192-222. https://doi.org/10.1287/isre.2.3.192",
    "Schepers, J., and Wetzels, M. (2007). A meta-analysis of the technology acceptance model: Investigating subjective norm and moderation effects. Information & Management, 44(1), 90-103. https://doi.org/10.1016/j.im.2006.10.007",
    "Steenkamp, N., Fisher, R., and Nesbit, T. (2024). Understanding accounting students' intentions to use digital badges to showcase employability skills. Accounting Education, 33(6), 906-934. https://doi.org/10.1080/09639284.2023.2276200",
    "Sykes, T. A., Venkatesh, V., and Johnson, J. L. (2014). Enterprise system implementation and employee job performance: Understanding the role of advice networks. MIS Quarterly, 38(1), 51-72. https://doi.org/10.25300/MISQ/2014/38.1.03",
    "Taylor, S., and Todd, P. A. (1995). Understanding information technology usage: A test of competing models. Information Systems Research, 6(2), 144-176. https://doi.org/10.1287/isre.6.2.144",
    "UNESCO. (2022). Towards a common definition of micro-credentials. UNESCO. https://unesdoc.unesco.org/ark:/48223/pf0000381668",
    "Ursavas, O. F., Yalcin, Y., and Bakir, E. (2019). The effect of subjective norms on preservice and in-service teachers' behavioural intentions to use technology: A multigroup multimodel study. British Journal of Educational Technology, 50(5), 2501-2519. https://doi.org/10.1111/bjet.12834",
    "Venkatesh, V. (2000). Determinants of perceived ease of use: Integrating control, intrinsic motivation, and emotion into the technology acceptance model. Information Systems Research, 11(4), 342-365. https://doi.org/10.1287/isre.11.4.342.11872",
    "Venkatesh, V., and Bala, H. (2008). Technology acceptance model 3 and a research agenda on interventions. Decision Sciences, 39(2), 273-315. https://doi.org/10.1111/j.1540-5915.2008.00192.x",
    "Venkatesh, V., and Davis, F. D. (2000). A theoretical extension of the technology acceptance model: Four longitudinal field studies. Management Science, 46(2), 186-204. https://doi.org/10.1287/mnsc.46.2.186.11926",
    "Venkatesh, V., Morris, M. G., Davis, G. B., and Davis, F. D. (2003). User acceptance of information technology: Toward a unified view. MIS Quarterly, 27(3), 425-478. https://doi.org/10.2307/30036540",
]


MASTER_METHOD_REFERENCES = [
    "Cronbach, L. J. (1951). Coefficient alpha and the internal structure of tests. Psychometrika, 16(3), 297-334. https://doi.org/10.1007/BF02310555",
    "Dawes, J. (2008). Do data characteristics change according to the number of scale points used? An experiment using 5-point, 7-point and 10-point scales. International Journal of Market Research, 50(1), 61-104. https://doi.org/10.1177/147078530805000106",
    "Hair, J. F., Black, W. C., Babin, B. J., and Anderson, R. E. (2019). Multivariate data analysis (8th ed.). Cengage.",
]

MASTER_REFERENCES = sorted(
    [
        reference
        for reference in REFERENCES
        if "A primer on partial least squares structural equation modeling" not in reference
    ]
    + MASTER_METHOD_REFERENCES,
    key=str.casefold,
)


assert len(ITEMS) == 40

SELECTED_CODES = ["PEOU", "REL", "PU", "SN", "BI"]
STUDY_CONSTRUCTS = [construct for construct in CONSTRUCTS if construct[1] in SELECTED_CODES]
STUDY_ITEMS = [item for item in ITEMS if item[0].rstrip("0123456789") in SELECTED_CODES]
assert len(STUDY_CONSTRUCTS) == 5
assert len(STUDY_ITEMS) == 18


OPERATIONAL_DEFINITIONS = {
    "PU": (
        "The extent to which a student believes that using a university-issued digital badge "
        "will improve the effectiveness and quality of a job application."
    ),
    "PEOU": (
        "The extent to which a student believes that using a university-issued digital badge "
        "will be clear, understandable, and free of effort."
    ),
    "CSE": "A student's belief in their ability to complete computer-based tasks.",
    "PEC": (
        "A student's perception that the resources, knowledge, opportunities, and support needed "
        "to use digital badges are available."
    ),
    "CPLAY": "The degree of cognitive spontaneity a student experiences when interacting with computers.",
    "CANX": "A student's apprehension or uneasiness when faced with using computers.",
    "SN": (
        "The extent to which a student perceives that important people and the university expect "
        "or support the use of digital badges in job applications."
    ),
    "IMG": (
        "The extent to which a student believes that using digital badges will enhance status, "
        "prestige, or profile."
    ),
    "REL": (
        "The extent to which a student believes that digital badges are applicable and pertinent "
        "to job-search and job-application activities."
    ),
    "RES": (
        "The extent to which a student believes that the results of using digital badges are "
        "tangible, observable, and communicable."
    ),
    "BI": (
        "A student's stated intention or plan to use university-issued digital badges when "
        "applying for jobs."
    ),
}


CONSTRUCT_THEORY_SOURCES = {
    "PU": "Davis (1989); Venkatesh and Davis (2000); Venkatesh and Bala (2008); Miao et al. (2024); Steenkamp et al. (2024)",
    "PEOU": "Davis (1989); Venkatesh and Davis (2000); Venkatesh and Bala (2008); Miao et al. (2024); Steenkamp et al. (2024)",
    "CSE": "Venkatesh (2000); Sykes et al. (2014); Steenkamp et al. (2024)",
    "PEC": "Venkatesh (2000); Venkatesh and Bala (2008); Steenkamp et al. (2024)",
    "CPLAY": "Venkatesh (2000); Venkatesh and Bala (2008); Steenkamp et al. (2024)",
    "CANX": "Venkatesh et al. (2003); Venkatesh and Bala (2008); Steenkamp et al. (2024)",
    "SN": "Taylor and Todd (1995); Venkatesh and Davis (2000); Schepers and Wetzels (2007); Venkatesh and Bala (2008); Steenkamp et al. (2024)",
    "IMG": "Venkatesh and Davis (2000); Venkatesh and Bala (2008); Steenkamp et al. (2024)",
    "REL": "Venkatesh and Davis (2000); Venkatesh and Bala (2008); Abdullah and Ward (2016); Abdullah et al. (2016); Steenkamp et al. (2024)",
    "RES": "Venkatesh and Bala (2008); Steenkamp et al. (2024)",
    "BI": "Fishbein and Ajzen (1975); Ajzen (1991); Venkatesh and Davis (2000); Venkatesh and Bala (2008); Steenkamp et al. (2024)",
}


THESIS_HYPOTHESES = HYPOTHESES

SOURCE_PATHS = [
    ["H1", "Perceived Ease of Use -> Perceived Usefulness", "Steenkamp et al. (2024), source H3", "Supported (beta = .671, p < .001)"],
    ["H2", "Job Application Relevance -> Perceived Usefulness", "Steenkamp et al. (2024), source H8", "Supported (beta = .229, p = .047)"],
    ["H3", "Perceived Usefulness -> Behavioural Intention to Use DB", "Steenkamp et al. (2024), source H1", "Supported (beta = .403, p = .010)"],
    ["H4", "Subjective Norm -> Behavioural Intention to Use DB", "Steenkamp et al. (2024), source H4", "Supported (beta = .593, p < .001)"],
]

SELECTED_SOURCE_HYPOTHESES = [
    ("H3", "Perceived ease of use has a positive direct effect on students' perceived usefulness of digital badges in applying for jobs."),
    ("H8", "Job application relevance has a positive direct effect on perceived usefulness of digital badges in applying for jobs."),
    ("H1", "Perceived usefulness has a positive direct effect on students' intention to use digital badges in applying for jobs."),
    ("H4", "Subjective norm has a positive direct effect on students' intention to use digital badges in applying for jobs."),
]

VARIABLE_LITERATURE = {
    "PEOU": (
        "Davis (1989) defined perceived ease of use as the belief that using a particular system would "
        "require little effort. Venkatesh and Davis (2000) retained it as a central technology-acceptance "
        "belief that can also increase perceived usefulness. Venkatesh and Bala (2008) incorporated the "
        "construct into TAM3 and examined the conditions that form ease-of-use judgements. In a "
        "micro-credential setting, Miao et al. (2024) treated perceived ease of use as the accessibility, "
        "adaptability, usability, and effortlessness of the supporting technology. Steenkamp et al. (2024) "
        "applied the same construct to university-issued digital badges. In this study, perceived ease of "
        "use means the extent to which a student believes that using a university-issued digital badge "
        "will be clear, understandable, and free of effort."
    ),
    "REL": (
        "Venkatesh and Davis (2000) described job relevance as the extent to which a target system is "
        "applicable to a person's work. Venkatesh and Bala (2008) positioned job relevance as a cognitive "
        "instrumental determinant of perceived usefulness. Abdullah and Ward (2016) identified job "
        "relevance as a commonly studied external factor in educational technology acceptance, and Abdullah "
        "et al. (2016) examined the relevance of technology to students' learning tasks in an e-portfolio "
        "setting. Steenkamp et al. (2024) used the exact contextual label Job Application Relevance for "
        "digital badges. In this study, job application relevance means the extent to which a student "
        "believes that using university-issued digital badges is important, relevant, and pertinent to "
        "job-search and job-application activities."
    ),
    "PU": (
        "Davis (1989) defined perceived usefulness as the belief that using a system will improve task or "
        "job performance. Venkatesh and Davis (2000) treated it as a direct determinant of behavioural "
        "intention, while Venkatesh and Bala (2008) retained this role in TAM3. Miao et al. (2024) applied "
        "perceived usefulness to technology used for micro-credential programmes and related it to improved "
        "learning performance and efficiency. Steenkamp et al. (2024) measured usefulness in terms of "
        "employability, job-application quality, effectiveness, and practical utility. In this study, "
        "perceived usefulness means the extent to which a student believes that using a university-issued "
        "digital badge will improve the effectiveness and quality of a job application."
    ),
    "SN": (
        "Fishbein and Ajzen (1975) introduced subjective norm as perceived social pressure from important "
        "others regarding a behaviour. Taylor and Todd (1995) applied the construct to information-technology "
        "use, and Venkatesh and Davis (2000) incorporated it into TAM2 as a social-influence process. "
        "Venkatesh and Bala (2008) retained subjective norm in TAM3, while Steenkamp et al. (2024) applied "
        "it to encouragement and support for using digital badges in job applications. In this study, "
        "subjective norm means the extent to which a student perceives that important people and the "
        "university expect or support the use of university-issued digital badges in job applications."
    ),
    "BI": (
        "Fishbein and Ajzen (1975) treated behavioural intention as a person's readiness or plan to perform "
        "a behaviour, and Ajzen (1991) retained intention as the closest antecedent of planned behaviour. "
        "Davis (1989) used intention to explain future system use, while Venkatesh and Davis (2000) and "
        "Venkatesh and Bala (2008) positioned it as the principal outcome of the technology-acceptance "
        "process. Steenkamp et al. (2024) specified the construct as Behavioural Intention to Use DB. In "
        "this study, it means a student's stated intention, prediction, and plan to use university-issued "
        "digital badges when applying for jobs."
    ),
}

RELATIONSHIP_LITERATURE = [
    (
        "2.4.1 Perceived Ease of Use and Perceived Usefulness",
        "The original TAM proposed that a system that is easier to use can be judged as more useful because "
        "less effort is required to obtain its benefits (Davis, 1989). Venkatesh and Davis (2000) and "
        "Venkatesh and Bala (2008) retained the positive perceived-ease-of-use to perceived-usefulness path "
        "in TAM2 and TAM3. Reviews of educational technology acceptance also report this relationship across "
        "e-learning settings (Abdullah & Ward, 2016; Granic & Marangunic, 2019). In the digital-badge study, "
        "Steenkamp et al. (2024) found a positive and significant effect (beta = .671, p < .001). Therefore, "
        "H1 proposes a positive effect of Perceived Ease of Use on Perceived Usefulness."
    ),
    (
        "2.4.2 Job Application Relevance and Perceived Usefulness",
        "TAM2 proposes that a technology is more useful when it is directly applicable to the task for which "
        "it is being considered (Venkatesh & Davis, 2000). TAM3 retained Job Relevance as a determinant of "
        "Perceived Usefulness (Venkatesh & Bala, 2008). Educational technology studies have likewise treated "
        "task or job relevance as an external factor shaping usefulness judgements (Abdullah & Ward, 2016; "
        "Abdullah et al., 2016). Steenkamp et al. (2024) contextualised the construct as Job Application "
        "Relevance and found a positive significant effect on Perceived Usefulness (beta = .229, p = .047). "
        "Therefore, H2 proposes a positive effect of Job Application Relevance on Perceived Usefulness."
    ),
    (
        "2.4.3 Perceived Usefulness and Behavioural Intention to Use DB",
        "Perceived Usefulness is a central predictor of Behavioural Intention in the original TAM (Davis, "
        "1989), TAM2 (Venkatesh & Davis, 2000), and TAM3 (Venkatesh & Bala, 2008). Systematic reviews support "
        "the continued relevance of this relationship in educational technologies (Abdullah & Ward, 2016; "
        "Granic & Marangunic, 2019). Miao et al. (2024) also found usefulness to be important in students' "
        "acceptance of technology for micro-credential programmes. In the closest digital-badge study, "
        "Steenkamp et al. (2024) reported a positive significant effect on intention (beta = .403, p = .010). "
        "Therefore, H3 proposes a positive effect of Perceived Usefulness on Behavioural Intention to Use DB."
    ),
    (
        "2.4.4 Subjective Norm and Behavioural Intention to Use DB",
        "Subjective Norm represents the influence of people or institutions that matter to the potential "
        "user. Taylor and Todd (1995) found that social influence contributed to technology-use intention. "
        "A meta-analysis by Schepers and Wetzels (2007) confirmed a significant relationship between "
        "Subjective Norm and Behavioural Intention, and Ursavas et al. (2019) found the relationship among "
        "pre-service and in-service teachers. Steenkamp et al. (2024) found a positive significant effect of "
        "Subjective Norm on Behavioural Intention to Use DB (beta = .593, p < .001). Therefore, H4 proposes "
        "a positive effect of Subjective Norm on Behavioural Intention to Use DB."
    ),
]


FIELD_ITEM_CORRECTIONS = {
    "BI3": (
        "If I had access to university-issued digital badges, I would plan to use them when I "
        "apply for jobs in the next 12 months."
    )
}


def prepare_source_excerpts() -> None:
    OUTPUTS.mkdir(parents=True, exist_ok=True)
    required = [
        SOURCE_MODEL_PAGE,
        SOURCE_APPENDIX_28,
        SOURCE_APPENDIX_29,
        SOURCE_TAM3_FIGURE,
        SOURCE_MIAO_PAGE,
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing source render(s): " + ", ".join(missing))

    Image.open(SOURCE_MODEL_PAGE).crop((105, 100, 1175, 1160)).save(MODEL_CROP)
    Image.open(SOURCE_APPENDIX_28).crop((110, 115, 1120, 1510)).save(APPENDIX_28_CROP)
    Image.open(SOURCE_APPENDIX_29).crop((110, 115, 1120, 1135)).save(APPENDIX_29_CROP)
    Image.open(SOURCE_TAM3_FIGURE).convert("RGB").save(TAM3_FIGURE, quality=95)
    Image.open(SOURCE_MIAO_PAGE).crop((245, 175, 1125, 668)).save(MIAO_FIGURE, quality=95)

    width, height = 1800, 820
    canvas = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(canvas)
    regular = base.load_font("regular", 29)
    bold = base.load_font("bold", 35)
    small_bold = base.load_font("bold", 27)
    small = base.load_font("regular", 23)

    boxes = {
        "PEOU": (55, 85, 575, 250, "Perceived Ease of Use", "#EAF2F8", "#365F91"),
        "REL": (55, 425, 575, 590, "Job Application\nRelevance", "#EAF2F8", "#365F91"),
        "PU": (705, 255, 1235, 455, "Perceived Usefulness", "#E2F0D9", "#548235"),
        "SN": (705, 620, 1235, 785, "Subjective Norm", "#EAF2F8", "#365F91"),
        "BI": (1430, 345, 1750, 575, "Behavioural\nIntention to\nUse DB", "#FFF2CC", "#BF9000"),
    }

    def draw_box(spec: tuple[int, int, int, int, str, str, str]) -> None:
        x1, y1, x2, y2, label, fill, outline = spec
        draw.rounded_rectangle((x1, y1, x2, y2), radius=16, fill=fill, outline=outline, width=4)
        lines = label.split("\n")
        line_height = 45
        y = (y1 + y2 - line_height * len(lines)) / 2
        for line in lines:
            bounds = draw.textbbox((0, 0), line, font=bold)
            x = (x1 + x2 - (bounds[2] - bounds[0])) / 2
            draw.text((x, y), line, fill="#111111", font=bold)
            y += line_height

    def arrow(start: tuple[int, int], end: tuple[int, int]) -> None:
        draw.line((start, end), fill="#1F4E79", width=7)
        sx, sy = start
        ex, ey = end
        dx, dy = ex - sx, ey - sy
        length = (dx * dx + dy * dy) ** 0.5
        ux, uy = dx / length, dy / length
        px, py = -uy, ux
        base_x, base_y = ex - 30 * ux, ey - 30 * uy
        draw.polygon(
            [
                (ex, ey),
                (base_x + 15 * px, base_y + 15 * py),
                (base_x - 15 * px, base_y - 15 * py),
            ],
            fill="#1F4E79",
        )

    for spec in boxes.values():
        draw_box(spec)

    paths = [
        ((575, 168), (705, 320), "H1", 615, 205),
        ((575, 508), (705, 390), "H2", 615, 440),
        ((1235, 355), (1430, 445), "H3", 1310, 365),
        ((1235, 702), (1430, 520), "H4", 1310, 585),
    ]
    for start, end, label, label_x, label_y in paths:
        arrow(start, end)
        draw.rounded_rectangle((label_x - 8, label_y - 5, label_x + 62, label_y + 34), radius=6, fill="white")
        draw.text((label_x, label_y), label, fill="#1F4E79", font=small_bold)

    for role, center_x in [
        ("Independent variables", 315),
        ("Intermediate variable", 970),
        ("Dependent variable", 1590),
    ]:
        role_bounds = draw.textbbox((0, 0), role, font=regular)
        draw.text((center_x - (role_bounds[2] - role_bounds[0]) / 2, 20), role, fill="#555555", font=regular)
    canvas.save(AUTHOR_FRAMEWORK, quality=95)


def configure_document(doc: Document, running_label: str) -> None:
    base.configure_document(doc)
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.paragraph_format.first_line_indent = None
    header.text = ""


def title_page(doc: Document, title: str, subtitle: str, label: str) -> None:
    for line in [
        "ASSUMPTION UNIVERSITY OF THAILAND",
        "Graduate School of Business and Advanced Technology Management",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = True
        r.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(15)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BY\nMR. MIN THIHA OO")
    r.bold = True
    r.font.size = Pt(12)

    for _ in range(2):
        doc.add_paragraph()
    for line in [
        "ITM7000 INDEPENDENT STUDY",
        "Submitted in Partial Fulfillment of the Requirements for the Degree of",
        "Master of Information Technology Management",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(11.5)
        if line == "ITM7000 INDEPENDENT STUDY":
            r.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("AUGUST 2026").font.size = Pt(12)
    doc.add_page_break()


def add_body(doc: Document, text: str, *, no_indent: bool = False) -> None:
    base.add_body(doc, text, no_indent=no_indent)


def add_bullets(doc: Document, entries: list[str]) -> None:
    for entry in entries:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(" " + entry)


def add_numbered(doc: Document, entries: list[str]) -> None:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "decimal")
    level.append(num_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    alignment = OxmlElement("w:lvlJc")
    alignment.set(qn("w:val"), "left")
    level.append(alignment)
    paragraph_properties = OxmlElement("w:pPr")
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "576")
    indentation.set(qn("w:hanging"), "288")
    paragraph_properties.append(indentation)
    level.append(paragraph_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    number.append(abstract_ref)
    numbering.append(number)

    for entry in entries:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.space_after = Pt(2)
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        level_ref = OxmlElement("w:ilvl")
        level_ref.set(qn("w:val"), "0")
        num_ref = OxmlElement("w:numId")
        num_ref.set(qn("w:val"), str(num_id))
        num_pr.extend([level_ref, num_ref])
        p_pr.append(num_pr)
        p.add_run(entry)


def add_note(doc: Document, text: str) -> None:
    base.add_note(doc, text)
    row = doc.tables[-1].rows[0]
    base.prevent_row_split(row)
    base.repeat_header(row)


def fix_table_geometry(table, widths: list[float]) -> None:
    total_dxa = round(sum(widths) * 1440)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)

    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def add_table(
    doc: Document,
    rows: list[list[str]],
    widths: list[float],
    *,
    caption: str,
    font_size: float = 9.0,
) -> None:
    if abs(sum(widths) - 6.2) > 0.01:
        raise ValueError(f"Table widths must total 6.2 inches: {widths}")
    base.add_table(doc, rows, widths, caption=caption, font_size=font_size)
    fix_table_geometry(doc.tables[-1], widths)


def add_figure(
    doc: Document,
    path: Path,
    caption: str,
    *,
    width: float = 5.9,
    alt_text: str | None = None,
) -> None:
    base.add_figure(doc, path, caption, width=width)
    if alt_text:
        doc_pr = doc.inline_shapes[-1]._inline.docPr
        doc_pr.set("title", "Conceptual framework")
        doc_pr.set("descr", alt_text)


def add_source_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GRAY


def add_hypotheses(doc: Document, hypotheses=HYPOTHESES) -> None:
    for code, statement in hypotheses:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.hanging_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{code}: ")
        r.bold = True
        p.add_run(statement)


def add_references(
    doc: Document,
    *,
    references: list[str] | None = None,
    font_size: float = 10.5,
    space_after: float = 3,
) -> None:
    reference_list = REFERENCES if references is None else references
    for reference in reference_list:
        p = doc.add_paragraph(reference)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(space_after)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.size = Pt(font_size)


def narrative_source_list(source_text: str) -> str:
    sources = source_text.split("; ")
    if len(sources) == 1:
        return sources[0]
    if len(sources) == 2:
        return " and ".join(sources)
    return ", ".join(sources[:-1]) + ", and " + sources[-1]


def questionnaire_item_text(code: str, source_text: str) -> str:
    return FIELD_ITEM_CORRECTIONS.get(code, source_text)


def add_questionnaire(doc: Document, *, include_sources: bool = True) -> None:
    scale_rows = [
        ["1", "2", "3", "4", "5"],
        [
            "Strongly disagree",
            "Disagree",
            "Neutral",
            "Agree",
            "Strongly agree",
        ],
    ]
    add_table(
        doc,
        scale_rows,
        [1.24, 1.24, 1.24, 1.24, 1.24],
        caption="Table 4.2: Five-point agreement scale",
        font_size=8.0,
    )

    grouped: dict[str, list[tuple[str, str, str]]] = {}
    for code, construct, source_text, prior_source in STUDY_ITEMS:
        grouped.setdefault(construct, []).append((code, source_text, prior_source))

    for table_number, (construct, group) in enumerate(grouped.items(), start=3):
        heading = "Behavioural Intention" if construct == "Behavioral Intention" else construct
        doc.add_heading(heading, level=3)
        rows = [["Code", "Questionnaire statement", "Measurement source"]]
        for code, source_text, prior_source in group:
            source = f"Steenkamp et al. (2024); {prior_source}" if include_sources else ""
            rows.append([code, questionnaire_item_text(code, source_text), source])
        add_table(
            doc,
            rows,
            [0.55, 4.1, 1.55],
            caption=f"Table 4.{table_number}: {heading} measurement items",
            font_size=8.0,
        )


def add_toc_page(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    for label, page in [
        ("CHAPTER 1: INTRODUCTION", 4),
        ("CHAPTER 2: LITERATURE REVIEW", 7),
        ("CHAPTER 3: RESEARCH FRAMEWORK", 11),
        ("CHAPTER 4: RESEARCH METHODOLOGY", 15),
        ("References", 22),
    ]:
        base.add_contents_entry(doc, label, page, bold=True)
    doc.add_page_break()


def build_master_full_replication_legacy() -> None:
    doc = Document()
    configure_document(doc, "Digital badges for IT micro-credentials | Thai IT students")
    title_page(doc, TITLE, SUBTITLE, "Master's Thesis")

    doc.add_heading("Abstract", level=1)
    add_body(
        doc,
        "University-issued digital badges can provide students with a verifiable way to present learning "
        "achievements from IT micro-credentials in job applications. Their practical value, however, "
        "depends partly on whether students intend to use them. This study examines the factors associated "
        "with Thai IT students' intentions to use university-issued digital badges when applying for jobs. "
        "The conceptual framework and questionnaire are adopted from the extended Technology Acceptance "
        "Model developed for digital badges by Steenkamp, Fisher, and Nesbit (2024). The model contains 11 "
        "constructs and 13 hypothesised relationships involving perceived usefulness, perceived ease of use, "
        "social influence, job relevance, facilitating conditions, computer-related beliefs, and behavioural "
        "intention. A quantitative cross-sectional survey will be conducted with students aged 18 years or "
        "older who are enrolled in IT-related university programmes in Thailand. Following a pilot study of "
        "40 eligible students, at least 384 usable responses will be collected for the main study. The "
        "measurement and structural models will be assessed using partial least squares structural equation "
        "modelling. The findings are expected to extend evidence on student adoption of digital credentials "
        "and inform Thai universities considering IT micro-credential badge initiatives.",
        no_indent=True,
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.add_run("Keywords: ").bold = True
    p.add_run(
        "digital badges, micro-credentials, technology acceptance model, behavioural intention, "
        "Thai IT students"
    )
    doc.add_page_break()
    add_toc_page(doc)

    doc.add_heading("CHAPTER 1: INTRODUCTION", level=1)
    doc.add_heading("1.1 Background of the Study", level=2)
    add_body(
        doc,
        "Micro-credentials certify learning outcomes acquired through a comparatively small volume of "
        "learning and are intended to support flexible learning and employability (Council of the European "
        "Union, 2022; UNESCO, 2022). A digital badge is a visual and potentially verifiable representation "
        "of an achievement. It can contain information about the issuer, recipient, assessment criteria, "
        "and learning outcome (Kiiskila et al., 2023; Steenkamp et al., 2024). Although the terms are related, "
        "the micro-credential is the record of learning, whereas the badge is one way of representing and "
        "sharing that record.",
    )
    add_body(
        doc,
        "For university students, a badge may provide an additional way to communicate skills in online "
        "profiles and job-application materials. The availability of a badge does not necessarily mean that "
        "students will use it. Adoption may depend on whether students regard the badge as useful and easy to "
        "use, whether important people and institutions support its use, whether it is relevant to job "
        "applications, and whether students have the confidence and resources needed to use it (Steenkamp et "
        "al., 2024).",
    )
    add_body(
        doc,
        "The Technology Acceptance Model (TAM) explains technology use through beliefs about usefulness and "
        "ease of use (Davis, 1989). Later extensions identify social influence, task relevance, perceived "
        "control, computer self-efficacy, computer anxiety, and computer playfulness as additional determinants "
        "of acceptance (Venkatesh, 2000; Venkatesh & Bala, 2008). Steenkamp et al. (2024) applied these "
        "relationships specifically to university-issued digital badges used in job applications. The present "
        "study applies that established model to IT students in Thailand.",
    )

    doc.add_heading("1.2 Statement of the Problem", level=2)
    add_body(
        doc,
        "Universities may invest in IT micro-credentials and digital-badge systems without knowing whether "
        "students will use the badges when seeking employment. Existing research has examined digital-badge "
        "adoption among accounting students in New Zealand (Steenkamp et al., 2024) and technology use for "
        "micro-credential programmes in other national settings (Miao et al., 2024). Evidence remains limited "
        "for students in IT-related programmes in Thailand. Without such evidence, Thai universities have "
        "little empirical guidance on the beliefs and conditions most closely associated with students' "
        "intention to use university-issued badges in job applications. This study addresses that gap by "
        "testing an established digital-badge acceptance model in the Thai IT-student context.",
    )

    doc.add_heading("1.3 Research Objectives", level=2)
    add_numbered(
        doc,
        [
            "To assess Thai IT students' intention to use university-issued digital badges when applying for jobs.",
            "To examine the effects of perceived usefulness, perceived ease of use, and subjective norm on behavioural intention.",
            "To examine the effects of perceived ease of use, subjective norm, image, job application relevance, and result demonstrability on perceived usefulness.",
            "To examine the effects of computer self-efficacy, perceptions of external control, computer anxiety, and computer playfulness on perceived ease of use.",
        ],
    )

    doc.add_heading("1.4 Research Questions", level=2)
    add_numbered(
        doc,
        [
            "To what extent do Thai IT students intend to use university-issued digital badges to showcase employability skills in job applications?",
            "How do perceived usefulness, perceived ease of use, and subjective norm influence behavioural intention?",
            "How do perceived ease of use, subjective norm, image, job application relevance, and result demonstrability influence perceived usefulness?",
            "How do computer self-efficacy, perceptions of external control, computer anxiety, and computer playfulness influence perceived ease of use?",
        ],
    )

    doc.add_heading("1.5 Scope of the Research", level=2)
    add_bullets(
        doc,
        [
            "Population: university students aged 18 or above who are enrolled in an IT-related program in Thailand.",
            "Focal object: a university-issued digital badge representing completion of an IT micro-credential and usable in job applications.",
            "Variables: 11 constructs from the extended Technology Acceptance Model applied by Steenkamp et al. (2024).",
            "Instrument: 40 construct items measured on a seven-point Likert agreement scale.",
            "Method: a pilot study followed by a cross-sectional online survey and PLS-SEM analysis.",
        ],
    )

    doc.add_heading("1.6 Research Limitations", level=2)
    add_body(
        doc,
        "The cross-sectional, self-report design can estimate associations and predictive paths but cannot "
        "prove actual badge use, employer recognition, employment, or causal effects. Non-probability "
        "recruitment may limit representativeness. Students may also have limited prior experience with "
        "digital badges; therefore a neutral definition will precede the construct items. Translation can "
        "introduce semantic differences, which will be managed through forward translation, independent "
        "back-translation, and the required pilot.",
    )

    doc.add_heading("1.7 Significance of the Study", level=2)
    add_body(
        doc,
        "The study extends the application of TAM3-related constructs to digital credentials in a Thai "
        "higher-education setting. It also provides practical evidence about the factors associated with "
        "students' willingness to use badges issued for IT micro-credentials. The findings may assist "
        "universities in planning badge communication, student support, training, and integration with "
        "career-development activities. The outcome measured is behavioural intention rather than actual "
        "employment or employer evaluation.",
    )

    doc.add_heading("1.8 Definition of Key Terms", level=2)
    definition_rows = [["Term", "Working definition"]]
    definition_rows += [
        ["Micro-credential", "A record of learning outcomes acquired through a small volume of learning (Council of the European Union, 2022; UNESCO, 2022)."],
        ["University-issued digital badge", "A verifiable visual token issued by a university to represent an achievement; in this study, completion of an IT micro-credential."],
        ["Perceived usefulness", OPERATIONAL_DEFINITIONS["PU"] + " (Davis, 1989)."],
        ["Perceived ease of use", OPERATIONAL_DEFINITIONS["PEOU"] + " (Davis, 1989)."],
        ["Behavioural intention", OPERATIONAL_DEFINITIONS["BI"] + " (Venkatesh & Davis, 2000)."],
    ]
    add_table(doc, definition_rows, [1.65, 4.55], caption="Table 1.1: Working definitions", font_size=9.4)

    doc.add_page_break()
    doc.add_heading("CHAPTER 2: LITERATURE REVIEW", level=1)
    doc.add_heading("2.1 Applied Theory: Technology Acceptance Model and TAM3", level=2)
    add_body(
        doc,
        "The Technology Acceptance Model proposes that perceived usefulness and perceived ease of use are "
        "central beliefs underlying a person's intention to use a technology (Davis, 1989). Perceived ease of "
        "use may also influence perceived usefulness because a technology that requires less effort can be "
        "more useful in practice (Venkatesh & Davis, 2000). Behavioural intention is consequently treated as "
        "the immediate outcome of the acceptance process.",
    )
    add_body(
        doc,
        "TAM3 extends the original model by explaining how social influence, task characteristics, control "
        "beliefs, intrinsic motivation, and anxiety contribute to usefulness and ease-of-use perceptions "
        "(Venkatesh, 2000; Venkatesh & Bala, 2008). Steenkamp et al. (2024) applied this extended structure to "
        "students' use of university-issued digital badges in job applications. Their model provides the "
        "theoretical framework for the present study because it addresses the same technology, decision, and "
        "intended-use outcome.",
    )

    doc.add_heading("2.2 Digital Badges and IT Micro-Credentials", level=2)
    add_body(
        doc,
        "Micro-credentials are short, assessed learning experiences that document specific learning outcomes "
        "(Council of the European Union, 2022; UNESCO, 2022). Digital badges can represent such achievements "
        "in a portable and verifiable format. Student perceptions of verifiable digital credentials include "
        "their usefulness for displaying competence, their credibility, and the ease with which they can be "
        "shared (Kiiskila et al., 2023). Research on micro-credential programmes also indicates that technology "
        "acceptance provides a suitable basis for studying student intention (Miao et al., 2024). In this study, "
        "the focal technology is a university-issued badge representing completion of an IT micro-credential."
    )

    doc.add_heading("2.3 Variable Definitions", level=2)
    for index, (name, code, _, _, _) in enumerate(CONSTRUCTS, start=1):
        display_name = "Behavioural Intention" if code == "BI" else name
        doc.add_heading(f"2.3.{index} {display_name}", level=3)
        add_body(
            doc,
            f"{display_name} is defined in this study as {OPERATIONAL_DEFINITIONS[code][0].lower() + OPERATIONAL_DEFINITIONS[code][1:]} "
            f"This definition follows {narrative_source_list(CONSTRUCT_THEORY_SOURCES[code])}.",
        )

    doc.add_heading("2.4 Relationships Between Variables", level=2)
    doc.add_heading("2.4.1 Determinants of Behavioural Intention", level=3)
    add_body(
        doc,
        "TAM proposes that perceived usefulness and perceived ease of use influence behavioural intention "
        "(Davis, 1989). Subjective norm can also influence intention when important social actors encourage "
        "technology use (Venkatesh & Davis, 2000). In the digital-badge study by Steenkamp et al. (2024), "
        "perceived usefulness and subjective norm had significant positive effects on behavioural intention, "
        "whereas the direct effect of perceived ease of use was not significant. All three theoretically "
        "specified paths are examined in the Thai context.",
    )

    doc.add_heading("2.4.2 Determinants of Perceived Usefulness", level=3)
    add_body(
        doc,
        "Perceived ease of use can increase perceived usefulness when lower effort improves the practical "
        "value of a system (Venkatesh & Davis, 2000). TAM3 further proposes that subjective norm and image "
        "represent social-influence processes, while job relevance and result demonstrability represent "
        "cognitive instrumental processes (Venkatesh & Bala, 2008). Steenkamp et al. (2024) found significant "
        "positive effects for perceived ease of use and job application relevance on perceived usefulness, "
        "and a significant effect of subjective norm on image. The remaining proposed relationships were not "
        "significant in their sample and are retested among Thai IT students.",
    )

    doc.add_heading("2.4.3 Determinants of Perceived Ease of Use", level=3)
    add_body(
        doc,
        "Early perceptions of ease of use may be anchored in computer self-efficacy, perceptions of external "
        "control, computer anxiety, and computer playfulness (Venkatesh, 2000; Venkatesh & Bala, 2008). "
        "Self-efficacy, external control, and playfulness are expected to increase perceived ease of use, "
        "whereas anxiety is expected to reduce it. Steenkamp et al. (2024) found a significant positive effect "
        "only for perceptions of external control. Testing all four relationships permits assessment of whether "
        "the pattern differs among Thai IT students.",
    )

    doc.add_heading("2.5 Previous Studies", level=2)
    prior_rows = [
        ["Study", "Context and method", "Main finding", "Relevance"],
        ["Davis (1989)", "Technology acceptance; scale development and validation.", "Perceived usefulness and perceived ease of use are central determinants of user acceptance.", "Establishes the core TAM constructs."],
        ["Venkatesh and Bala (2008)", "Longitudinal research integrating determinants of usefulness and ease of use.", "TAM3 explains acceptance through social influence, cognitive processes, control, motivation, and emotion.", "Provides the extended theoretical structure and many measurement scales."],
        ["Kiiskila et al. (2023)", "Higher-education students; qualitative study of verifiable digital credentials.", "Students identified value in verifiable and shareable credentials while also raising implementation concerns.", "Supports the digital-credential and student context."],
        ["Miao et al. (2024)", "University students; quantitative PLS-SEM study of technology use for micro-credential programmes.", "A TAM-based model was used to explain intention in a micro-credential setting.", "Supports the use of technology-acceptance theory for micro-credentials."],
        ["Steenkamp et al. (2024)", "New Zealand accounting students; online survey; extended TAM3; PLS-SEM; n = 57.", "Six of 13 hypothesised relationships were supported, including usefulness and subjective norm as predictors of intention.", "Provides the conceptual framework, hypotheses, and questionnaire used in this study."],
    ]
    add_table(doc, prior_rows, [1.2, 1.75, 1.75, 1.5], caption="Table 2.1: Summary of previous studies", font_size=7.8)

    doc.add_page_break()
    doc.add_heading("CHAPTER 3: RESEARCH FRAMEWORK", level=1)
    doc.add_heading("3.1 Theoretical Framework", level=2)
    add_body(
        doc,
        "This study is grounded in TAM and TAM3. The core TAM relationships connect perceived usefulness and "
        "perceived ease of use with behavioural intention (Davis, 1989). TAM3 explains these beliefs through "
        "social influence, cognitive instrumental processes, control beliefs, intrinsic motivation, and "
        "computer anxiety (Venkatesh, 2000; Venkatesh & Bala, 2008). The combined framework is appropriate "
        "because the decision under investigation is whether students intend to use a digital technology for "
        "a specific task: presenting university-issued badges in job applications.",
    )
    add_body(
        doc,
        "Within this framework, perceived usefulness and perceived ease of use are the principal acceptance "
        "beliefs. Subjective norm and image represent social-influence processes, while job application "
        "relevance and result demonstrability represent cognitive judgements about the value of the badge for "
        "the task. Computer self-efficacy, external control, anxiety, and playfulness explain students' initial "
        "ease-of-use perceptions (Venkatesh, 2000; Venkatesh & Bala, 2008).",
    )
    add_body(
        doc,
        "Behavioural intention is used as the dependent variable because the proposed badge is not assumed to "
        "be available to every respondent at the time of data collection. The framework therefore explains "
        "students' stated intention to use a badge, rather than actual badge use or employment outcomes."
    )

    doc.add_heading("3.2 Conceptual Framework", level=2)
    add_body(
        doc,
        "The conceptual framework is adopted from the digital-badge acceptance model developed and tested by "
        "Steenkamp et al. (2024). It includes the same 11 constructs and 13 directional relationships. The "
        "population examined in the present study is Thai students enrolled in IT-related university "
        "programmes, and the badge represents completion of an IT micro-credential.",
    )
    add_figure(
        doc,
        MODEL_CROP,
        "Figure 3.1: Conceptual framework for students' intention to use digital badges",
        width=5.75,
        alt_text=(
            "Conceptual framework adopted from Steenkamp et al. (2024), showing 11 constructs and "
            "13 directional relationships predicting perceived usefulness, perceived ease of use, "
            "image, and behavioural intention to use digital badges."
        ),
    )
    add_source_note(
        doc,
        "Source: Steenkamp, Fisher, and Nesbit (2024, Figure 1, p. 10).",
    )
    add_body(
        doc,
        "Perceived usefulness, perceived ease of use, and subjective norm are positioned as predictors of "
        "behavioural intention. Perceived ease of use, subjective norm, image, job application relevance, and "
        "result demonstrability are positioned as predictors of perceived usefulness. Computer self-efficacy, "
        "perceptions of external control, computer anxiety, and computer playfulness are positioned as "
        "predictors of perceived ease of use. Subjective norm is also linked to image. In Figure 3.1, DB "
        "denotes digital badges.",
    )

    doc.add_heading("3.3 Research Hypotheses", level=2)
    add_body(
        doc,
        "Based on the theoretical relationships described above and the digital-badge model of Steenkamp et "
        "al. (2024, pp. 11-13), the following hypotheses are proposed:",
    )
    add_hypotheses(doc, THESIS_HYPOTHESES)

    heading = doc.add_heading("3.4 Operationalization", level=2)
    heading.paragraph_format.page_break_before = True
    op_rows = [["Variable", "Operational definition", "Items", "Scale", "Source/reference"]]
    for name, code, _, codes, _ in CONSTRUCTS:
        display_name = "Behavioural Intention" if code == "BI" else name
        op_rows.append(
            [
                display_name,
                OPERATIONAL_DEFINITIONS[code],
                codes,
                "7-point Likert",
                CONSTRUCT_THEORY_SOURCES[code],
            ]
        )
    add_table(
        doc,
        op_rows,
        [1.15, 2.35, 0.75, 0.75, 1.2],
        caption="Table 3.1: Operationalization of the study variables",
        font_size=7.4,
    )
    add_source_note(
        doc,
        "The complete questionnaire statements and their measurement sources are presented in Section 4.3.",
    )

    doc.add_page_break()
    doc.add_heading("CHAPTER 4: RESEARCH METHODOLOGY", level=1)
    doc.add_heading("4.1 Research Design", level=2)
    add_body(
        doc,
        "The study will use a quantitative, explanatory, cross-sectional survey design. This design is "
        "appropriate for measuring students' perceptions and testing the hypothesised relationships among "
        "the 11 latent variables at one point in time. The construct indicators will be specified as "
        "reflective, consistent with Steenkamp et al. (2024), and the model will be estimated using PLS-SEM.",
    )

    doc.add_heading("4.2 Sampling Procedure", level=2)
    doc.add_heading("4.2.1 Target Population", level=3)
    add_body(
        doc,
        "The target population consists of students aged 18 years or older who are currently enrolled in an "
        "IT-related undergraduate or postgraduate programme at a university in Thailand. Relevant fields may "
        "include information technology, computer science, software engineering, information systems, data "
        "science, cybersecurity, and closely related programmes.",
    )
    doc.add_heading("4.2.2 Sampling Method and Sample Size", level=3)
    add_body(
        doc,
        "Non-probability convenience and snowball sampling will be used because a complete national sampling "
        "frame of Thai IT students is not available. The online questionnaire will be distributed through "
        "IT-related university programmes, student groups, and academic networks. The Krejcie and Morgan "
        "(1970) table identifies 384 as a conventional sample-size benchmark for a large population under "
        "probability-sampling assumptions. Because the present study uses non-probability sampling, this "
        "benchmark does not imply a 5% margin of error for the achieved sample. It is used as a conservative "
        "target that also exceeds the minimum normally required for a PLS-SEM model of this complexity (Hair "
        "et al., 2022). The main study will therefore require at least 384 usable responses, with a recruitment "
        "target of 400 to allow for incomplete or ineligible responses.",
    )
    doc.add_heading("4.2.3 Inclusion and Exclusion Criteria", level=3)
    add_body(
        doc,
        "Respondents will be included if they provide informed consent, are at least 18 years old, are "
        "currently enrolled in an IT-related university programme in Thailand, and complete the construct "
        "measures. Responses will be excluded for failed eligibility screening, duplicate submission, "
        "substantial missing data, or implausibly short completion time. Pilot participants will not be "
        "included in the main-study dataset.",
    )

    doc.add_heading("4.3 Research Instrument and Questionnaire Design", level=2)
    add_body(
        doc,
        "Data will be collected using a structured online questionnaire. The conceptual measures and 40 "
        "construct items are adopted from Steenkamp et al. (2024, Appendix, pp. 28-29), who drew the measures "
        "from established TAM, TAM3, UTAUT, and computer self-efficacy scales. The focal referent throughout "
        "the construct section is a university-issued digital badge representing completion of an IT "
        "micro-credential and used in a job application.",
    )
    add_body(
        doc,
        "BI3 contains a grammatical correction from 'use the when' in the published Appendix to 'use them "
        "when'; the construct meaning, referent, and 12-month timeframe are unchanged.",
    )
    instrument_rows = [
        ["Section", "Content", "Purpose"],
        ["A", "Participant information, consent, age, enrolment status, and field of study", "Confirm consent and eligibility"],
        ["B", "Neutral definition of an IT micro-credential and a university-issued digital badge", "Provide a common referent"],
        ["C", "Age group, gender, study level, field, year of study, institution type, and prior badge awareness or use", "Describe the sample"],
        ["D", "Forty statements measuring the 11 model constructs", "Test the measurement and structural models"],
    ]
    add_table(doc, instrument_rows, [0.65, 3.65, 1.9], caption="Table 4.1: Structure of the questionnaire", font_size=8.6)
    add_body(
        doc,
        "Construct items will use a seven-point Likert agreement scale ranging from 1 (strongly disagree) to "
        "7 (strongly agree). The full English questionnaire statements and their published measurement sources "
        "are presented below.",
    )
    add_questionnaire(doc)

    doc.add_heading("4.4 Translation and Pilot Study", level=2)
    add_body(
        doc,
        "The questionnaire will be translated from English into Thai and independently back-translated into "
        "English. Differences will be reconciled to preserve the intended meaning of each construct (Brislin, "
        "1970). Before the main survey, the Thai questionnaire will be pilot tested with 40 eligible Thai IT "
        "students. A sample of this size is appropriate for an initial assessment of item comprehension and "
        "scale performance (Hertzog, 2008; Johanson & Brooks, 2010). The pilot will assess eligibility logic, "
        "clarity of instructions and items, completion time, missing responses, and preliminary internal "
        "consistency. Pilot participants will be excluded from the main survey.",
    )

    doc.add_heading("4.5 Data Collection", level=2)
    add_body(
        doc,
        "Following academic and ethical approval, an online survey link will be distributed through the "
        "identified university and student channels. The first page will provide the participant information "
        "and consent statement. Participation will be voluntary, and respondents may leave the survey before "
        "submission. No directly identifying information will be required. Responses will be stored securely "
        "and used only for the purposes described in the participant information.",
    )

    doc.add_heading("4.6 Statistical Treatment of Data", level=2)
    add_numbered(
        doc,
        [
            "Screen records for consent, eligibility, duplicates, missingness, completion quality, and coding errors; reverse-code negatively phrased indicators before construct analysis.",
            "Describe respondent characteristics using frequencies and percentages, and summarize construct responses using means and standard deviations.",
            "Assess the reflective measurement model using indicator loadings, Cronbach's alpha, composite reliability, average variance extracted, and discriminant validity. All 40 indicators will be administered; any indicator removal will be based on the Thai-sample results and theoretical content and will be reported.",
            "Assess structural-model collinearity using variance inflation factors and evaluate the explanatory power of endogenous constructs using R-squared values.",
            "Estimate the 13 hypothesised paths in SmartPLS using 5,000 bootstrap samples and report path coefficients, confidence intervals, p-values, and effect sizes (Hair et al., 2022).",
            "Determine support for each hypothesis at the 5% significance level and compare the resulting pattern with the findings of Steenkamp et al. (2024).",
        ],
    )

    doc.add_heading("4.7 Ethical Considerations", level=2)
    add_body(
        doc,
        "The study will follow the university's research-ethics requirements. Participants will receive "
        "information about the study purpose, eligibility requirements, voluntary participation, "
        "confidentiality, data use, and withdrawal before submission. Only respondents who provide informed "
        "consent will proceed. Results will be reported in aggregate form.",
    )

    doc.add_page_break()
    doc.add_heading("References", level=1)
    add_references(doc)

    doc.save(MASTER_OUT)


def build_evidence_pack_full_replication_legacy() -> None:
    doc = Document()
    configure_document(doc, "Evidence pack | direct source audit")
    title_page(
        doc,
        "Published Framework, Hypotheses, Variables, and Questionnaire Evidence",
        "Direct-model replication source pack",
        "Evidence Pack",
    )

    doc.add_heading("1. Decision", level=1)
    add_note(
        doc,
        "Pass under the strict source-fidelity rule. One published article supplies the actual model, "
        "all 13 hypotheses, all 11 constructs, the 40 coded construct items, the seven-point scale, and "
        "the PLS-SEM procedure. The Thai population and IT micro-credential badge context are disclosed "
        "replication changes.",
    )
    add_body(
        doc,
        "Primary source: Steenkamp, N., Fisher, R., and Nesbit, T. (2024), Understanding "
        "accounting students' intentions to use digital badges to showcase employability skills, "
        "Accounting Education, 33(6), 906-934. DOI: 10.1080/09639284.2023.2276200.",
    )

    doc.add_heading("2. Strict Source Audit", level=1)
    audit_rows = [
        ["Required component", "Published location", "Status", "Current use"],
        ["Actual framework", "Figure 1, article p. 10", "Direct", "Complete figure reproduced; no redrawing or path selection."],
        ["Variables", "Figure 1, Appendix, Table 2", "Direct", "All 11 source constructs and codes retained."],
        ["Hypotheses", "H1-H13, article pp. 11-13", "Direct", "All paths and directions reproduced."],
        ["Questionnaire", "Appendix, article pp. 28-29", "Direct", "All 40 coded source items transcribed."],
        ["Scale", "Appendix note, article p. 29", "Direct", "Seven-point agreement scale retained."],
        ["Analysis", "Method/results, article pp. 13-17", "Direct-method guide", "Reflective PLS-SEM and 5,000 bootstrap samples."],
    ]
    add_table(doc, audit_rows, [1.25, 1.55, 0.85, 2.55], caption="Table 2.1: Pass/fail audit", font_size=8.3)

    doc.add_page_break()
    doc.add_heading("3. Actual Published Framework", level=1)
    add_body(
        doc,
        "The image below is a direct crop from the source PDF. It is not a reconstruction. The source "
        "boxes, labels, arrows, and caption are unchanged.",
    )
    add_figure(
        doc,
        MODEL_CROP,
        "Figure 3.1: Direct source reproduction of Steenkamp et al. (2024, Figure 1, article p. 10).",
        width=5.8,
    )

    doc.add_heading("4. Exact Published Hypotheses", level=1)
    add_body(
        doc,
        "The statements below reproduce the source hypotheses. H11's grammar is retained to avoid "
        "silently rewriting the published statement.",
    )
    add_hypotheses(doc)

    doc.add_heading("5. Exact Variable and Item-Count Match", level=1)
    count_rows = [["Published construct", "Code", "Published codes", "Count", "Thesis match"]]
    count_rows += [[name, code, codes, count, "Exact"] for name, code, _, codes, count in CONSTRUCTS]
    count_rows.append(["Total", "", "", "40", "Exact"])
    add_table(doc, count_rows, [1.8, 0.55, 1.55, 0.55, 1.75], caption="Table 5.1: Construct and item-count audit", font_size=8.4)

    doc.add_heading("6. Actual Published Questionnaire Appendix", level=1)
    add_body(
        doc,
        "The next two images are direct crops of the source Appendix pages. They show the item codes, "
        "wording, prior sources identified by Steenkamp et al., and seven-point response scale.",
    )
    add_figure(
        doc,
        APPENDIX_28_CROP,
        "Figure 6.1: Direct source reproduction of the first construct-item Appendix page, Steenkamp et al. (2024, article p. 28).",
        width=5.35,
    )
    doc.add_page_break()
    add_figure(
        doc,
        APPENDIX_29_CROP,
        "Figure 6.2: Direct source reproduction of the second construct-item Appendix page, Steenkamp et al. (2024, article p. 29).",
        width=6.0,
    )

    doc.add_page_break()
    doc.add_heading("7. Item-Level Provenance", level=1)
    add_body(
        doc,
        "All 40 coded statements below are verbatim from the source Appendix, including BI3's apparent "
        "typo; any correction requires approval before fielding.",
    )
    provenance_rows = [["Code", "Construct", "Exact source wording", "Master status"]]
    for code, construct, source_text, _ in ITEMS:
        status = "Exact"
        if code == "BI3":
            status = "Exact; typo noted"
        provenance_rows.append([code, construct, source_text, status])
    add_table(doc, provenance_rows, [0.55, 1.35, 3.55, 0.75], caption="Table 7.1: Forty-item direct provenance matrix", font_size=7.6)

    doc.add_heading("8. Source Anomalies and Decisions", level=1)
    anomaly_rows = [
        ["Source detail", "Evidence", "Decision"],
        ["Unnumbered employer sentence after SN4", "The methods state 40 items; the Appendix has 40 coded items; Table 2 reports SN1-SN4 only.", "Do not invent SN5 or score the unnumbered sentence."],
        ["PEC4, CPLAY4, RES4 removed in source analysis", "Source measurement-model section reports outer loadings below .40.", "Administer all 40 source items; make any Thai-sample deletion from Thai data and report it."],
        ["BI3 source typo", "Appendix prints 'use the when'.", "Preserve source wording in evidence; seek approval for the field correction 'them'."],
        ["Behavioural/Behavioral label variation", "Figure and Appendix use different English spelling/label length.", "Retain the figure label in the model and the Appendix label in the item evidence; treat both as BI."],
    ]
    add_table(doc, anomaly_rows, [1.35, 2.55, 2.3], caption="Table 8.1: No-hidden-decisions register", font_size=8.4)

    doc.add_heading("9. What Changes and What Does Not", level=1)
    change_rows = [
        ["Element", "Source study", "Thai study", "Classification"],
        ["Population", "Accounting students at one New Zealand university", "Students in IT-related programs in Thailand", "Disclosed context change"],
        ["Badge context", "University-issued badges for employability skills", "University-issued badges representing IT micro-credentials for employability skills", "Disclosed context specification"],
        ["Model, constructs, paths", "Figure 1; 11 constructs; 13 paths", "Same complete model", "No change"],
        ["Construct items", "40 coded Appendix items", "Same English source master", "No conceptual change"],
        ["Response scale", "Seven-point agreement", "Same", "No change"],
        ["Language", "English", "English audit master plus approved Thai translation", "Necessary cross-language procedure"],
        ["Pilot", "Not reported as a separate pilot", "Required local pilot of 40", "Local process requirement"],
    ]
    add_table(doc, change_rows, [1.15, 1.85, 2.05, 1.15], caption="Table 9.1: Replication boundary", font_size=8.2)

    doc.add_heading("10. References", level=1)
    add_references(doc, font_size=9, space_after=1)
    doc.save(EVIDENCE_OUT)


def add_toc_page_author(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    for label, page in [
        ("CHAPTER 1: INTRODUCTION", 4),
        ("CHAPTER 2: LITERATURE REVIEW", 8),
        ("CHAPTER 3: RESEARCH FRAMEWORK", 13),
        ("CHAPTER 4: RESEARCH METHODOLOGY", 18),
        ("References", 23),
    ]:
        base.add_contents_entry(doc, label, page, bold=True)
    doc.add_page_break()


def build_master() -> None:
    doc = Document()
    configure_document(doc, "")
    title_page(doc, TITLE, SUBTITLE, "Independent Study")

    doc.add_heading("Abstract", level=1)
    add_body(
        doc,
        "University-issued digital badges provide students with a portable way to present learning "
        "achievements from IT micro-credentials in job applications. Their value depends partly on whether "
        "students regard them as useful and intend to use them. This proposed study examines four published "
        "relationships explaining Thai IT students' behavioural intention to use university-issued digital "
        "badges in job applications. The conceptual framework draws on the Technology Acceptance Model, "
        "Technology Acceptance Model 3, research on technology use for micro-credential programmes, and the "
        "digital-badge framework tested by Steenkamp, Fisher, and Nesbit (2024). The framework contains five "
        "constructs: Perceived Ease of Use, Job Application Relevance, Perceived Usefulness, Subjective Norm, "
        "and Behavioural Intention to Use DB. A quantitative cross-sectional questionnaire survey will be "
        "conducted with students aged 18 years or older who are enrolled in IT-related university programmes "
        "in Thailand. The questionnaire retains 18 measurement statements from Steenkamp et al. (2024) and "
        "uses a five-point Likert agreement scale. After a pilot test with 40 eligible students, at least "
        "384 usable responses will be collected. Data will be analysed using descriptive statistics, "
        "Cronbach's alpha, collinearity diagnostics, and two multiple linear regression models. The study is "
        "intended to identify the beliefs most closely associated with students' intended use of digital "
        "badges and to inform university micro-credential and career-support practices.",
        no_indent=True,
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.add_run("Keywords: ").bold = True
    p.add_run(
        "digital badges, micro-credentials, perceived usefulness, subjective norm, behavioural intention, "
        "Thai IT students"
    )
    doc.add_page_break()
    add_toc_page_author(doc)

    doc.add_heading("CHAPTER 1: INTRODUCTION", level=1)
    doc.add_heading("1.1 Background of the Study", level=2)
    add_body(
        doc,
        "Micro-credentials certify learning outcomes acquired through a comparatively small volume of "
        "learning and are intended to support flexible learning, upskilling, and employability (Council of "
        "the European Union, 2022; UNESCO, 2022). A digital badge is a visual and potentially verifiable "
        "representation of an achievement. It may contain information about the issuer, recipient, learning "
        "outcomes, and assessment criteria (Kiiskila et al., 2023; Steenkamp et al., 2024). In the present "
        "study, the micro-credential is the learning achievement and the university-issued digital badge is "
        "the electronic representation that a student can store and present.",
    )
    add_body(
        doc,
        "Digital badges may be especially relevant to students in IT-related programmes because technical "
        "skills and tools change rapidly and students may complete focused learning in areas such as data "
        "analytics, cybersecurity, cloud computing, and software development. A badge can give students an "
        "additional way to present this learning in online profiles and job applications. The availability "
        "of a badge, however, does not ensure that students will use it. Intended use may depend on whether "
        "the badge is easy to use, relevant to applying for jobs, useful for strengthening an application, "
        "and supported by people and institutions that matter to the student.",
    )
    add_body(
        doc,
        "The Technology Acceptance Model (TAM) explains technology use through Perceived Usefulness and "
        "Perceived Ease of Use (Davis, 1989). TAM2 and TAM3 added social influence and task-relevance "
        "processes, including Subjective Norm and Job Relevance (Venkatesh & Davis, 2000; Venkatesh & Bala, "
        "2008). These models have been applied widely in educational technology research (Abdullah & Ward, "
        "2016; Granic & Marangunic, 2019). Miao et al. (2024) used TAM-related variables to study students' "
        "technology acceptance for micro-credential programmes, while Steenkamp et al. (2024) applied TAM3 "
        "constructs directly to university-issued digital badges used in job applications.",
    )
    add_body(
        doc,
        "The present study uses these published foundations to examine a smaller model suited to a master's "
        "thesis. It retains exact construct names and four relationships shown and tested by Steenkamp et al. "
        "(2024). Perceived Ease of Use and Job Application Relevance are proposed to influence Perceived "
        "Usefulness. Perceived Usefulness and Subjective Norm are proposed to influence Behavioural Intention "
        "to Use DB. The model is examined among Thai students in IT-related university programmes.",
    )

    doc.add_heading("1.2 Statement of the Problem", level=2)
    add_body(
        doc,
        "Universities may introduce IT micro-credentials and digital-badge systems without knowing whether "
        "students will use the badges when seeking employment. If students do not see the badges as relevant "
        "to job applications, useful for presenting their skills, easy to use, or supported by important "
        "people and their university, the badges may remain unused even when the underlying learning is "
        "valuable. Universities therefore need evidence about the beliefs associated with students' intended "
        "use before investing further in badge communication, platforms, and career-support activities.",
    )
    add_body(
        doc,
        "Existing studies provide useful but geographically limited evidence. Steenkamp et al. (2024) "
        "examined accounting students at one New Zealand university, while Miao et al. (2024) examined "
        "technology use for micro-credential programmes among private-university students in Pakistan. "
        "Evidence remains limited for Thai university students in IT-related programmes and for the specific "
        "use of university-issued badges in job applications. This study addresses that contextual gap by "
        "testing a focused set of previously published relationships and measures among Thai IT students."
    )

    doc.add_heading("1.3 Research Objectives", level=2)
    add_numbered(
        doc,
        [
            "To examine the effect of Perceived Ease of Use on Perceived Usefulness of university-issued digital badges among Thai IT students.",
            "To examine the effect of Job Application Relevance on Perceived Usefulness of university-issued digital badges among Thai IT students.",
            "To examine the effect of Perceived Usefulness on Behavioural Intention to Use DB among Thai IT students.",
            "To examine the effect of Subjective Norm on Behavioural Intention to Use DB among Thai IT students.",
        ],
    )

    doc.add_heading("1.4 Research Questions", level=2)
    add_numbered(
        doc,
        [
            "How does Perceived Ease of Use affect Perceived Usefulness of university-issued digital badges among Thai IT students?",
            "How does Job Application Relevance affect Perceived Usefulness of university-issued digital badges among Thai IT students?",
            "How does Perceived Usefulness affect Behavioural Intention to Use DB among Thai IT students?",
            "How does Subjective Norm affect Behavioural Intention to Use DB among Thai IT students?",
        ],
    )

    doc.add_heading("1.5 Scope of the Research", level=2)
    add_body(
        doc,
        "The study focuses on students aged 18 years or older who are enrolled in undergraduate or "
        "postgraduate IT-related programmes at universities in Thailand. The focal object is a "
        "university-issued digital badge representing successful completion of an IT micro-credential and "
        "capable of being presented in a job application. The model contains five constructs and four "
        "hypothesised relationships. Data will be collected using an 18-item, five-point Likert-scale online "
        "questionnaire. Descriptive statistics, Cronbach's alpha, variance inflation factors, and two multiple "
        "linear regression models will be used for analysis. The study concerns students' perceptions and intended "
        "use; it does not measure employers' evaluations, actual employment outcomes, or actual badge use."
    )

    doc.add_heading("1.6 Research Limitations", level=2)
    add_body(
        doc,
        "First, the cross-sectional design measures variables at one point in time and cannot establish "
        "causality. Second, convenience and snowball recruitment may limit the representativeness of the "
        "sample. Third, self-reported intention may not lead to actual badge use in future job applications. "
        "Fourth, some respondents may have limited experience with digital badges; a neutral definition will "
        "therefore be provided before the measurement items. These limits will be considered when the results "
        "are interpreted."
    )

    doc.add_heading("1.7 Significance of the Study", level=2)
    add_body(
        doc,
        "The study contributes to technology-acceptance and digital-credential research by testing a focused "
        "set of published TAM3 relationships in a Thai higher-education and IT-student setting. It also "
        "examines whether relationships supported in Steenkamp et al.'s (2024) New Zealand digital-badge "
        "study are observed in a different discipline and national context. The study does not propose new "
        "constructs; its contribution is the contextual test and the focused combination of established "
        "constructs in one model.",
    )
    add_body(
        doc,
        "The findings may help universities identify where practical support is most useful. Evidence about "
        "Perceived Ease of Use can guide platform instructions and student support. Evidence about Job "
        "Application Relevance and Perceived Usefulness can guide how badges are connected to employability "
        "and career services. Evidence about Subjective Norm can guide communication by lecturers, career "
        "advisers, and universities. Students may benefit from clearer information about how badges can be "
        "stored and presented, while credential providers and policymakers may gain evidence about the "
        "conditions associated with learner acceptance."
    )

    doc.add_heading("1.8 Definition of Key Terms", level=2)
    definition_rows = [
        ["Term", "Definition used in this study"],
        ["Behavioural Intention to Use DB", OPERATIONAL_DEFINITIONS["BI"] + " (Steenkamp et al., 2024)."],
        ["Digital badge", "A visual and potentially verifiable electronic representation of an achievement, including information about the issuer, recipient, criteria, and learning outcome (Kiiskila et al., 2023; Steenkamp et al., 2024)."],
        ["Job Application Relevance", OPERATIONAL_DEFINITIONS["REL"] + " (Steenkamp et al., 2024)."],
        ["Micro-credential", "A record of learning outcomes acquired through a small volume of learning (Council of the European Union, 2022; UNESCO, 2022)."],
        ["Perceived Ease of Use", OPERATIONAL_DEFINITIONS["PEOU"] + " (Davis, 1989; Steenkamp et al., 2024)."],
        ["Perceived Usefulness", OPERATIONAL_DEFINITIONS["PU"] + " (Davis, 1989; Steenkamp et al., 2024)."],
        ["Subjective Norm", OPERATIONAL_DEFINITIONS["SN"] + " (Venkatesh & Davis, 2000; Steenkamp et al., 2024)."],
        ["University-issued digital badge", "A digital badge issued by a university to represent successful completion of an IT micro-credential and capable of being presented in a job application."],
    ]
    add_table(doc, definition_rows, [1.65, 4.55], caption="Table 1.1: Definitions of key terms", font_size=8.9)

    doc.add_page_break()
    doc.add_heading("CHAPTER 2: LITERATURE REVIEW", level=1)
    doc.add_heading("2.1 Applied Theory", level=2)
    doc.add_heading("2.1.1 Technology Acceptance Model", level=3)
    add_body(
        doc,
        "Davis (1989) developed the Technology Acceptance Model to explain why individuals accept or reject "
        "information technology. TAM proposed that Perceived Usefulness and Perceived Ease of Use shape "
        "attitudes and intentions toward use. Perceived Usefulness concerns expected performance benefits, "
        "whereas Perceived Ease of Use concerns expected effort. The model was tested through scale "
        "development and empirical studies of computer systems and became one of the most widely applied "
        "models of technology acceptance. Reviews show that TAM and its extensions have continued to explain "
        "the use of learning systems, e-portfolios, and other educational technologies (Abdullah & Ward, "
        "2016; Granic & Marangunic, 2019)."
    )
    doc.add_heading("2.1.2 TAM2 and TAM3", level=3)
    add_body(
        doc,
        "Venkatesh and Davis (2000) extended TAM to TAM2 by adding social-influence processes and cognitive "
        "instrumental processes. Subjective Norm represented perceived expectations from important others, "
        "while Job Relevance concerned the applicability of a system to the user's task. Venkatesh and Bala "
        "(2008) later combined TAM2 with determinants of Perceived Ease of Use to produce TAM3. TAM3 retained "
        "the central paths to Behavioural Intention and organised the antecedents of Perceived Usefulness and "
        "Perceived Ease of Use in one integrated framework. The model was examined using longitudinal data "
        "from organisational technology implementations and was intended to provide more practical guidance "
        "than the original, more general TAM."
    )
    doc.add_heading("2.1.3 Suitability for the Present Study", level=3)
    add_body(
        doc,
        "TAM3 is suitable because the present study concerns students' intention to use a digital tool for a "
        "specific task. The selected constructs cover the practical effort of using a badge, its relevance and "
        "usefulness for job applications, and the influence of important people and the university. Miao et "
        "al. (2024) demonstrated the relevance of TAM variables to technology used for micro-credential "
        "programmes. More directly, Steenkamp et al. (2024) applied TAM3 constructs and questionnaire items "
        "to students' use of university-issued digital badges for employability. These studies provide both "
        "the theoretical and same-field basis for the present framework."
    )

    doc.add_heading("2.2 Micro-Credentials and University-Issued Digital Badges", level=2)
    add_body(
        doc,
        "Micro-credentials are smaller units of assessed learning that can support flexible education and "
        "lifelong learning (Council of the European Union, 2022; UNESCO, 2022). Their digital representation "
        "can make achievement information more portable and verifiable. Kiiskila et al. (2023) found that "
        "students and university administrators associated digital credentials with the ability to present "
        "skills, maintain records of learning, and share evidence with employers, while also identifying "
        "implementation and recognition concerns. These findings indicate that the value of a digital "
        "credential depends partly on how learners expect to use it."
    )
    add_body(
        doc,
        "In the present study, the badge is limited to one clear purpose: representing completion of an IT "
        "micro-credential and being available for use in a job application. This common referent reduces the "
        "risk that respondents imagine unrelated badge types. The study measures perceptions of the badge "
        "and intention to use it rather than claiming that the badge is already recognised by employers or "
        "that it produces employment outcomes."
    )

    doc.add_heading("2.3 Variable Definitions", level=2)
    variable_order = ["PEOU", "REL", "PU", "SN", "BI"]
    variable_names = {
        "PEOU": "Perceived Ease of Use",
        "REL": "Job Application Relevance",
        "PU": "Perceived Usefulness",
        "SN": "Subjective Norm",
        "BI": "Behavioural Intention to Use DB",
    }
    for index, code in enumerate(variable_order, start=1):
        doc.add_heading(f"2.3.{index} {variable_names[code]}", level=3)
        add_body(doc, VARIABLE_LITERATURE[code])

    doc.add_heading("2.4 Relationships Between Variables", level=2)
    for heading_text, paragraph_text in RELATIONSHIP_LITERATURE:
        doc.add_heading(heading_text, level=3)
        add_body(doc, paragraph_text)

    doc.add_heading("2.5 Summary of Previous Studies", level=2)
    previous_rows = [
        ["Study", "Context and method", "Framework or finding", "Use in this study"],
        ["Davis (1989)", "Technology-acceptance scale development and empirical validation.", "Established Perceived Usefulness, Perceived Ease of Use, and intention.", "Core theory and original PU/PEOU measures."],
        ["Venkatesh and Davis (2000)", "Four longitudinal field studies of technology adoption.", "TAM2 added Subjective Norm and Job Relevance as determinants of acceptance.", "Theory for SN, job relevance, PU, and BI relationships."],
        ["Venkatesh and Bala (2008)", "Longitudinal organisational studies; integrated TAM3.", "Combined the determinants of usefulness and ease of use in one framework.", "Foundational framework and original scales."],
        ["Abdullah and Ward (2016)", "Meta-analytic model of e-learning acceptance factors.", "Confirmed recurring external determinants of PU and PEOU in education.", "Educational-technology relationship evidence."],
        ["Kiiskila et al. (2023)", "Interviews with students and university administrators about verifiable credentials.", "Identified value and use-intention factors for digital credentials.", "Micro-credential and digital-credential context."],
        ["Miao et al. (2024)", "Survey of 474 private-university students in Pakistan; PLS-SEM.", "Used TAM and self-determination variables to explain technology acceptance for micro-credential programmes.", "Same-field framework and student context."],
        ["Steenkamp et al. (2024)", "Online survey of 57 accounting students in New Zealand; PLS-SEM.", "Tested 11 constructs and 13 paths for using university-issued digital badges in job applications.", "Exact labels, four selected paths, and all 18 questionnaire items."],
    ]
    add_table(doc, previous_rows, [1.15, 1.75, 1.75, 1.55], caption="Table 2.1: Summary of previous studies", font_size=7.3)

    doc.add_page_break()
    doc.add_heading("CHAPTER 3: RESEARCH FRAMEWORK", level=1)
    doc.add_heading("3.1 Theoretical Frameworks from Previous Studies", level=2)
    add_body(
        doc,
        "The conceptual framework was developed from three published frameworks. The first provides the "
        "general TAM3 structure, the second applies technology-acceptance variables to micro-credential "
        "programmes, and the third applies the variables and questionnaire directly to university-issued "
        "digital badges. The original published diagrams are presented before the framework used in this study."
    )

    doc.add_heading("3.1.1 Technology Acceptance Model 3", level=3)
    add_body(
        doc,
        "Venkatesh and Bala (2008) combined TAM2 and the determinants of Perceived Ease of Use to develop "
        "TAM3. They tested the integrated model using longitudinal survey observations from organisational "
        "technology implementations. The model included Subjective Norm and Job Relevance as determinants "
        "of Perceived Usefulness and retained Perceived Usefulness and Behavioural Intention as central "
        "acceptance constructs. Figure 3.1 reproduces the authors' published TAM3 model."
    )
    add_figure(
        doc,
        TAM3_FIGURE,
        "Figure 3.1: Technology Acceptance Model 3",
        width=4.45,
        alt_text="Published TAM3 framework showing determinants of perceived usefulness, perceived ease of use, behavioural intention, and use behaviour.",
    )
    add_source_note(doc, "Source: Venkatesh and Bala (2008, Figure 2).")

    doc.add_heading("3.1.2 Technology Use for Micro-Credential Programmes", level=3)
    add_body(
        doc,
        "Miao et al. (2024) combined TAM and self-determination theory to study technology used for "
        "micro-credential programmes. They collected questionnaire data from 474 private-university students "
        "in Karachi using convenience sampling and analysed the model with PLS-SEM. Perceived Usefulness "
        "significantly influenced Attitude towards using technology for Micro-Credential Programs. Perceived "
        "Ease of Use was reported as supported at p = .051, and attitude significantly influenced Behavioral "
        "Intention. The proposed labour-market "
        "moderation was not significant. Figure 3.2 reproduces their published conceptual framework."
    )
    add_figure(
        doc,
        MIAO_FIGURE,
        "Figure 3.2: Conceptual framework for technology use in micro-credential programmes",
        width=5.85,
        alt_text="Published Miao et al. framework linking TAM and self-determination variables to attitude and behavioral intention for micro-credential programs.",
    )
    add_source_note(doc, "Source: Miao et al. (2024, Figure 1, p. 956).")

    doc.add_heading("3.1.3 Digital-Badge Acceptance Framework", level=3)
    add_body(
        doc,
        "Steenkamp et al. (2024) developed an extended TAM3 framework for students' intention to use "
        "university-issued digital badges in job applications. They collected online questionnaire data from "
        "57 accounting students at a New Zealand university and analysed 11 reflective constructs with "
        "PLS-SEM. Six of 13 hypothesised paths were supported. The four paths selected for the present study "
        "were all significant: Perceived Ease of Use to Perceived Usefulness, Job Application Relevance to "
        "Perceived Usefulness, Perceived Usefulness to Behavioural Intention to Use DB, and Subjective Norm "
        "to Behavioural Intention to Use DB. Figure 3.3 is a direct reproduction of their published model."
    )
    add_figure(
        doc,
        MODEL_CROP,
        "Figure 3.3: Research model for students' intention to use digital badges",
        width=5.25,
        alt_text="Published Steenkamp et al. digital-badge framework containing 11 constructs and 13 directional relationships.",
    )
    add_source_note(doc, "Source: Steenkamp, Fisher, and Nesbit (2024, Figure 1, p. 10).")

    doc.add_heading("3.2 Conceptual Framework of the Study", level=2)
    add_body(
        doc,
        "The conceptual framework in Figure 3.4 was constructed for this study by selecting five exact "
        "constructs and four previously tested paths from Steenkamp et al. (2024). TAM3 provides the "
        "underlying theoretical structure, and Miao et al. (2024) provides additional evidence that Perceived "
        "Ease of Use, Perceived Usefulness, and Behavioral Intention are relevant in a micro-credential "
        "technology context. No construct in the final framework has been renamed from Steenkamp et al.'s "
        "published digital-badge framework."
    )
    add_figure(
        doc,
        AUTHOR_FRAMEWORK,
        "Figure 3.4: Conceptual framework of the study",
        width=6.0,
        alt_text="Author-constructed framework in which perceived ease of use and job application relevance predict perceived usefulness, while perceived usefulness and subjective norm predict behavioural intention to use digital badges.",
    )
    add_source_note(
        doc,
        "Source: Constructed by the author based on Venkatesh and Bala (2008), Miao et al. (2024), and Steenkamp et al. (2024).",
    )
    framework_rows = [
        ["Study construct", "Exact-label source", "Selected published path", "Additional framework support"],
        ["Perceived Ease of Use", "Steenkamp et al. (2024), Figure 1", "PEOU -> PU (source H3)", "Venkatesh & Bala (2008); Miao et al. (2024)"],
        ["Job Application Relevance", "Steenkamp et al. (2024), Figure 1", "REL -> PU (source H8)", "Job Relevance in Venkatesh & Bala (2008)"],
        ["Perceived Usefulness", "Steenkamp et al. (2024), Figure 1", "PU -> BI (source H1)", "Venkatesh & Bala (2008); Miao et al. (2024)"],
        ["Subjective Norm", "Steenkamp et al. (2024), Figure 1", "SN -> BI (source H4)", "Venkatesh & Bala (2008)"],
        ["Behavioural Intention to Use DB", "Steenkamp et al. (2024), Figure 1", "Dependent construct", "Behavioral Intention in Venkatesh & Bala (2008) and Miao et al. (2024)"],
    ]
    add_table(doc, framework_rows, [1.35, 1.5, 1.35, 2.0], caption="Table 3.1: Framework provenance", font_size=7.2)

    doc.add_heading("3.3 Research Hypotheses", level=2)
    add_body(
        doc,
        "The following hypotheses correspond directly to the four arrows in Figure 3.4. The numbering is "
        "specific to the present study; Table 3.1 identifies the original Steenkamp et al. (2024) hypothesis "
        "number for each selected path."
    )
    add_hypotheses(doc, THESIS_HYPOTHESES)

    heading = doc.add_heading("3.4 Operationalization of the Variables", level=2)
    heading.paragraph_format.page_break_before = True
    code_to_definition = {code: OPERATIONAL_DEFINITIONS[code] for code in SELECTED_CODES}
    code_to_name = {
        "PEOU": "Perceived Ease of Use",
        "REL": "Job Application Relevance",
        "PU": "Perceived Usefulness",
        "SN": "Subjective Norm",
        "BI": "Behavioural Intention to Use DB",
    }
    op_rows = [["Variable and operational definition", "Code", "Measurement statement", "Scale", "Source"]]
    seen_codes: set[str] = set()
    for item_code, _, source_text, _ in STUDY_ITEMS:
        construct_code = item_code.rstrip("0123456789")
        variable_cell = ""
        if construct_code not in seen_codes:
            variable_cell = f"{code_to_name[construct_code]}: {code_to_definition[construct_code]}"
            seen_codes.add(construct_code)
        op_rows.append(
            [
                variable_cell,
                item_code,
                questionnaire_item_text(item_code, source_text),
                "5-point Likert",
                "Steenkamp et al. (2024)",
            ]
        )
    add_table(
        doc,
        op_rows,
        [1.55, 0.45, 2.65, 0.6, 0.95],
        caption="Table 3.2: Operationalization and full measurement items",
        font_size=7.0,
    )
    add_source_note(
        doc,
        "BI3 corrects the source Appendix's apparent typographical phrase 'use the when' to 'use them when'; no construct meaning or timeframe is changed.",
    )

    doc.add_page_break()
    doc.add_heading("CHAPTER 4: RESEARCH METHODOLOGY", level=1)
    doc.add_heading("4.1 Research Design", level=2)
    add_body(
        doc,
        "The study will use a quantitative, explanatory, cross-sectional survey design. This design is "
        "appropriate for measuring student perceptions and estimating the four hypothesised relationships at "
        "one point in time. The five constructs will be measured using the published multi-item statements "
        "selected from Steenkamp et al. (2024). Construct scores will be formed after internal-consistency "
        "testing, and the four hypotheses will be examined using two multiple linear regression models. The design can "
        "identify associations and predictive relationships but cannot establish causal effects."
    )

    doc.add_heading("4.2 Sampling Procedure", level=2)
    doc.add_heading("4.2.1 Target Population", level=3)
    add_body(
        doc,
        "The target population consists of students aged 18 years or older who are currently enrolled in an "
        "IT-related undergraduate or postgraduate programme at a university in Thailand. Relevant fields "
        "include information technology, computer science, software engineering, information systems, data "
        "science, cybersecurity, computer engineering, and closely related programmes."
    )
    doc.add_heading("4.2.2 Sampling Method and Sample Size", level=3)
    add_body(
        doc,
        "Convenience and snowball sampling will be used because a complete national sampling frame of Thai IT "
        "students is not available. The questionnaire will be distributed through university programmes, "
        "student groups, and academic networks. The Krejcie and Morgan (1970) table provides 384 as a "
        "conventional benchmark for a large population under probability-sampling assumptions. The present "
        "non-probability sample cannot claim the corresponding margin of error, but 384 is retained as a "
        "conservative minimum for the descriptive and regression analyses. The recruitment target will be "
        "400 usable responses to allow for incomplete or ineligible records."
    )
    doc.add_heading("4.2.3 Inclusion and Exclusion Criteria", level=3)
    add_body(
        doc,
        "Respondents will be included if they provide informed consent, are at least 18 years old, are "
        "currently enrolled in an IT-related university programme in Thailand, and complete the construct "
        "measures. Records will be excluded for failed eligibility screening, duplicate submission, "
        "substantial missing data, or implausibly short completion time. Pilot participants will not be "
        "included in the main-study dataset."
    )

    doc.add_heading("4.3 Research Instrument and Questionnaire Design", level=2)
    add_body(
        doc,
        "Data will be collected through a structured online questionnaire. The 18 construct items are the "
        "items printed by Steenkamp et al. (2024, Appendix, pp. 28-29) for the five exact constructs retained "
        "in the conceptual framework. Their wording already refers to university-issued digital badges and "
        "job applications, so no contextual replacement is required. BI3 contains only the grammatical "
        "correction described in Section 3.4."
    )
    instrument_rows = [
        ["Section", "Content", "Purpose"],
        ["A", "Participant information, consent, age, enrolment status, and field", "Confirm consent and eligibility"],
        ["B", "Neutral definition of an IT micro-credential and university-issued digital badge", "Provide a common referent"],
        ["C", "Age group, gender, study level, field, year, institution type, and prior badge awareness or use", "Describe the sample"],
        ["D", "Eighteen statements measuring five model constructs", "Assess reliability and test the hypotheses"],
    ]
    add_table(doc, instrument_rows, [0.65, 3.65, 1.9], caption="Table 4.1: Structure of the questionnaire", font_size=8.5)
    add_body(
        doc,
        "Steenkamp et al. (2024) administered the source statements with seven response categories. The "
        "present study retains the 18 measurement statements but uses a five-point Likert agreement scale "
        "ranging from 1 (strongly disagree) to 5 (strongly agree). This is a documented response-format "
        "adaptation; the construct definitions, item assignments, and item wording remain unchanged except "
        "for the disclosed BI3 grammatical correction. Five- and seven-point formats are both established "
        "options for Likert-type measurement (Dawes, 2008). The complete English instrument is presented below."
    )
    add_questionnaire(doc)

    doc.add_heading("4.4 Translation and Pilot Testing", level=2)
    add_body(
        doc,
        "The approved English questionnaire will be translated into Thai and independently back-translated "
        "into English. Differences will be reconciled to preserve each construct's intended meaning (Brislin, "
        "1970). Before the main survey, the Thai questionnaire will be pilot tested with 40 eligible Thai IT "
        "students, consistent with the programme examples and published pilot-study guidance (Hertzog, 2008; "
        "Johanson & Brooks, 2010). The pilot will assess screening logic, instructions, item comprehension, "
        "completion time, missing responses, and preliminary internal consistency. Pilot participants will be "
        "excluded from the main study."
    )

    doc.add_heading("4.5 Data Collection", level=2)
    add_body(
        doc,
        "After academic and ethical approval, an online survey link will be distributed through the selected "
        "university and student channels. The first page will provide participant information and request "
        "informed consent. Participation will be voluntary. Respondents may leave before submitting the "
        "questionnaire, and no directly identifying information will be required. Survey records will be "
        "stored securely and used only for the stated academic purpose."
    )

    doc.add_heading("4.6 Statistical Treatment of Data", level=2)
    add_body(
        doc,
        "After data collection, records will be screened for consent, eligibility, duplicate submissions, "
        "missing data, completion quality, and coding errors. Statistical analyses will be conducted using "
        "jamovi. Frequencies and percentages will describe respondent demographics and prior digital-badge "
        "experience. The mean and standard deviation will be reported for every measurement item. After the "
        "reliability assessment, a composite score for each construct will be calculated as the mean of its "
        "assigned items."
    )
    add_body(
        doc,
        "Internal consistency will be assessed separately for Perceived Ease of Use, Job Application "
        "Relevance, Perceived Usefulness, Subjective Norm, and Behavioural Intention to Use DB using "
        "Cronbach's alpha (Cronbach, 1951). Alpha values of .70 or higher will be treated as acceptable for "
        "the study, while item-total information and the theoretical coverage of each published scale will "
        "also be considered before any item is removed (Hair et al., 2019). Reliability will be reported for "
        "both the pilot sample and the final sample."
    )
    add_body(
        doc,
        "Two multiple linear regression models will test the four hypotheses. In the first model, Perceived "
        "Usefulness will be the dependent variable, with Perceived Ease of Use and Job Application Relevance "
        "as independent variables; this model will test H1 and H2. In the second model, Behavioural Intention "
        "to Use DB will be the dependent variable, with Perceived Usefulness and Subjective Norm as "
        "independent variables; this model will test H3 and H4. Variance inflation factors will be examined "
        "for collinearity, and residual diagnostics will be reviewed for the assumptions of multiple linear "
        "regression (Hair et al., 2019). Each model will report R-squared, adjusted R-squared, the overall "
        "F-test, unstandardized and standardized coefficients, standard errors, t-values, p-values, and "
        "confidence intervals. A hypothesis will be supported when its coefficient is positive and its "
        "p-value is below .05."
    )

    doc.add_heading("4.7 Ethical Considerations", level=2)
    add_body(
        doc,
        "The study will follow the university's research-ethics requirements. The consent page will explain "
        "the study purpose, eligibility requirements, voluntary participation, foreseeable risks, "
        "confidentiality, data use, and withdrawal before submission. Only respondents who provide consent "
        "will continue. Findings will be reported in aggregate form and will not be presented as evidence of "
        "actual employment outcomes or employer acceptance."
    )

    doc.add_page_break()
    doc.add_heading("References", level=1)
    add_references(doc, references=MASTER_REFERENCES)
    doc.save(MASTER_OUT)


def build_evidence_pack() -> None:
    doc = Document()
    configure_document(doc, "")
    title_page(
        doc,
        "Framework, Hypotheses, Variables, and Questionnaire Evidence",
        "Source Record for the Author-Constructed Digital-Badge Model",
        "Evidence Pack",
    )

    doc.add_heading("1. Purpose", level=1)
    add_body(
        doc,
        "This document records the published frameworks used to construct the study model, the direct source "
        "of every variable and relationship, and the original questionnaire source. It accompanies the "
        "chaptered thesis manuscript."
    )
    add_note(
        doc,
        "The study does not reproduce Steenkamp et al.'s complete model. It selects five exact constructs and "
        "four supported paths from that published model, adds no new construct, and uses only the 18 printed "
        "questionnaire items belonging to those constructs.",
    )

    doc.add_heading("2. Published Frameworks", level=1)
    doc.add_heading("2.1 Technology Acceptance Model 3", level=2)
    add_figure(
        doc,
        TAM3_FIGURE,
        "Figure 2.1: Technology Acceptance Model 3",
        width=4.5,
        alt_text="Published TAM3 framework showing determinants of perceived usefulness, perceived ease of use, behavioural intention, and use behaviour.",
    )
    add_source_note(doc, "Source: Venkatesh and Bala (2008, Figure 2).")
    doc.add_heading("2.2 Micro-Credential Technology-Acceptance Framework", level=2)
    add_figure(
        doc,
        MIAO_FIGURE,
        "Figure 2.2: Technology use for micro-credential programmes",
        width=5.85,
        alt_text="Published micro-credential framework combining technology acceptance and self-determination variables to explain attitude and behavioural intention.",
    )
    add_source_note(doc, "Source: Miao et al. (2024, Figure 1, p. 956).")
    doc.add_heading("2.3 Digital-Badge Acceptance Framework", level=2)
    add_figure(
        doc,
        MODEL_CROP,
        "Figure 2.3: Published digital-badge research model",
        width=4.5,
        alt_text="Published Steenkamp digital-badge framework containing eleven constructs and thirteen directional relationships.",
    )
    add_source_note(doc, "Source: Steenkamp, Fisher, and Nesbit (2024, Figure 1, p. 10).")

    doc.add_heading("3. Conceptual Framework Used in the Study", level=1)
    add_figure(
        doc,
        AUTHOR_FRAMEWORK,
        "Figure 3.1: Conceptual framework of the study",
        width=6.0,
        alt_text="Author-constructed framework in which perceived ease of use and job application relevance predict perceived usefulness, while perceived usefulness and subjective norm predict behavioural intention to use digital badges.",
    )
    add_source_note(
        doc,
        "Source: Constructed by the author based on Venkatesh and Bala (2008), Miao et al. (2024), and Steenkamp et al. (2024).",
    )
    path_rows = [["Study hypothesis", "Exact path", "Direct published source", "Published result"]] + SOURCE_PATHS
    add_table(doc, path_rows, [0.8, 2.2, 1.55, 1.65], caption="Table 3.1: Direct path provenance", font_size=8.0)
    add_hypotheses(doc, THESIS_HYPOTHESES)

    doc.add_heading("4. Exact Variable Provenance", level=1)
    variable_rows = [
        ["Study variable", "Exact source label", "Published framework location", "Questionnaire codes"],
        ["Perceived Ease of Use", "Perceived Ease of Use", "Steenkamp et al. (2024), Figure 1", "PEOU1-PEOU4"],
        ["Job Application Relevance", "Job Application Relevance", "Steenkamp et al. (2024), Figure 1", "REL1-REL3"],
        ["Perceived Usefulness", "Perceived Usefulness", "Steenkamp et al. (2024), Figure 1", "PU1-PU4"],
        ["Subjective Norm", "Subjective Norm", "Steenkamp et al. (2024), Figure 1", "SN1-SN4"],
        ["Behavioural Intention to Use DB", "Behavioural Intention to Use DB", "Steenkamp et al. (2024), Figure 1", "BI1-BI3"],
    ]
    add_table(doc, variable_rows, [1.55, 1.55, 2.1, 1.0], caption="Table 4.1: Exact-label audit", font_size=8.1)

    doc.add_heading("5. Published Questionnaire Evidence", level=1)
    add_body(
        doc,
        "Steenkamp et al. (2024, Appendix, pp. 28-29) printed all construct codes, questionnaire "
        "statements, prior measurement sources, and the seven-point agreement scale. The two images below are "
        "direct crops of those Appendix pages."
    )
    add_figure(
        doc,
        APPENDIX_28_CROP,
        "Figure 5.1: Published questionnaire Appendix, first page",
        width=5.3,
        alt_text="Direct crop of the first published Steenkamp questionnaire Appendix page, including construct codes, item statements, and adaptation sources.",
    )
    add_source_note(doc, "Source: Steenkamp et al. (2024, Appendix, p. 28).")
    doc.add_page_break()
    add_figure(
        doc,
        APPENDIX_29_CROP,
        "Figure 5.2: Published questionnaire Appendix, second page",
        width=5.9,
        alt_text="Direct crop of the second published Steenkamp questionnaire Appendix page, including job application relevance, behavioural intention, and the seven-point response scale.",
    )
    add_source_note(doc, "Source: Steenkamp et al. (2024, Appendix, p. 29).")

    item_rows = [["Code", "Exact construct", "Published wording", "Field wording/status"]]
    for code, construct, source_text, _ in STUDY_ITEMS:
        field_text = questionnaire_item_text(code, source_text)
        status = "Exact wording"
        if field_text != source_text:
            status = "Typographical correction only: " + field_text
        item_rows.append([code, construct, source_text, status])
    add_table(doc, item_rows, [0.55, 1.35, 2.8, 1.5], caption="Table 5.1: Eighteen-item provenance matrix", font_size=7.2)
    add_source_note(
        doc,
        "Response scale: 1 = Strongly disagree, 2 = Moderately disagree, 3 = Somewhat disagree, 4 = Neutral, 5 = Somewhat agree, 6 = Moderately agree, and 7 = Strongly agree.",
    )

    doc.add_heading("6. Requirement Check", level=1)
    requirement_rows = [
        ["Professor's requested component", "Where it is provided", "Status"],
        ["Actual frameworks borrowed from previous studies", "Figures 2.1-2.3 and thesis Section 3.1", "Provided as published source images"],
        ["Author's conceptual framework", "Figure 3.1 and thesis Figure 3.4", "Constructed from cited frameworks"],
        ["Variables with exact names", "Table 4.1 and thesis Table 3.1", "All exact Steenkamp labels"],
        ["Hypotheses supported by previous studies", "Table 3.1 and thesis Sections 2.4 and 3.3", "Four direct published paths"],
        ["Borrowed questionnaire", "Figures 5.1-5.2 and Table 5.1", "All 18 items directly traceable"],
        ["Operationalization", "Thesis Table 3.2", "Full definitions, item wording, scale, and source"],
        ["Pilot", "Thesis Section 4.4", "Forty eligible pilot participants"],
    ]
    add_table(doc, requirement_rows, [2.0, 2.75, 1.45], caption="Table 6.1: Supervisor-requirement check", font_size=8.1)

    doc.add_heading("7. References", level=1)
    add_references(doc, font_size=9.2, space_after=2)
    doc.save(EVIDENCE_OUT)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build the current thesis documents.")
    parser.add_argument(
        "--master-only",
        action="store_true",
        help="Build only the supervisor-facing thesis manuscript.",
    )
    args = parser.parse_args()

    prepare_source_excerpts()
    build_master()
    print(MASTER_OUT)
    if not args.master_only:
        build_evidence_pack()
        print(EVIDENCE_OUT)
    print(MODEL_CROP)
    print(APPENDIX_28_CROP)
    print(APPENDIX_29_CROP)


if __name__ == "__main__":
    main()
