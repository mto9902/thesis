from __future__ import annotations

from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = PROJECT_ROOT / "outputs"
SOURCE_DIR = PROJECT_ROOT / "source"

MASTER_OUT = OUTPUTS / "Master_Thesis_EVC_Digital_Credentials_Thai_IT_Students.docx"
EVIDENCE_OUT = OUTPUTS / "Evidence_Pack_Published_Framework_and_Questionnaire.docx"
SOURCE_FIGURE = OUTPUTS / "Published_Framework_Kiiskila_Figure_1_Direct_Crop.png"
CONCEPTUAL_FIGURE = OUTPUTS / "Conceptual_Framework_Exact_EVC_Variables.png"

SOURCE_PAGE = SOURCE_DIR / "kiiskila_page_366.png"

TITLE = (
    "Factors Influencing Intention to Use Digital Credentials for IT "
    "Micro-Credentials among Thai University Students"
)
SUBTITLE = "An Expectancy-Value-Cost Study"

FONT_CANDIDATES = {
    "regular": [
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "C:/Windows/Fonts/times.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
    ],
    "bold": [
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf",
        "C:/Windows/Fonts/timesbd.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
    ],
}

HEADER_FILL = "D9E2F3"
SUBHEADER_FILL = "E2F0D9"
NOTE_FILL = "FFF2CC"
LIGHT_FILL = "F2F2F2"


REFERENCES = [
    (
        "Ajzen, I. (1991). The theory of planned behavior. Organizational Behavior "
        "and Human Decision Processes, 50(2), 179-211. "
        "https://doi.org/10.1016/0749-5978(91)90020-T"
    ),
    (
        "Barron, K. E., & Hulleman, C. S. (2015). Expectancy-value-cost model of "
        "motivation. In J. D. Wright (Ed.), International encyclopedia of the social "
        "& behavioral sciences (2nd ed., Vol. 8, pp. 503-509). Elsevier. "
        "https://doi.org/10.1016/B978-0-08-097086-8.26099-6"
    ),
    (
        "Brislin, R. W. (1970). Back-translation for cross-cultural research. "
        "Journal of Cross-Cultural Psychology, 1(3), 185-216. "
        "https://doi.org/10.1177/135910457000100301"
    ),
    (
        "Council of the European Union. (2022). Council Recommendation of 16 June "
        "2022 on a European approach to micro-credentials for lifelong learning and "
        "employability (2022/C 243/02). Official Journal of the European Union, "
        "C 243, 10-25. https://eur-lex.europa.eu/eli/C/2022/243/oj"
    ),
    (
        "Durik, A. M., Vida, M., & Eccles, J. S. (2006). Task values and ability "
        "beliefs as predictors of high school literacy choices: A developmental "
        "analysis. Journal of Educational Psychology, 98(2), 382-393. "
        "https://doi.org/10.1037/0022-0663.98.2.382"
    ),
    (
        "Eccles, J. S. (2005). Subjective task value and the Eccles et al. model of "
        "achievement-related choices. In A. J. Elliot & C. S. Dweck (Eds.), "
        "Handbook of competence and motivation (pp. 105-121). Guilford Press."
    ),
    (
        "Eccles, J. S., Adler, T. F., Futterman, R., Goff, S. B., Kaczala, C. M., "
        "Meece, J. L., & Midgley, C. (1983). Expectancies, values, and academic "
        "behaviors. In J. T. Spence (Ed.), Achievement and achievement motives "
        "(pp. 75-146). W. H. Freeman."
    ),
    (
        "Eccles, J. S., & Wigfield, A. (2002). Motivational beliefs, values, and "
        "goals. Annual Review of Psychology, 53, 109-132. "
        "https://doi.org/10.1146/annurev.psych.53.100901.135153"
    ),
    (
        "Fishbein, M., & Ajzen, I. (1975). Belief, attitude, intention and behavior: "
        "An introduction to theory and research. Addison-Wesley."
    ),
    (
        "Flake, J. K., Barron, K. E., Hulleman, C. S., McCoach, D. B., & Welsh, "
        "M. E. (2015). Measuring cost: The forgotten component of expectancy-value "
        "theory. Contemporary Educational Psychology, 41, 232-244. "
        "https://doi.org/10.1016/j.cedpsych.2015.03.002"
    ),
    (
        "Hertzog, M. A. (2008). Considerations in determining sample size for pilot "
        "studies. Research in Nursing & Health, 31(2), 180-191. "
        "https://doi.org/10.1002/nur.20247"
    ),
    (
        "Hulleman, C. S., Barron, K. E., Kosovich, J. J., & Lazowski, R. A. (2016). "
        "Student motivation: Current theories, constructs, and interventions within "
        "an expectancy-value framework. In A. A. Lipnevich, F. Preckel, & R. D. "
        "Roberts (Eds.), Psychosocial skills and school systems in the 21st century "
        "(pp. 241-278). Springer. "
        "https://doi.org/10.1007/978-3-319-28606-8_10"
    ),
    (
        "Johanson, G. A., & Brooks, G. P. (2010). Initial scale development: Sample "
        "size for pilot studies. Educational and Psychological Measurement, 70(3), "
        "394-400. https://doi.org/10.1177/0013164409355692"
    ),
    (
        "Kato, S., Galan-Muros, V., & Weko, T. (2020). The emergence of alternative "
        "credentials (OECD Education Working Papers No. 216). OECD Publishing. "
        "https://doi.org/10.1787/b741f39e-en"
    ),
    (
        "Kiiskila, P., Kukkonen, A., & Pirkkalainen, H. (2023). Are micro-credentials "
        "valuable for students? Perspective on verifiable digital credentials. "
        "SN Computer Science, 4, Article 366. "
        "https://doi.org/10.1007/s42979-023-01797-y"
    ),
    (
        "Kosovich, J. J., Hulleman, C. S., Barron, K. E., & Getty, S. (2015). A "
        "practical measure of student motivation: Establishing validity evidence for "
        "the expectancy-value-cost scale in middle school. The Journal of Early "
        "Adolescence, 35(5-6), 790-816. "
        "https://doi.org/10.1177/0272431614556890"
    ),
    (
        "Krejcie, R. V., & Morgan, D. W. (1970). Determining sample size for "
        "research activities. Educational and Psychological Measurement, 30(3), "
        "607-610. https://doi.org/10.1177/001316447003000308"
    ),
    (
        "Miao, M., Ahmed, F., Ahsan, M., & Qamar, B. (2024). Intention to use "
        "technology for micro-credential programs: Evidence from technology "
        "acceptance and self-determination model. International Journal of "
        "Educational Management, 38(4), 948-977. "
        "https://doi.org/10.1108/IJEM-02-2023-0066"
    ),
    (
        "Ministry of Higher Education, Science, Research and Innovation. (2022). "
        "Micro-credentials in Thai higher education: Opportunities, challenges and "
        "outlooks from ASEAN and Europe. "
        "https://www.mhesi.go.th/index.php/en/news-and-announce-all/news-all/"
        "executive-ps-news/7614-micro-credentials-in-thai-higher-education-"
        "opportunities-challenges-and-outlooks-from-asean-and-europe.html"
    ),
    (
        "Ministry of Higher Education, Science, Research and Innovation. (2023). "
        "Regional policy seminar on micro-credentials in ASEAN higher education. "
        "https://www.mhesi.go.th/index.php/en/news-and-announce-all/news-all/"
        "executive-ps-news/9589-unesco-bangkok-korean-educational-development-"
        "institute-kedi-micro-credentials.html"
    ),
    (
        "Oliver, B. (2021). Micro-credentials: A learner value framework: "
        "Provocation. Journal of Teaching and Learning for Graduate Employability, "
        "12(1), 48-51. https://doi.org/10.21153/jtlge2021vol12no1art1456"
    ),
    (
        "Perez, T., Wormington, S. V., Barger, M. M., Schwartz-Bloom, R. D., Lee, "
        "Y. K., & Linnenbrink-Garcia, L. (2019). Science expectancy, value, and cost "
        "profiles and their proximal and distal relations to undergraduate science, "
        "technology, engineering, and math persistence. Science Education, 103(2), "
        "264-286. https://doi.org/10.1002/sce.21490"
    ),
    (
        "Shang, S., & Lyv, W. (2022). Understanding the impact of quality elements "
        "on MOOCs continuance intention. Education and Information Technologies, "
        "27, 10949-10976. https://doi.org/10.1007/s10639-022-11063-y"
    ),
    (
        "UNESCO. (2022). Towards a common definition of micro-credentials. UNESCO. "
        "https://unesdoc.unesco.org/ark:/48223/pf0000381668"
    ),
    (
        "Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User "
        "acceptance of information technology: Toward a unified view. MIS Quarterly, "
        "27(3), 425-478. https://doi.org/10.2307/30036540"
    ),
    (
        "Wigfield, A., & Eccles, J. S. (2000). Expectancy-value theory of "
        "achievement motivation. Contemporary Educational Psychology, 25(1), 68-81. "
        "https://doi.org/10.1006/ceps.1999.1015"
    ),
]


EVC_ITEMS = [
    (
        "EOS1",
        "Expectation of success in using digital credentials",
        "I know I can learn how to use digital credentials issued for IT micro-credentials.",
        "Kosovich et al. (2015), E1; adapted referent",
    ),
    (
        "EOS2",
        "Expectation of success in using digital credentials",
        "I believe I can be successful in using digital credentials issued for IT micro-credentials.",
        "Kosovich et al. (2015), E2; adapted referent",
    ),
    (
        "EOS3",
        "Expectation of success in using digital credentials",
        "I am confident that I can understand how to use digital credentials issued for IT micro-credentials.",
        "Kosovich et al. (2015), E3; adapted referent",
    ),
    (
        "PV1",
        "Perceived value from using digital credentials",
        "I think using digital credentials issued for IT micro-credentials is important.",
        "Kosovich et al. (2015), V1; adapted referent",
    ),
    (
        "PV2",
        "Perceived value from using digital credentials",
        "I value using digital credentials issued for IT micro-credentials.",
        "Kosovich et al. (2015), V2; adapted referent",
    ),
    (
        "PV3",
        "Perceived value from using digital credentials",
        "I think using digital credentials issued for IT micro-credentials is useful.",
        "Kosovich et al. (2015), V3; adapted referent",
    ),
    (
        "PC1",
        "Perceived costs to use digital credentials",
        "Using digital credentials issued for IT micro-credentials would require too much of my time.",
        "Kosovich et al. (2015), C1; adapted referent",
    ),
    (
        "PC2",
        "Perceived costs to use digital credentials",
        "Because of other things that I do, I would not have enough time to use digital credentials issued for IT micro-credentials.",
        "Kosovich et al. (2015), C2; adapted referent",
    ),
    (
        "PC3",
        "Perceived costs to use digital credentials",
        "I would be unable to put in the time needed to use digital credentials issued for IT micro-credentials effectively.",
        "Kosovich et al. (2015), C3; adapted referent",
    ),
    (
        "PC4",
        "Perceived costs to use digital credentials",
        "I would have to give up too much to use digital credentials issued for IT micro-credentials.",
        "Kosovich et al. (2015), C4; adapted referent",
    ),
    (
        "INT1",
        "Intention to use digital credentials",
        "I intend to use digital credentials issued for IT micro-credentials in the future.",
        "Venkatesh et al. (2003), BI1; adapted system and time referent",
    ),
    (
        "INT2",
        "Intention to use digital credentials",
        "I predict that I will use digital credentials issued for IT micro-credentials in the future.",
        "Venkatesh et al. (2003), BI2; adapted system and time referent",
    ),
    (
        "INT3",
        "Intention to use digital credentials",
        "I plan to use digital credentials issued for IT micro-credentials in the future.",
        "Venkatesh et al. (2003), BI3; adapted system and time referent",
    ),
]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    tr_pr.append(marker)


def set_repeat_table_header(row) -> None:
    repeat_header(row)


def set_cell_text(cell, value: str, *, bold=False, center=False, size=9.5) -> None:
    cell.text = value
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.left_indent = Inches(0)
        paragraph.paragraph_format.right_indent = Inches(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(size)
            run.bold = bold


def add_table(
    doc: Document,
    rows: list[list[str]],
    widths: list[float],
    *,
    caption: str | None = None,
    header_fill: str = HEADER_FILL,
    font_size: float = 9.5,
) -> None:
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(caption)
        r.bold = True
        r.font.size = Pt(10.5)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    for c_idx, width in enumerate(widths):
        table.columns[c_idx].width = Inches(width)
    repeat_header(table.rows[0])
    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        for c_idx, value in enumerate(values):
            cell = row.cells[c_idx]
            cell.width = Inches(widths[c_idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(widths[c_idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_text(
                cell,
                value,
                bold=(r_idx == 0),
                center=(r_idx == 0),
                size=font_size,
            )
            if r_idx == 0:
                shade(cell, header_fill)
            elif r_idx % 2 == 0:
                shade(cell, "FAFAFA")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Table of contents updates automatically when opened in Word."
    separate.append(placeholder)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_contents_entry(doc: Document, label: str, page: int, *, level=0, bold=False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Inches(0.22 * level)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    r = p.add_run(label)
    r.bold = bold
    r.font.size = Pt(8.5)
    pr = p.add_run(f"\t{page}")
    pr.bold = bold
    pr.font.size = Pt(8.5)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number(section.footer.paragraphs[0])

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Inches(0.5)

    for name, size in (("Title", 16), ("Heading 1", 14), ("Heading 2", 12.5), ("Heading 3", 12)):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = None
        style.paragraph_format.first_line_indent = None
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def add_body(doc: Document, text: str, *, no_indent=False, bold_lead: str | None = None):
    p = doc.add_paragraph()
    if no_indent:
        p.paragraph_format.first_line_indent = None
    if bold_lead and text.startswith(bold_lead):
        first, rest = text[: len(bold_lead)], text[len(bold_lead) :]
        r = p.add_run(first)
        r.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_numbered(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(f"{idx}. ").bold = True
        p.add_run(item)


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, NOTE_FILL)
    set_cell_text(cell, text, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, path: Path, caption: str, *, width=6.15) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.paragraph_format.first_line_indent = None
    cp.paragraph_format.line_spacing = 1.0
    cp.paragraph_format.space_after = Pt(6)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption)
    r.font.size = Pt(10)


def add_references(doc: Document) -> None:
    for reference in REFERENCES:
        p = doc.add_paragraph(reference)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(5)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def normalize_diacritics(doc: Document) -> None:
    replacements = {
        "Kiiskila": "Kiiskilä",
        "Galan-Muros": "Galán-Muros",
    }

    def fix_paragraph(paragraph) -> None:
        for run in paragraph.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)

    for paragraph in doc.paragraphs:
        fix_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fix_paragraph(paragraph)


def add_title_page(doc: Document, *, evidence=False) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PUBLISHED FRAMEWORK AND QUESTIONNAIRE EVIDENCE PACK" if evidence else TITLE.upper())
    r.bold = True
    r.font.size = Pt(16)
    if not evidence:
        p2 = doc.add_paragraph()
        p2.paragraph_format.first_line_indent = None
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(SUBTITLE)
        r2.bold = True
        r2.font.size = Pt(14)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Min Thiha Oo")
    r.font.size = Pt(13)
    r.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("August 2026").font.size = Pt(12)
    doc.add_page_break()


def load_font(weight: str, size: int) -> ImageFont.FreeTypeFont:
    for candidate in FONT_CANDIDATES[weight]:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    raise FileNotFoundError(
        "No compatible serif font was found. Install Times New Roman or DejaVu Serif."
    )


def prepare_figures() -> None:
    OUTPUTS.mkdir(parents=True, exist_ok=True)
    source = Image.open(SOURCE_PAGE).convert("RGB")
    # Direct crop of the published page. The diagram and original figure caption are retained.
    crop = source.crop((105, 105, 1135, 625))
    crop.save(SOURCE_FIGURE, quality=95)

    width, height = 1900, 1050
    canvas = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(canvas)
    regular = load_font("regular", 36)
    bold = load_font("bold", 38)
    small_bold = load_font("bold", 31)
    small = load_font("regular", 28)

    left_boxes = [
        (70, 70, 760, 260, "Expectation of success in\nusing digital credentials"),
        (70, 430, 760, 620, "Perceived value from using\ndigital credentials"),
        (70, 790, 760, 980, "Perceived costs to use\ndigital credentials"),
    ]
    right_box = (1260, 365, 1840, 685, "Intention to use\ndigital credentials")

    for x1, y1, x2, y2, label in left_boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill="#EAF2F8", outline="#345F8C", width=4)
        lines = label.split("\n")
        line_h = 48
        y = (y1 + y2 - line_h * len(lines)) / 2
        for line in lines:
            box = draw.textbbox((0, 0), line, font=bold)
            draw.text(((x1 + x2 - (box[2] - box[0])) / 2, y), line, fill="#111111", font=bold)
            y += line_h

    x1, y1, x2, y2, label = right_box
    draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill="#E2F0D9", outline="#548235", width=4)
    lines = label.split("\n")
    line_h = 52
    y = (y1 + y2 - line_h * len(lines)) / 2
    for line in lines:
        box = draw.textbbox((0, 0), line, font=bold)
        draw.text(((x1 + x2 - (box[2] - box[0])) / 2, y), line, fill="#111111", font=bold)
        y += line_h

    starts = [(760, 165), (760, 525), (760, 885)]
    ends = [(1260, 430), (1260, 525), (1260, 620)]
    labels = [("H1 (+)", 950, 250), ("H2 (+)", 965, 470), ("H3 (-)", 950, 700)]

    def arrow(start, end):
        draw.line((start, end), fill="#1F4E79", width=8)
        sx, sy = start
        ex, ey = end
        dx, dy = ex - sx, ey - sy
        length = (dx * dx + dy * dy) ** 0.5
        ux, uy = dx / length, dy / length
        px, py = -uy, ux
        tip = (ex, ey)
        base = (ex - 34 * ux, ey - 34 * uy)
        left = (base[0] + 17 * px, base[1] + 17 * py)
        right = (base[0] - 17 * px, base[1] - 17 * py)
        draw.polygon([tip, left, right], fill="#1F4E79")

    for start, end in zip(starts, ends):
        arrow(start, end)
    for text_value, x, y in labels:
        draw.rounded_rectangle((x - 12, y - 8, x + 150, y + 45), radius=8, fill="white")
        draw.text((x, y), text_value, fill="#1F4E79", font=small_bold)

    footer = "Selected direct paths from Kiiskilä, Kukkonen, and Pirkkalainen (2023, Fig. 1)"
    box = draw.textbbox((0, 0), footer, font=small)
    draw.text(((width - (box[2] - box[0])) / 2, 1005), footer, fill="#444444", font=small)
    canvas.save(CONCEPTUAL_FIGURE, quality=95)


def build_master() -> None:
    doc = Document()
    configure_document(doc)
    add_title_page(doc)

    doc.add_heading("Draft Status", level=1)
    add_note(
        doc,
        "This is a pre-data working thesis. Chapters 1-4 and the proposed questionnaire are "
        "drafted for approval. Chapters 5-6 contain reporting structures only. No pilot or main-study "
        "results are claimed.",
    )
    doc.add_heading("Abstract", level=1)
    add_body(
        doc,
        "The final abstract will be written in the past tense after the pilot, main data collection, "
        "and statistical analysis have been completed. It will report the purpose, sample, method, "
        "principal findings, and conclusions without presenting results that do not yet exist.",
    )
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    contents = [
        ("CHAPTER 1: INTRODUCTION", 4, 0, True),
        ("1.1 Background of the Study", 4, 1, False),
        ("1.2 Statement of the Problem", 5, 1, False),
        ("1.3 Research Objectives", 5, 1, False),
        ("1.4 Research Questions", 5, 1, False),
        ("1.5 Scope of the Research", 6, 1, False),
        ("1.6 Research Limitations", 6, 1, False),
        ("1.7 Significance of the Study", 6, 1, False),
        ("1.8 Definition of Key Terms", 7, 1, False),
        ("CHAPTER 2: LITERATURE REVIEW", 8, 0, True),
        ("2.1 Introduction", 8, 1, False),
        ("2.2 Micro-Credentials and Digital Credentials", 8, 1, False),
        ("2.3 Applied Theory: Expectancy-Value-Cost", 9, 1, False),
        ("2.4 Definitions of the Study Variables", 10, 1, False),
        ("2.5 Relationships Between the Variables", 11, 1, False),
        ("2.6 Previous Empirical Studies", 12, 1, False),
        ("2.7 Research Gap", 13, 1, False),
        ("2.8 Chapter Summary", 13, 1, False),
        ("CHAPTER 3: THEORETICAL FRAMEWORK, CONCEPTUAL FRAMEWORK, AND HYPOTHESES", 14, 0, True),
        ("3.1 Introduction", 14, 1, False),
        ("3.2 Published Theoretical Framework", 14, 1, False),
        ("3.3 Conceptual Framework", 15, 1, False),
        ("3.4 Hypotheses Development", 15, 1, False),
        ("3.5 Operationalization of Variables", 16, 1, False),
        ("3.6 Chapter Summary", 17, 1, False),
        ("CHAPTER 4: RESEARCH METHODOLOGY", 18, 0, True),
        ("4.1-4.4 Design, Population, and Sampling", 18, 1, False),
        ("4.5 Research Instrument and Questionnaire Design", 19, 1, False),
        ("4.6 Translation and Pilot Test", 20, 1, False),
        ("4.7-4.9 Data Collection, Analysis, and Ethics", 21, 1, False),
        ("4.10 Chapter Summary", 22, 1, False),
        ("CHAPTER 5: RESEARCH RESULTS", 22, 0, True),
        ("CHAPTER 6: DISCUSSION AND RECOMMENDATIONS", 23, 0, True),
        ("REFERENCES", 24, 0, True),
        ("APPENDIX A: PROPOSED QUESTIONNAIRE", 26, 0, True),
    ]
    for label, page, level, bold in contents:
        add_contents_entry(doc, label, page, level=level, bold=bold)
    doc.add_page_break()

    # Chapter 1
    doc.add_heading("CHAPTER 1: INTRODUCTION", level=1)
    doc.add_heading("1.1 Background of the Study", level=2)
    add_body(
        doc,
        "Higher education is increasingly using short, focused learning experiences to support "
        "upskilling and lifelong learning. UNESCO (2022) and the Council of the European Union "
        "(2022) describe micro-credentials as records of assessed learning outcomes from relatively "
        "small learning experiences. Kato et al. (2020) similarly position alternative credentials as "
        "smaller than conventional degrees and diverse in provider, duration, delivery, and purpose. "
        "These definitions indicate that a micro-credential is not merely a short online course: it "
        "must identify what was learned and how that learning was assessed."
    )
    add_body(
        doc,
        "The learning experience and the record issued for it are related but distinct. Kiiskila et "
        "al. (2023) explain that a verifiable digital credential can carry structured information "
        "about an award, including learning outcomes, skills, assessment, and issuer details, and can "
        "be stored and shared electronically. This distinction matters because students may value the "
        "content of a short course but remain uncertain about whether they can use its digital "
        "credential, whether the credential is worthwhile, and whether the time and effort required "
        "to manage it are acceptable."
    )
    add_body(
        doc,
        "These questions can be examined through the expectancy-value-cost framework. Eccles et al. "
        "(1983) and Wigfield and Eccles (2000) explain achievement-related choices through beliefs "
        "about likely success and the value attached to an activity. Later EVC work treats perceived "
        "cost as a separate component because a person may avoid an activity even when success and "
        "value are judged positively if the required cost is too high (Barron & Hulleman, 2015; "
        "Hulleman et al., 2016). Kiiskila et al. (2023) applied this logic directly to digital "
        "credentials and published a model containing four core constructs: expectation of success "
        "in using digital credentials, perceived value from using digital credentials, perceived "
        "costs to use digital credentials, and intention to use digital credentials."
    )
    add_body(
        doc,
        "The Thai higher-education context makes this question timely without assuming that all "
        "institutions or employers already treat micro-credentials in the same way. Thailand's "
        "Ministry of Higher Education, Science, Research and Innovation has hosted regional policy "
        "discussions on micro-credentials, flexible learning, and recognition (MHESI, 2022, 2023). "
        "For IT students, short credentials may be particularly relevant because digital skills "
        "change quickly. However, the usefulness of a digital credential depends partly on whether "
        "students believe they can use it, see value in using it, and consider its costs manageable."
    )

    doc.add_heading("1.2 Statement of the Problem", level=2)
    add_body(
        doc,
        "Current research does not yet provide a clear quantitative account of these judgments among "
        "Thai IT students. Micro-credential research often studies general benefits, employability, "
        "platform acceptance, or willingness to continue online learning. For example, Miao et al. "
        "(2024) examine behavioral intention toward technology-based micro-credential programs, "
        "whereas Shang and Lyv (2022) examine MOOC continuance intention. These studies are useful, "
        "but they do not test the exact digital-credential EVC model published by Kiiskila et al. "
        "(2023) in Thailand."
    )
    add_body(
        doc,
        "Kiiskila et al. (2023) provide unusually close theoretical support because their framework "
        "uses the same four variables required by the present study. Their evidence was qualitative "
        "and developed in a European higher-education setting. A quantitative study among Thai IT "
        "students can therefore test whether the three direct relationships shown in that published "
        "framework are observable in this population. The study does not claim to measure employer "
        "acceptance, actual employment outcomes, or actual future use. It examines students' reported "
        "expectations, value, costs, and intention at one point in time."
    )

    doc.add_heading("1.3 Research Objectives", level=2)
    add_numbered(
        doc,
        [
            "To examine the effect of expectation of success in using digital credentials on intention to use digital credentials among Thai university students in IT-related fields.",
            "To examine the effect of perceived value from using digital credentials on intention to use digital credentials among Thai university students in IT-related fields.",
            "To examine the effect of perceived costs to use digital credentials on intention to use digital credentials among Thai university students in IT-related fields.",
        ],
    )

    doc.add_heading("1.4 Research Questions", level=2)
    add_numbered(
        doc,
        [
            "How does expectation of success in using digital credentials affect intention to use digital credentials among Thai university students in IT-related fields?",
            "How does perceived value from using digital credentials affect intention to use digital credentials among Thai university students in IT-related fields?",
            "How do perceived costs to use digital credentials affect intention to use digital credentials among Thai university students in IT-related fields?",
        ],
    )

    doc.add_heading("1.5 Scope of the Research", level=2)
    add_body(
        doc,
        "The study is limited to students aged 18 or above who are currently enrolled in an "
        "undergraduate or postgraduate IT-related program at a university in Thailand. Relevant "
        "fields include information technology, information systems, computer science, software "
        "engineering, computer engineering, data science, and closely related digital disciplines. "
        "The research uses a quantitative cross-sectional questionnaire. The independent variables "
        "are expectation of success in using digital credentials, perceived value from using digital "
        "credentials, and perceived costs to use digital credentials. The dependent variable is "
        "intention to use digital credentials. All four names are retained exactly from the core "
        "constructs in Kiiskila et al. (2023). The questionnaire refers specifically to digital "
        "credentials issued for IT micro-credentials and does not evaluate a named provider or brand."
    )

    doc.add_heading("1.6 Research Limitations", level=2)
    add_body(
        doc,
        "Four limitations define the interpretation of the study. First, non-probability recruitment "
        "through accessible universities and student networks will limit statistical generalization "
        "to every Thai university. Second, a cross-sectional survey can identify statistical "
        "associations but cannot establish that the measured beliefs cause later behavior. Third, "
        "self-reported intention may not become actual credential use. Fourth, the study focuses on "
        "students in IT-related fields and digital credentials for IT micro-credentials; findings may "
        "not transfer directly to other academic fields, employers, or traditional qualifications."
    )

    doc.add_heading("1.7 Significance of the Study", level=2)
    add_body(
        doc,
        "For universities and continuing-education units, the study can identify whether students' "
        "intention is more strongly associated with confidence in using a credential, perceived "
        "benefit, or perceived burden. This can guide practical support such as clearer instructions "
        "for storing and sharing credentials, better explanation of credential benefits, and simpler "
        "credential-management processes. The results will not show that any credential guarantees "
        "employment or academic recognition; they will show which learner-side judgments require "
        "attention."
    )
    add_body(
        doc,
        "For students, the study can help providers design communication and support around the "
        "questions students actually face: whether they can use the credential, why doing so is "
        "worthwhile, and what time or opportunity costs are involved. For Thai higher-education "
        "policymakers, the findings can add student evidence to ongoing discussion of flexible "
        "learning and digital recognition without assuming uniform adoption across institutions. For "
        "researchers, the study offers a quantitative test of an exact published digital-credential "
        "framework in a new national and disciplinary setting and documents a transparent adaptation "
        "of a validated EVC questionnaire."
    )

    doc.add_heading("1.8 Definition of Key Terms", level=2)
    key_terms = [
        (
            "Micro-credential. ",
            "A record of assessed learning outcomes from a focused learning experience that is smaller "
            "in scope than a conventional degree. This working definition follows UNESCO (2022), the "
            "Council of the European Union (2022), and Kato et al. (2020).",
        ),
        (
            "Digital credential. ",
            "An electronic, shareable record of an educational achievement that can contain information "
            "about the issuer, learner, learning outcomes, assessment, and verification. The definition "
            "used here follows Kiiskila et al. (2023).",
        ),
        (
            "Use of digital credentials. ",
            "Receiving, storing, managing, verifying, and sharing digital credentials issued for IT "
            "micro-credentials. This study-specific boundary is based on the functions discussed by "
            "Kiiskila et al. (2023).",
        ),
        (
            "Expectation of success in using digital credentials. ",
            "A student's belief that they can understand and competently use digital credentials now "
            "and in the future (Eccles & Wigfield, 2002; Kiiskila et al., 2023).",
        ),
        (
            "Perceived value from using digital credentials. ",
            "A student's judgment that using digital credentials is important, useful, and worthwhile "
            "for their learning or future goals (Kosovich et al., 2015; Kiiskila et al., 2023).",
        ),
        (
            "Perceived costs to use digital credentials. ",
            "A student's perception of time, effort, competing obligations, and sacrificed alternatives "
            "associated with using digital credentials (Flake et al., 2015; Kiiskila et al., 2023).",
        ),
        (
            "Intention to use digital credentials. ",
            "A student's stated plan, prediction, and intention to use digital credentials issued for IT "
            "micro-credentials in the future (Venkatesh et al., 2003; Kiiskila et al., 2023).",
        ),
    ]
    for lead, rest in key_terms:
        add_body(doc, lead + rest, bold_lead=lead)

    # Chapter 2
    doc.add_page_break()
    doc.add_heading("CHAPTER 2: LITERATURE REVIEW", level=1)
    doc.add_heading("2.1 Introduction", level=2)
    add_body(
        doc,
        "This chapter reviews micro-credentials and digital credentials, explains the "
        "expectancy-value-cost framework in detail, provides multiple prior definitions for each "
        "study variable, reviews evidence for each proposed relationship, and identifies the specific "
        "research gap addressed by the study."
    )

    doc.add_heading("2.2 Micro-Credentials and Digital Credentials", level=2)
    doc.add_heading("2.2.1 Concept and Boundaries", level=3)
    add_body(
        doc,
        "Definitions of micro-credentials vary, but recent policy sources converge on several "
        "features. UNESCO (2022) emphasizes a record of focused learning achievement that is "
        "assessed against stated standards. The Council of the European Union (2022) similarly "
        "emphasizes learning outcomes from a small volume of learning and transparent standards. "
        "Kato et al. (2020) place micro-credentials within the broader category of alternative "
        "credentials and show that providers, duration, assessment, and recognition can differ. "
        "Therefore, duration alone does not determine whether a Coursera-style course or another "
        "short course is a micro-credential; assessment and a meaningful record of achievement are "
        "also required."
    )
    add_body(
        doc,
        "Kiiskila et al. (2023) distinguish the learning offering from the digital credential that "
        "records it. A digital credential can contain richer and more verifiable information than a "
        "simple completion certificate. In the present study, the micro-credential is the assessed "
        "short learning experience, while the digital credential is the electronic record that the "
        "student may receive, manage, verify, and share. This boundary prevents the questionnaire "
        "from treating course enrollment, learning quality, employer acceptance, and credential use "
        "as if they were the same construct."
    )

    doc.add_heading("2.2.2 Learner Value and Use", level=3)
    add_body(
        doc,
        "Oliver (2021) proposes that learner value in micro-credentials depends on benefits weighed "
        "against costs, including time and effort. Kiiskila et al. (2023) report that students saw "
        "digital credentials as useful for presenting skills, reviewing learning, and supporting "
        "future education or employment-related activities. At the same time, they identify personal, "
        "institutional, and external factors that may enable or constrain use. Miao et al. (2024) "
        "confirm that behavioral intention is already an important outcome in quantitative "
        "micro-credential research, although their model is based on technology acceptance and "
        "self-determination rather than the EVC constructs used here."
    )

    doc.add_heading("2.2.3 Thai Higher-Education Context", level=3)
    add_body(
        doc,
        "Thai policy discussions have considered micro-credentials as part of flexible learning and "
        "lifelong-learning development. MHESI (2022) convened a regional meeting on opportunities, "
        "challenges, and outlooks for micro-credentials in Thai higher education, and MHESI (2023) "
        "reported a regional policy seminar on micro-credentials in ASEAN higher education. These "
        "activities establish policy relevance, but they do not demonstrate that all Thai students "
        "understand or intend to use digital credentials. Student-level evidence is still required."
    )

    doc.add_heading("2.3 Applied Theory: Expectancy-Value-Cost", level=2)
    doc.add_heading("2.3.1 Origin and Development", level=3)
    add_body(
        doc,
        "The expectancy-value tradition explains choices through two broad judgments: whether a "
        "person expects to succeed and how much the activity is valued. Eccles et al. (1983) "
        "developed a model of achievement-related choices in which expectations and subjective task "
        "values are shaped by prior experience, goals, social context, and interpretations of the "
        "task. Wigfield and Eccles (2000) later reviewed the model as a framework for explaining "
        "achievement motivation and choice. Eccles and Wigfield (2002) further clarify that value is "
        "multifaceted and that expectancies concern anticipated future performance."
    )
    add_body(
        doc,
        "Later EVC work gives cost independent analytical importance. Barron and Hulleman (2015) and "
        "Hulleman et al. (2016) explain that high expectancy and high value may still be insufficient "
        "when an activity demands unacceptable effort, time, emotional burden, or lost alternatives. "
        "Flake et al. (2015) provide measurement evidence for cost as a multidimensional and distinct "
        "construct. Kosovich et al. (2015) then validate a short practical scale with three separate "
        "factors: expectancy, value, and cost."
    )

    doc.add_heading("2.3.2 Prior and Current Applications", level=3)
    add_body(
        doc,
        "EVC has been used to explain educational choices and persistence across different levels and "
        "subjects. Durik et al. (2006) show that task values and ability beliefs predict later literacy "
        "choices. Perez et al. (2019) apply expectancy, value, and cost profiles to undergraduate STEM "
        "persistence. Kosovich et al. (2015) demonstrate that brief EVC measures can relate to future "
        "academic interest. Most important for this study, Kiiskila et al. (2023) apply a modified EVC "
        "model directly to the value and intended use of verifiable digital credentials."
    )

    doc.add_heading("2.3.3 Suitability for the Present Study", level=3)
    add_body(
        doc,
        "EVC is suitable because the research question is not whether students accept a particular "
        "software platform. It is whether they intend to use digital credentials when they consider "
        "their own ability, the value of using the credential, and the associated cost. Kiiskila et "
        "al. (2023) already publish these exact constructs in this exact domain. Using their labels "
        "without renaming avoids an unsupported equivalence between different theories or variables. "
        "The present study selects the three direct predictor-to-intention paths in their published "
        "model and tests them quantitatively among Thai IT students."
    )

    doc.add_heading("2.4 Definitions of the Study Variables", level=2)
    doc.add_heading("2.4.1 Expectation of Success in Using Digital Credentials", level=3)
    add_body(
        doc,
        "Eccles et al. (1983) treat expectancy as a belief about likely success in an upcoming task. "
        "Wigfield and Eccles (2000) describe expectancy beliefs as anticipated performance and "
        "competence in achievement activities. Eccles and Wigfield (2002) distinguish these future "
        "expectations from current ability beliefs while recognizing their close empirical relation. "
        "Kosovich et al. (2015) operationalize expectancy through confidence in learning and "
        "understanding material. In the digital-credential field, Kiiskila et al. (2023) define the "
        "domain as confidence in one's ability to use digital credentials now and in the future."
    )
    add_body(
        doc,
        "Definition adopted in this study. Expectation of success in using digital credentials is a "
        "Thai IT student's belief that they can learn, understand, and successfully use digital "
        "credentials issued for IT micro-credentials.",
        bold_lead="Definition adopted in this study. ",
    )

    doc.add_heading("2.4.2 Perceived Value from Using Digital Credentials", level=3)
    add_body(
        doc,
        "Eccles et al. (1983) explain subjective task value as the importance and attractiveness of "
        "engaging in an activity. Eccles and Wigfield (2002) organize value into attainment, intrinsic, "
        "and utility components. Kosovich et al. (2015) use concise importance, value, and usefulness "
        "items to represent a broad value factor. Oliver (2021) applies a learner-value perspective to "
        "micro-credentials by comparing expected benefits with costs. Kiiskila et al. (2023) apply "
        "value directly to using digital credentials for presenting skills, reviewing learning, and "
        "supporting future goals."
    )
    add_body(
        doc,
        "Definition adopted in this study. Perceived value from using digital credentials is a Thai "
        "IT student's judgment that using digital credentials issued for IT micro-credentials is "
        "important, worthwhile, and useful.",
        bold_lead="Definition adopted in this study. ",
    )

    doc.add_heading("2.4.3 Perceived Costs to Use Digital Credentials", level=3)
    add_body(
        doc,
        "Eccles (2005) discusses cost as negative consequences attached to an achievement choice. "
        "Eccles and Wigfield (2002) include effort and the loss of alternative activities within this "
        "judgment. Flake et al. (2015) distinguish effort cost, opportunity cost, and emotional cost. "
        "Kosovich et al. (2015) measure cost through excessive time, competing obligations, inability "
        "to invest enough time, and excessive sacrifice. Kiiskila et al. (2023) apply cost to the time "
        "and effort associated with adopting and using digital credentials."
    )
    add_body(
        doc,
        "Definition adopted in this study. Perceived costs to use digital credentials are a Thai IT "
        "student's perceived time, effort, competing obligations, and sacrificed alternatives involved "
        "in using digital credentials issued for IT micro-credentials.",
        bold_lead="Definition adopted in this study. ",
    )

    doc.add_heading("2.4.4 Intention to Use Digital Credentials", level=3)
    add_body(
        doc,
        "Fishbein and Ajzen (1975) define behavioral intention as the strength of a person's intention "
        "to perform a specified behavior. Ajzen (1991) treats intention as an indication of readiness "
        "and planned effort toward behavior. Venkatesh et al. (2003) operationalize technology-use "
        "intention through intention, prediction, and planning regarding future system use. Miao et "
        "al. (2024) apply behavioral intention to students' future use of technology for "
        "micro-credential programs. Kiiskila et al. (2023) define the exact domain-specific outcome as "
        "intention to use digital credentials."
    )
    add_body(
        doc,
        "Definition adopted in this study. Intention to use digital credentials is a Thai IT student's "
        "stated intention, prediction, and plan to use digital credentials issued for IT "
        "micro-credentials in the future.",
        bold_lead="Definition adopted in this study. ",
    )

    doc.add_heading("2.5 Relationships Between the Variables", level=2)
    doc.add_heading("2.5.1 Expectation of Success and Intention to Use", level=3)
    add_body(
        doc,
        "EVC theory predicts that people are more likely to choose activities in which they expect to "
        "succeed (Eccles et al., 1983; Wigfield & Eccles, 2000). Durik et al. (2006) show that ability "
        "beliefs predict later educational choices, while Kosovich et al. (2015) report evidence "
        "linking expectancy with future interest. Kiiskila et al. (2023) place a direct path from "
        "expectation of success in using digital credentials to intention to use digital credentials. "
        "Together, these studies support a positive relationship in the present context."
    )

    doc.add_heading("2.5.2 Perceived Value and Intention to Use", level=3)
    add_body(
        doc,
        "Subjective value is a central predictor of choice in the Eccles et al. (1983) model, and "
        "Eccles and Wigfield (2002) explain that utility, importance, and interest shape whether a "
        "person engages. Durik et al. (2006) show that task value predicts later choices, and Kosovich "
        "et al. (2015) link value with future interest. In online learning, Shang and Lyv (2022) find "
        "that perceived value contributes to the experience leading to MOOC continuance intention. "
        "Kiiskila et al. (2023) show the exact direct path from perceived value from using digital "
        "credentials to intention to use digital credentials."
    )

    doc.add_heading("2.5.3 Perceived Costs and Intention to Use", level=3)
    add_body(
        doc,
        "Cost can prevent participation even when expectancy and value are favorable (Barron & "
        "Hulleman, 2015). Flake et al. (2015) establish cost as a distinct EVC construct and explain "
        "its negative relation to motivation and choice. Kosovich et al. (2015) show that cost is "
        "separate from expectancy and value and relates negatively to future interest, while Perez et "
        "al. (2019) connect EVC profiles including cost with undergraduate STEM persistence. "
        "Kiiskila et al. (2023) conclude that lower perceived costs support stronger adoption of "
        "digital credentials, which supports a negative relationship between perceived costs and "
        "intention."
    )

    doc.add_heading("2.6 Previous Empirical Studies", level=2)
    previous_rows = [["Study", "Context and method", "Variables relevant here", "Main contribution"]]
    previous_rows += [
        [
            "Durik et al. (2006)",
            "Longitudinal educational study",
            "Ability beliefs, task values, later choices",
            "Shows that EVC-related beliefs predict later educational choices.",
        ],
        [
            "Kosovich et al. (2015)",
            "Middle-school math and science; 10-item survey; CFA",
            "Expectancy, value, cost, future interest",
            "Validates a brief three-factor EVC scale and its 6-point response format.",
        ],
        [
            "Flake et al. (2015)",
            "Educational measurement studies",
            "Effort, opportunity, and emotional cost",
            "Supports cost as a separate component rather than an unnamed part of value.",
        ],
        [
            "Perez et al. (2019)",
            "Undergraduate STEM persistence",
            "Science expectancy, value, and cost profiles",
            "Extends EVC evidence to university STEM students and persistence.",
        ],
        [
            "Shang & Lyv (2022)",
            "555 MOOC learners; PLS-SEM",
            "Perceived value and continuance intention",
            "Provides online-learning evidence that perceived value contributes to continued use.",
        ],
        [
            "Kiiskila et al. (2023)",
            "19 learners and 19 administrators from 11 higher-education institutions; interviews",
            "Exact four digital-credential constructs used in this thesis",
            "Publishes the direct source framework and adapts EVC questions to digital credentials.",
        ],
        [
            "Miao et al. (2024)",
            "474 university students; 5-point survey; PLS-SEM",
            "Behavioral intention toward micro-credential technology",
            "Confirms that a three-item intention construct is reliable in micro-credential research.",
        ],
    ]
    add_table(doc, previous_rows, [1.15, 1.65, 1.55, 1.85], caption="Table 2.1: Previous studies relevant to the proposed model", font_size=8.8)

    doc.add_heading("2.7 Research Gap", level=2)
    add_body(
        doc,
        "The literature supports the EVC constructs, the three proposed directions, and a practical "
        "questionnaire. Nevertheless, three gaps remain. First, Kiiskila et al.'s (2023) exact "
        "digital-credential model was developed qualitatively rather than tested through a large "
        "quantitative student survey. Second, the model has not been tested specifically among Thai "
        "university students in IT-related fields. Third, the relationship between the model's exact "
        "constructs and intention to use digital credentials for IT micro-credentials remains "
        "uncertain in Thailand. The present study addresses these gaps without adding or renaming "
        "variables that are absent from the borrowed framework."
    )

    doc.add_heading("2.8 Chapter Summary", level=2)
    add_body(
        doc,
        "This chapter established the distinction between micro-credentials and digital credentials, "
        "explained the origin and development of EVC, supplied multiple definitions for every study "
        "variable, and reviewed evidence for each relationship. Chapter 3 presents the published "
        "framework directly, traces its exact labels into the conceptual framework, and states the "
        "hypotheses."
    )

    # Chapter 3
    doc.add_page_break()
    doc.add_heading("CHAPTER 3: THEORETICAL FRAMEWORK, CONCEPTUAL FRAMEWORK, AND HYPOTHESES", level=1)
    doc.add_heading("3.1 Introduction", level=2)
    add_body(
        doc,
        "This chapter presents the actual published framework borrowed for the study, explains the "
        "selected direct paths, shows the study's conceptual framework, develops the hypotheses, and "
        "provides operational definitions. The framework variables are not renamed."
    )

    doc.add_heading("3.2 Published Theoretical Framework Borrowed from Previous Research", level=2)
    add_body(
        doc,
        "The primary borrowed framework is Figure 1 in Kiiskila, Kukkonen, and Pirkkalainen (2023). "
        "Their peer-reviewed open-access study applies expectancy-value theory to verifiable digital "
        "credentials in higher education. The researchers interviewed 19 learners and 19 "
        "administrators from 11 higher-education institutions. Their learner questions were based on "
        "EVC surveys including Kosovich et al. (2015), and the published findings model was modified "
        "from Hulleman et al. (2016)."
    )
    add_figure(
        doc,
        SOURCE_FIGURE,
        "Figure 3.1: Direct reproduction of Kiiskila et al. (2023, Figure 1, p. 366). "
        "The image is cropped from the original published PDF; no boxes, labels, or arrows were redrawn. "
        "The article is licensed CC BY 4.0.",
    )
    add_body(
        doc,
        "The source figure includes three direct arrows into intention to use digital credentials. It "
        "also includes correlations among the EVC predictors and personal, institutional, and external "
        "enabling factors. The present study selects only the three direct predictor-to-intention paths "
        "because these provide a focused three-independent-variable, one-dependent-variable model. "
        "Selecting a subset changes the scope of the test but does not rename any selected construct."
    )
    trace_rows = [
        ["Exact wording in Kiiskila et al. (2023)", "Role in this study", "Name changed?"],
        ["Expectation of success in using digital credentials", "Independent variable", "No"],
        ["Perceived value from using digital credentials", "Independent variable", "No"],
        ["Perceived costs to use digital credentials", "Independent variable", "No"],
        ["Intention to use digital credentials", "Dependent variable", "No"],
    ]
    add_table(doc, trace_rows, [3.3, 1.65, 1.2], caption="Table 3.1: Exact-name traceability from the published framework", font_size=9.5)

    doc.add_heading("3.3 Conceptual Framework", level=2)
    add_figure(
        doc,
        CONCEPTUAL_FIGURE,
        "Figure 3.2: Conceptual framework of the present study. Author-created diagram selecting the "
        "three direct paths in Kiiskila et al. (2023, Figure 1). All construct labels are unchanged and "
        "all hypothesis arrows are straight and labeled.",
    )
    add_body(
        doc,
        "Figure 3.2 is an adapted study diagram, not a screenshot from the source. Its function is to "
        "show exactly which published paths will be tested. H1 and H2 are expected to be positive, "
        "while H3 is expected to be negative."
    )

    doc.add_heading("3.4 Hypotheses Development", level=2)
    doc.add_heading("3.4.1 Expectation of Success in Using Digital Credentials", level=3)
    add_body(
        doc,
        "Expectancy is expected to support choice because people are more willing to engage when they "
        "believe that success is attainable (Eccles et al., 1983; Wigfield & Eccles, 2000). Ability "
        "beliefs predict later educational choices (Durik et al., 2006), and brief expectancy measures "
        "relate to future interest (Kosovich et al., 2015). Kiiskila et al. (2023) depict the exact "
        "positive direct path for digital credentials."
    )
    add_body(
        doc,
        "H1: Expectation of success in using digital credentials has a positive effect on intention to "
        "use digital credentials among Thai university students in IT-related fields.",
        no_indent=True,
        bold_lead="H1: ",
    )

    doc.add_heading("3.4.2 Perceived Value from Using Digital Credentials", level=3)
    add_body(
        doc,
        "Value is expected to support intention because importance and utility guide educational "
        "choice (Eccles et al., 1983; Eccles & Wigfield, 2002). Task values predict later learning "
        "choices (Durik et al., 2006), while perceived value contributes to continued online-learning "
        "use (Shang & Lyv, 2022). Kiiskila et al. (2023) depict the exact direct path in the "
        "digital-credential context."
    )
    add_body(
        doc,
        "H2: Perceived value from using digital credentials has a positive effect on intention to use "
        "digital credentials among Thai university students in IT-related fields.",
        no_indent=True,
        bold_lead="H2: ",
    )

    doc.add_heading("3.4.3 Perceived Costs to Use Digital Credentials", level=3)
    add_body(
        doc,
        "Cost is expected to reduce intention because time, effort, and lost alternatives can deter a "
        "choice even when expectancy and value are positive (Barron & Hulleman, 2015). Cost is "
        "empirically distinct from expectancy and value (Flake et al., 2015; Kosovich et al., 2015), "
        "and EVC profiles are related to undergraduate STEM persistence (Perez et al., 2019). "
        "Kiiskila et al. (2023) conclude that low cost supports stronger digital-credential adoption."
    )
    add_body(
        doc,
        "H3: Perceived costs to use digital credentials have a negative effect on intention to use "
        "digital credentials among Thai university students in IT-related fields.",
        no_indent=True,
        bold_lead="H3: ",
    )

    doc.add_heading("3.5 Operationalization of Variables", level=2)
    op_rows = [["Construct (exact published name)", "Role", "Operational definition", "Items and source"]]
    op_rows += [
        [
            "Expectation of success in using digital credentials",
            "IV",
            "Belief that the student can learn, understand, and successfully use digital credentials issued for IT micro-credentials.",
            "EOS1-EOS3; adapted from Kosovich et al. (2015), E1-E3.",
        ],
        [
            "Perceived value from using digital credentials",
            "IV",
            "Judgment that using digital credentials issued for IT micro-credentials is important, valuable, and useful.",
            "PV1-PV3; adapted from Kosovich et al. (2015), V1-V3.",
        ],
        [
            "Perceived costs to use digital credentials",
            "IV",
            "Perceived time, competing obligations, and sacrifice involved in using digital credentials issued for IT micro-credentials.",
            "PC1-PC4; adapted from Kosovich et al. (2015), C1-C4.",
        ],
        [
            "Intention to use digital credentials",
            "DV",
            "Stated intention, prediction, and plan to use digital credentials issued for IT micro-credentials in the future.",
            "INT1-INT3; adapted from Venkatesh et al. (2003), BI1-BI3; same-field support from Miao et al. (2024).",
        ],
    ]
    add_table(doc, op_rows, [1.65, 0.65, 2.05, 1.85], caption="Table 3.2: Operationalization of the exact framework constructs", font_size=8.7)

    doc.add_heading("3.6 Chapter Summary", level=2)
    add_body(
        doc,
        "This chapter reproduced the actual published source framework, demonstrated exact-name "
        "traceability, presented a conceptual diagram with straight labeled arrows, supported each "
        "hypothesis with prior studies, and mapped each construct to borrowed questionnaire items. "
        "Chapter 4 explains how the proposed quantitative survey will be conducted."
    )

    # Chapter 4
    doc.add_page_break()
    doc.add_heading("CHAPTER 4: RESEARCH METHODOLOGY", level=1)
    doc.add_heading("4.1 Introduction", level=2)
    add_body(
        doc,
        "This chapter presents the proposed quantitative design, target population, sample, "
        "questionnaire, pilot test, data collection procedure, statistical analysis, and ethical "
        "considerations. The study contains no interviews and no experimental manipulation."
    )

    doc.add_heading("4.2 Research Design", level=2)
    add_body(
        doc,
        "The study uses a quantitative, explanatory, cross-sectional survey design. A structured "
        "questionnaire is appropriate because all four constructs are measured through multi-item "
        "agreement scales and the research questions concern the statistical effects of three "
        "independent variables on one dependent variable. The design is observational: respondents "
        "are not randomly assigned to different credential descriptions, and no credential feature is "
        "manipulated. Consequently, regression coefficients will be interpreted as predictive "
        "associations rather than proof of causal effects."
    )

    doc.add_heading("4.3 Population and Sample Size", level=2)
    add_body(
        doc,
        "The target population comprises university students aged 18 or above who are currently "
        "enrolled in IT-related undergraduate or postgraduate programs in Thailand. Because a complete "
        "national sampling frame is not available, the study will seek at least 384 usable main-study "
        "responses and will target 400 to allow for unusable or incomplete submissions. The minimum of "
        "384 follows the large-population value commonly reported by Krejcie and Morgan (1970) for a "
        "95% confidence level and a 5% margin under conservative assumptions. A separate pilot of 40 "
        "eligible students will be conducted before the main survey, and pilot responses will not be "
        "included in the main dataset."
    )

    doc.add_heading("4.4 Sampling Procedure", level=2)
    add_body(
        doc,
        "Purposive screening and convenience recruitment will be combined. Survey invitations will be "
        "distributed through accessible university programs, lecturers, student groups, and online "
        "communities. The opening screening questions will confirm age, current enrollment at a Thai "
        "university, and an IT-related field. Respondents who do not meet all three conditions will not "
        "continue to the construct items. Recruitment will seek more than one university and more than "
        "one IT discipline to reduce dependence on a single program, while the limitations of "
        "non-probability sampling will be reported."
    )

    doc.add_heading("4.5 Research Instrument and Questionnaire Design", level=2)
    add_body(
        doc,
        "The questionnaire has four parts. Part A contains consent and three eligibility questions. "
        "Part B gives a short, neutral definition of the study object: a digital credential issued "
        "after an assessed IT micro-credential and capable of being stored, verified, and shared. "
        "This description is an orientation statement, not an experimental vignette. Part C contains "
        "the 13 construct items. Part D contains categorical demographic and background questions. "
        "No brand, provider, guaranteed academic credit, employer endorsement, or employment outcome "
        "is stated."
    )
    add_body(
        doc,
        "The ten EVC items are adapted from the published EVC Scale of Kosovich et al. (2015). Only "
        "the referent is changed from a school subject to digital credentials issued for IT "
        "micro-credentials. This adaptation is additionally supported by Kiiskila et al. (2023), who "
        "state that their digital-credential questions were based on EVC surveys including Kosovich et "
        "al. The three intention items retain the intention-prediction-plan structure of Venkatesh et "
        "al. (2003). Miao et al. (2024) provide same-field support for a three-item behavioral-intention "
        "construct in micro-credential research."
    )
    add_body(
        doc,
        "All 13 construct items use the original EVC six-point agreement format: 1 = Strongly disagree, "
        "2 = Disagree, 3 = Slightly disagree, 4 = Slightly agree, 5 = Agree, and 6 = Strongly agree. "
        "Using one response format reduces switching between scales. Higher scores indicate more of the "
        "named construct. Cost items are not reverse-worded during administration; their expected "
        "regression coefficient is negative."
    )
    item_rows = [["Code", "Exact construct", "Proposed questionnaire item", "Borrowed source"]]
    item_rows += [list(item) for item in EVC_ITEMS]
    add_table(doc, item_rows, [0.7, 1.55, 2.55, 1.4], caption="Table 4.1: Proposed construct items and source traceability", font_size=8.3)

    doc.add_heading("4.6 Translation and Pilot Test", level=2)
    add_body(
        doc,
        "The English master questionnaire will be translated into Thai and independently "
        "back-translated into English following the logic of Brislin (1970). Differences affecting "
        "construct meaning will be corrected before the pilot. The exact English construct labels and "
        "item codes will remain the reference version for analysis and reporting."
    )
    add_body(
        doc,
        "A pilot test will be conducted with 40 students who meet the main eligibility criteria. This "
        "size is at the upper end of the pilot samples evaluated by Hertzog (2008) and exceeds Johanson "
        "and Brooks's (2010) reasonable minimum of 30 representative participants for preliminary "
        "survey or scale development. The pilot will check eligibility routing, missing responses, "
        "completion time, wording clarity, and internal consistency. Cronbach's alpha will be calculated "
        "for each construct, with .70 used as the minimum working criterion. Any wording revision will "
        "be documented against the original item code. Pilot respondents will be excluded from the "
        "main survey, and no pilot result is reported until the pilot has actually occurred."
    )

    doc.add_heading("4.7 Data Collection Procedure", level=2)
    add_body(
        doc,
        "After approval of the framework and questionnaire, the pilot will be administered first. "
        "Following any documented wording corrections, the final online questionnaire will be "
        "distributed to the target population. The first page will explain the study, voluntary "
        "participation, anonymity, expected completion time, and the right to stop. No name, student "
        "number, email address, or other direct identifier will be requested. Collection will continue "
        "until at least 384 usable responses are obtained, with 400 usable responses as the target."
    )

    doc.add_heading("4.8 Statistical Treatment of Data", level=2)
    add_body(
        doc,
        "The data will be analyzed in Jamovi. Screening will remove ineligible cases, duplicate "
        "submissions when identifiable without retaining personal data, substantially incomplete "
        "questionnaires, and cases failing any prespecified attention check. Exclusion rules will be "
        "fixed before hypothesis testing. Frequencies and percentages will summarize categorical "
        "characteristics, while means and standard deviations will summarize items and construct "
        "scores."
    )
    add_body(
        doc,
        "Internal consistency will be assessed separately for EOS, PV, PC, and INT using Cronbach's "
        "alpha and corrected item-total correlations. Construct scores will be calculated as the mean "
        "of their items when reliability is acceptable. Pearson correlations will provide preliminary "
        "bivariate evidence. The main test will be multiple linear regression: INT = b0 + b1(EOS) + "
        "b2(PV) + b3(PC) + error. The analysis will report unstandardized and standardized "
        "coefficients, t values, p values, confidence intervals, the F test, adjusted R-squared, and "
        "variance-inflation factors. Linearity, residual normality, homoscedasticity, influential cases, "
        "and multicollinearity will be checked. H1 and H2 are supported when their coefficients are "
        "positive and statistically significant at p < .05; H3 is supported when its coefficient is "
        "negative and statistically significant at p < .05."
    )

    doc.add_heading("4.9 Ethical Considerations", level=2)
    add_body(
        doc,
        "Participation will be voluntary and limited to adults. The consent page will state the "
        "purpose, procedures, minimal foreseeable risk, confidentiality arrangements, and right to "
        "withdraw before submission. Data will be stored securely and used only for academic research. "
        "Recruitment messages will avoid pressure from lecturers or institutions. The study will obtain "
        "any university approval required before pilot or main data collection."
    )

    doc.add_heading("4.10 Chapter Summary", level=2)
    add_body(
        doc,
        "The proposed study is a quantitative cross-sectional survey of Thai university students in "
        "IT-related fields. It uses a 40-participant pilot, a main target of 400 usable responses, a "
        "six-point Likert questionnaire transparently adapted from published scales, and multiple "
        "linear regression for the three hypotheses."
    )

    # Chapters 5-6 reporting shells
    doc.add_page_break()
    doc.add_heading("CHAPTER 5: RESEARCH RESULTS", level=1)
    add_note(
        doc,
        "Reporting shell only. This chapter will be completed after the pilot and main study. No "
        "numbers or hypothesis decisions should be inserted until they come from the final dataset.",
    )
    for heading, text_value in [
        ("5.1 Data Screening and Response Profile", "Report invitations, submissions, exclusions, and final usable sample."),
        ("5.2 Pilot and Main-Study Reliability", "Report alpha and corrected item-total statistics separately; do not mix pilot and main samples."),
        ("5.3 Descriptive Analysis", "Report demographic frequencies and item/construct means and standard deviations."),
        ("5.4 Correlation Analysis", "Report the correlation matrix with significance levels."),
        ("5.5 Multiple Linear Regression", "Report assumptions, model fit, coefficients, confidence intervals, and VIF."),
        ("5.6 Hypothesis Summary", "State the H1-H3 decisions using the prespecified directions and the .05 significance level."),
    ]:
        doc.add_heading(heading, level=2)
        add_body(doc, text_value)

    doc.add_page_break()
    doc.add_heading("CHAPTER 6: DISCUSSION AND RECOMMENDATIONS", level=1)
    add_note(
        doc,
        "Reporting shell only. Interpret the observed results here after Chapter 5 is complete; do not "
        "write findings in advance.",
    )
    for heading, text_value in [
        ("6.1 Discussion of Findings", "Interpret H1-H3 in order and compare each result with the studies reviewed in Chapter 2."),
        ("6.2 Theoretical Implications", "Explain whether the Thai IT-student evidence supports or qualifies the selected Kiiskila et al. paths."),
        ("6.3 Practical Implications", "Give evidence-based recommendations to universities, credential providers, students, and policymakers."),
        ("6.4 Limitations of the Study", "Revisit sampling, cross-sectional self-report, intention-versus-behavior, and disciplinary scope in light of the completed study."),
        ("6.5 Suggestions for Future Research", "Identify specific extensions such as longitudinal use, employer samples, other fields, or additional published enabling factors."),
        ("6.6 Conclusion", "Answer the three research questions directly without claiming outcomes beyond the data."),
    ]:
        doc.add_heading(heading, level=2)
        add_body(doc, text_value)

    doc.add_page_break()
    doc.add_heading("REFERENCES", level=1)
    add_references(doc)

    doc.add_page_break()
    doc.add_heading("APPENDIX A: PROPOSED QUESTIONNAIRE", level=1)
    doc.add_heading("Participant Information and Consent", level=2)
    add_body(
        doc,
        "This questionnaire studies Thai university students' views about digital credentials issued "
        "for IT micro-credentials. Participation is voluntary and anonymous. You must be at least 18, "
        "currently enrolled at a university in Thailand, and studying an IT-related field. The survey "
        "contains no interview and no open-ended question. By selecting 'I agree' and continuing, you "
        "confirm that you understand the information and voluntarily participate.",
        no_indent=True,
    )
    doc.add_heading("Part A: Screening", level=2)
    screening = [
        "Are you 18 years old or above? (Yes / No)",
        "Are you currently enrolled at a university in Thailand? (Yes / No)",
        "Is your current program in an IT-related field? (Yes / No)",
    ]
    add_numbered(doc, screening)
    doc.add_heading("Part B: Reference Definition", level=2)
    add_note(
        doc,
        "In this survey, an IT micro-credential is a focused learning experience with stated learning "
        "outcomes and assessment. A digital credential is the electronic record issued after successful "
        "completion. It can contain information about the issuer, learning outcomes, assessment, and "
        "verification, and it can be stored and shared. 'Use' includes receiving, storing, managing, "
        "verifying, and sharing the digital credential.",
    )
    doc.add_heading("Part C: Agreement Items", level=2)
    add_body(
        doc,
        "Please select one response for every statement: 1 = Strongly disagree; 2 = Disagree; "
        "3 = Slightly disagree; 4 = Slightly agree; 5 = Agree; 6 = Strongly agree.",
        no_indent=True,
    )
    questionnaire_rows = [["Code", "Statement", "1", "2", "3", "4", "5", "6"]]
    for code, _, item_text, _ in EVC_ITEMS:
        questionnaire_rows.append([code, item_text, "o", "o", "o", "o", "o", "o"])
    add_table(doc, questionnaire_rows, [0.75, 3.5, 0.33, 0.33, 0.33, 0.33, 0.33, 0.33], caption="Table A.1: Proposed English master questionnaire", font_size=8.7)
    doc.add_heading("Part D: Background Questions", level=2)
    background = [
        "Age group: 18-20 / 21-25 / 26-30 / 31 or above",
        "Gender: Woman / Man / Non-binary or another identity / Prefer not to say",
        "Level of study: Undergraduate / Master's / Doctoral / Other postgraduate",
        "Primary field: Information Technology / Information Systems / Computer Science / Software Engineering / Computer Engineering / Data Science / Other related field",
        "University type: Public / Private / Other",
        "Have you completed any online short course? Yes / No / Not sure",
        "Have you earned a micro-credential or digital badge before? Yes / No / Not sure",
        "Before this survey, how familiar were you with digital credentials? Not at all / Slightly / Moderately / Very / Extremely",
    ]
    add_numbered(doc, background)

    normalize_diacritics(doc)
    doc.save(MASTER_OUT)


def build_evidence_pack() -> None:
    doc = Document()
    configure_document(doc)
    add_title_page(doc, evidence=True)

    doc.add_heading("1. Purpose of This Pack", level=1)
    add_body(
        doc,
        "This pack shows, in one place, the actual published framework borrowed for the study, the "
        "exact transfer of its variable names into the conceptual framework, the published "
        "questionnaire sources, every proposed adapted item, and the evidence supporting each "
        "hypothesis. It accompanies the chaptered thesis draft and is not a separate theoretical model."
    )
    add_note(
        doc,
        "Main correction: the previous recognition, stackability, industry-endorsement, trust, and "
        "perceived-value experiment has been removed. Those labels did not appear together as variables "
        "in any borrowed published framework. The replacement uses four exact labels from one "
        "same-field published model.",
    )

    doc.add_heading("2. Source Integrity: What Is Original and What Is Adapted", level=1)
    audit_rows = [
        ["Visual or table", "Status", "How it was produced"],
        [
            "Kiiskila et al. (2023), Figure 1",
            "Direct source reproduction",
            "Cropped directly from page 366 of the original Springer PDF. No boxes, labels, or arrows were generated or redrawn. CC BY 4.0.",
        ],
        [
            "Conceptual framework for this thesis",
            "Author-created adaptation",
            "New diagram selecting three direct paths from the source. Exact variable labels retained; H1-H3 added; straight arrows used.",
        ],
        [
            "Questionnaire wording",
            "Documented adaptation",
            "Item structures borrowed from published scales; the referent is changed to digital credentials issued for IT micro-credentials.",
        ],
    ]
    add_table(doc, audit_rows, [1.8, 1.45, 3.0], caption="Table 2.1: Source audit", font_size=9.2)

    doc.add_heading("3. Actual Published Framework Borrowed", level=1)
    add_body(
        doc,
        "Primary source: Kiiskila, P., Kukkonen, A., and Pirkkalainen, H. (2023), Are "
        "Micro-Credentials Valuable for Students? Perspective on Verifiable Digital Credentials, "
        "SN Computer Science, 4, Article 366. DOI: 10.1007/s42979-023-01797-y."
    )
    add_figure(
        doc,
        SOURCE_FIGURE,
        "Figure 3.1: Direct crop of the source's Figure 1 on p. 366. This is the published diagram, not "
        "a recreation. Source: Kiiskila et al. (2023), CC BY 4.0.",
    )
    add_body(
        doc,
        "Why it fits: the source is peer-reviewed, open access, situated in higher education, and "
        "specifically concerns verifiable digital credentials associated with micro-credentials. The "
        "authors use EVC as their theoretical framework, state that their questions were based on "
        "surveys including Kosovich et al. (2015), and publish the same four core variables used in the "
        "new conceptual framework."
    )
    source_rows = [
        ["Source construct, copied exactly", "Direct source path selected", "Hypothesis"],
        ["Expectation of success in using digital credentials", "To intention to use digital credentials", "H1 positive"],
        ["Perceived value from using digital credentials", "To intention to use digital credentials", "H2 positive"],
        ["Perceived costs to use digital credentials", "To intention to use digital credentials", "H3 negative"],
        ["Intention to use digital credentials", "Dependent variable receiving all three paths", "DV"],
    ]
    add_table(doc, source_rows, [2.45, 2.45, 1.25], caption="Table 3.1: Exact constructs and selected source paths", font_size=9.2)

    doc.add_heading("4. Conceptual Framework Used in the Thesis", level=1)
    add_figure(
        doc,
        CONCEPTUAL_FIGURE,
        "Figure 4.1: Author-created conceptual framework. It selects three straight direct paths from "
        "Kiiskila et al. (2023, Figure 1); it does not rename the variables.",
    )
    exact_rows = [
        ["Published source wording", "Thesis wording", "Match"],
        ["Expectation of success in using digital credentials", "Expectation of success in using digital credentials", "Exact"],
        ["Perceived value from using digital credentials", "Perceived value from using digital credentials", "Exact"],
        ["Perceived costs to use digital credentials", "Perceived costs to use digital credentials", "Exact"],
        ["Intention to use digital credentials", "Intention to use digital credentials", "Exact"],
    ]
    add_table(doc, exact_rows, [2.75, 2.75, 0.7], caption="Table 4.1: Variable-name comparison", font_size=9.1)

    doc.add_heading("5. Questionnaire Evidence", level=1)
    doc.add_heading("5.1 Expectancy-Value-Cost Items", level=2)
    add_body(
        doc,
        "Kosovich et al. (2015) published a brief 10-item EVC Scale: three expectancy items (E1-E3), "
        "three value items (V1-V3), and four cost items (C1-C4). The Appendix uses a six-point scale "
        "from Strongly disagree to Strongly agree. Confirmatory factor analysis supported separate "
        "expectancy, value, and cost factors, and the paper reports good omega reliability. The "
        "published Appendix is available at DOI 10.1177/0272431614556890."
    )
    add_body(
        doc,
        "The adaptation is narrow: the source's math/science-class referent is replaced with digital "
        "credentials issued for IT micro-credentials. The sentence purpose and item order remain "
        "traceable through E1-E3, V1-V3, and C1-C4. Kiiskila et al. (2023) independently support this "
        "move because they state that their digital-credential questions were adapted from EVC surveys "
        "including Kosovich et al."
    )
    doc.add_heading("5.2 Intention Items", level=2)
    add_body(
        doc,
        "Venkatesh et al. (2003) publish three behavioral-intention items (BI1-BI3) built around "
        "intention, prediction, and planning to use a system in the future. The present items change the "
        "system referent to digital credentials issued for IT micro-credentials and use 'in the future' "
        "as the common time referent. The source article is available through the AIS eLibrary and DOI "
        "10.2307/30036540. Miao et al. (2024) provide field-specific corroboration: their "
        "micro-credential study used three behavioral-intention items and reported alpha = .870, "
        "composite reliability = .920, and AVE = .794. Miao et al. identify the number, source, and "
        "measurement performance but do not print the full wording, so they are used as contextual "
        "evidence rather than as the wording source."
    )

    doc.add_heading("5.3 Full Proposed Questionnaire Items", level=2)
    item_rows = [["Code", "Exact construct", "Proposed item", "Source mapping"]]
    item_rows += [list(item) for item in EVC_ITEMS]
    add_table(doc, item_rows, [0.7, 1.55, 2.55, 1.4], caption="Table 5.1: Final English items ready for Thai translation and pilot", font_size=8.3)
    add_note(
        doc,
        "Response scale for every construct item: 1 = Strongly disagree; 2 = Disagree; "
        "3 = Slightly disagree; 4 = Slightly agree; 5 = Agree; 6 = Strongly agree. No interviews and "
        "no open-ended questions are proposed.",
    )

    doc.add_heading("6. Evidence for Each Hypothesis", level=1)
    relationship_rows = [
        ["Hypothesis", "Direct framework evidence", "Additional relationship evidence"],
        [
            "H1: Expectation of success -> intention (+)",
            "Kiiskila et al. (2023), Figure 1 and discussion",
            "Eccles et al. (1983); Wigfield & Eccles (2000); Durik et al. (2006); Kosovich et al. (2015)",
        ],
        [
            "H2: Perceived value -> intention (+)",
            "Kiiskila et al. (2023), Figure 1 and discussion",
            "Eccles et al. (1983); Eccles & Wigfield (2002); Durik et al. (2006); Shang & Lyv (2022)",
        ],
        [
            "H3: Perceived costs -> intention (-)",
            "Kiiskila et al. (2023), Figure 1 and discussion",
            "Barron & Hulleman (2015); Flake et al. (2015); Kosovich et al. (2015); Perez et al. (2019)",
        ],
    ]
    add_table(doc, relationship_rows, [1.8, 1.9, 2.5], caption="Table 6.1: Relationship-evidence matrix", font_size=9.1)

    doc.add_heading("7. Definition Evidence Required for Chapter 2", level=1)
    definition_rows = [
        ["Exact variable", "Four or more definition sources", "Study definition"],
        [
            "Expectation of success in using digital credentials",
            "Eccles et al. (1983); Wigfield & Eccles (2000); Eccles & Wigfield (2002); Kosovich et al. (2015); Kiiskila et al. (2023)",
            "Belief that the student can learn, understand, and successfully use digital credentials issued for IT micro-credentials.",
        ],
        [
            "Perceived value from using digital credentials",
            "Eccles et al. (1983); Eccles & Wigfield (2002); Kosovich et al. (2015); Oliver (2021); Kiiskila et al. (2023)",
            "Judgment that using the digital credential is important, valuable, and useful.",
        ],
        [
            "Perceived costs to use digital credentials",
            "Eccles (2005); Eccles & Wigfield (2002); Flake et al. (2015); Kosovich et al. (2015); Kiiskila et al. (2023)",
            "Perceived time, competing obligations, and sacrifice involved in use.",
        ],
        [
            "Intention to use digital credentials",
            "Fishbein & Ajzen (1975); Ajzen (1991); Venkatesh et al. (2003); Kiiskila et al. (2023); Miao et al. (2024)",
            "Stated intention, prediction, and plan to use the digital credential in the future.",
        ],
    ]
    add_table(doc, definition_rows, [1.75, 2.75, 1.7], caption="Table 7.1: Variable-definition evidence", font_size=8.8)

    doc.add_heading("8. Pilot and Analysis Evidence", level=1)
    method_rows = [
        ["Decision", "Planned treatment", "Evidence"],
        ["Pilot", "40 eligible Thai IT students; excluded from main sample", "Hertzog (2008) evaluates pilot samples up to 40; Johanson & Brooks (2010) recommend 30 representative participants as a reasonable minimum."],
        ["Main sample", "Minimum 384 usable; target 400", "Krejcie & Morgan (1970) large-population table under conventional 95%/5% assumptions."],
        ["Response scale", "Six-point agreement for all 13 construct items", "Preserves Kosovich et al. (2015) EVC response format."],
        ["Primary analysis", "Multiple linear regression in Jamovi", "One continuous DV and three continuous IV composite scores."],
        ["Claim boundary", "Associations and prediction, not causal effects or actual employment outcomes", "Cross-sectional observational design."],
    ]
    add_table(doc, method_rows, [1.2, 2.5, 2.5], caption="Table 8.1: Method decisions and evidence", font_size=9.0)

    doc.add_heading("9. Source Access List", level=1)
    source_access = [
        ["Purpose", "Source", "Where the evidence appears"],
        ["Exact published framework", "Kiiskila et al. (2023)", "Figure 1, p. 366; research framework and discussion; DOI 10.1007/s42979-023-01797-y"],
        ["EVC questionnaire", "Kosovich et al. (2015)", "Published Appendix, E1-E3, V1-V3, C1-C4; DOI 10.1177/0272431614556890"],
        ["Intention questionnaire", "Venkatesh et al. (2003)", "Behavioral intention items BI1-BI3, p. 460; DOI 10.2307/30036540"],
        ["Micro-credential intention context", "Miao et al. (2024)", "Table 1 source mapping and Table 3 reliability; DOI 10.1108/IJEM-02-2023-0066"],
        ["Pilot-size evidence", "Hertzog (2008); Johanson & Brooks (2010)", "Aim-specific pilot guidance and a 30-participant reasonable minimum; DOIs 10.1002/nur.20247 and 10.1177/0013164409355692"],
    ]
    add_table(doc, source_access, [1.45, 1.65, 3.1], caption="Table 9.1: Quick source-location guide", font_size=9.0)

    doc.add_heading("References", level=1)
    add_references(doc)
    normalize_diacritics(doc)
    doc.save(EVIDENCE_OUT)


def main() -> None:
    prepare_figures()
    build_master()
    build_evidence_pack()
    print(MASTER_OUT)
    print(EVIDENCE_OUT)
    print(SOURCE_FIGURE)
    print(CONCEPTUAL_FIGURE)


if __name__ == "__main__":
    main()
