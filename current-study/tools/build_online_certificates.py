from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_direct_replication as layout
import build_documents as base


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
SOURCE = ROOT / "source"

MASTER_OUT = OUTPUTS / "Master_Thesis_Online_IT_Professional_Certificates_Thai_Students.docx"
AUTHOR_FRAMEWORK = OUTPUTS / "Conceptual_Framework_TPB_Online_IT_Certificates.png"
HUNSINGER_FRAMEWORK = SOURCE / "Hunsinger_Smith_2008_Figure_2_Direct_Crop.png"
WANG_FRAMEWORK = SOURCE / "Wang_2023_Figure_1_Direct_Crop.png"

TITLE = "Factors Influencing Thai IT Students' Intentions to Pursue Short Online IT Professional Certificates"
SUBTITLE = "A Theory of Planned Behavior Study"


HYPOTHESES = [
    (
        "H1",
        "Attitude toward pursuing a short online IT professional certificate is positively related to behavioral intention to pursue such a certificate within the next 12 months.",
    ),
    (
        "H2",
        "Subjective norm concerning pursuit of a short online IT professional certificate is positively related to behavioral intention to pursue such a certificate within the next 12 months.",
    ),
    (
        "H3",
        "Perceived behavioral control over pursuing a short online IT professional certificate is positively related to behavioral intention to pursue such a certificate within the next 12 months.",
    ),
]


REFERENCES = [
    "Ajzen, I. (1991). The theory of planned behavior. Organizational Behavior and Human Decision Processes, 50(2), 179-211. https://doi.org/10.1016/0749-5978(91)90020-T",
    "Council of the European Union. (2022). Council Recommendation of 16 June 2022 on a European approach to micro-credentials for lifelong learning and employability (2022/C 243/02). Official Journal of the European Union, C 243, 10-25. https://eur-lex.europa.eu/eli/C/2022/243/oj",
    "Cronbach, L. J. (1951). Coefficient alpha and the internal structure of tests. Psychometrika, 16(3), 297-334. https://doi.org/10.1007/BF02310555",
    "Fianu, E., Blewett, C., Ampong, G. O. A., & Ofori, K. S. (2018). Factors affecting MOOC usage by students in selected Ghanaian universities. Education Sciences, 8(2), Article 70. https://doi.org/10.3390/educsci8020070",
    "Hair, J. F., Black, W. C., Babin, B. J., & Anderson, R. E. (2019). Multivariate data analysis (8th ed.). Cengage.",
    "Hertzog, M. A. (2008). Considerations in determining sample size for pilot studies. Research in Nursing & Health, 31(2), 180-191. https://doi.org/10.1002/nur.20247",
    "Hunsinger, D. S., & Smith, M. A. (2008). Factors that influence information systems undergraduates to pursue IT certification. Journal of Information Technology Education, 7, 247-265. https://doi.org/10.28945/188",
    "Ilias, A., Baidi, N., Ghani, E. K., & Abdul Rahman, R. (2022). Factors driving the intention to pursue internal auditing certification and career among future graduates in Malaysia. Universal Journal of Accounting and Finance, 10(2), 549-558. https://doi.org/10.13189/ujaf.2022.100219",
    "Johanson, G. A., & Brooks, G. P. (2010). Initial scale development: Sample size for pilot studies. Educational and Psychological Measurement, 70(3), 394-400. https://doi.org/10.1177/0013164409355692",
    "Krejcie, R. V., & Morgan, D. W. (1970). Determining sample size for research activities. Educational and Psychological Measurement, 30(3), 607-610. https://doi.org/10.1177/001316447003000308",
    "Miao, M., Ahmed, M., Ahsan, N., & Qamar, B. (2024). Intention to use technology for micro-credential programs: Evidence from technology acceptance and self-determination model. International Journal of Educational Management, 38(4), 948-977. https://doi.org/10.1108/IJEM-02-2023-0066",
    "Ministry of Higher Education, Science, Research and Innovation. (2022). Micro-credentials in Thai higher education: Opportunities, challenges and outlooks from ASEAN and Europe. https://www.mhesi.go.th/index.php/en/news-and-announce-all/news-all/executive-ps-news/7614-micro-credentials-in-thai-higher-education-opportunities-challenges-and-outlooks-from-asean-and-europe.html",
    "Mohan, M. M., Upadhyaya, P., & Pillai, K. R. (2020). Intention and barriers to use MOOCs: An investigation among the post graduate students in India. Education and Information Technologies, 25, 5017-5031. https://doi.org/10.1007/s10639-020-10215-2",
    "Steenkamp, N., Fisher, R., & Nesbit, T. (2024). Understanding accounting students' intentions to use digital badges to showcase employability skills. Accounting Education, 33(6), 906-934. https://doi.org/10.1080/09639284.2023.2276200",
    "UNESCO. (2022). Towards a common definition of micro-credentials. https://unesdoc.unesco.org/ark:/48223/pf0000381668",
    "Wang, K. (2023). The perception and behavioral intention toward MOOCs: Undergraduates in China. International Review of Research in Open and Distributed Learning, 24(1), 22-46. https://doi.org/10.19173/irrodl.v24i1.6677",
]


REFERENTS = [
    ("IT managers", "IT managers"),
    ("My professors", "my professors"),
    ("Hiring managers", "hiring managers"),
    ("My advisors", "my advisors"),
    ("My parents", "my parents"),
    ("The general public", "the general public"),
]


CONTROLS = [
    ("learning ability", "have the learning ability"),
    ("knowledge", "have the knowledge"),
    ("skills", "have the skills"),
    ("money and resources", "have the money and resources"),
]


def font_path(bold: bool = False) -> str:
    candidates = (
        ["/System/Library/Fonts/Supplemental/Arial Bold.ttf", "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf"]
        if bold
        else ["/System/Library/Fonts/Supplemental/Arial.ttf", "/System/Library/Fonts/Supplemental/Times New Roman.ttf"]
    )
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    raise FileNotFoundError("A suitable system font was not found")


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#1f4e79") -> None:
    draw.line([start, end], fill=color, width=5)
    x2, y2 = end
    draw.polygon([(x2, y2), (x2 - 18, y2 - 11), (x2 - 18, y2 + 11)], fill=color)


def make_author_framework() -> None:
    canvas = Image.new("RGB", (1800, 820), "white")
    draw = ImageDraw.Draw(canvas)
    regular = ImageFont.truetype(font_path(), 38)
    bold = ImageFont.truetype(font_path(True), 40)
    small = ImageFont.truetype(font_path(True), 30)

    left_x1, left_x2 = 90, 760
    right_x1, right_x2 = 1180, 1710
    boxes = [
        ((left_x1, 110, left_x2, 270), "Attitude"),
        ((left_x1, 330, left_x2, 490), "Subjective Norm"),
        ((left_x1, 550, left_x2, 710), "Perceived Behavioral\nControl"),
        ((right_x1, 290, right_x2, 535), "Behavioral Intention\nto Pursue"),
    ]
    for box, label in boxes:
        draw.rounded_rectangle(box, radius=14, fill="#f7f9fb", outline="#1f4e79", width=4)
        lines = label.split("\n")
        heights = [draw.textbbox((0, 0), line, font=bold)[3] for line in lines]
        total_h = sum(heights) + (12 if len(lines) > 1 else 0)
        y = (box[1] + box[3] - total_h) / 2
        for line, height in zip(lines, heights):
            bounds = draw.textbbox((0, 0), line, font=bold)
            width = bounds[2] - bounds[0]
            draw.text(((box[0] + box[2] - width) / 2, y), line, fill="#17212b", font=bold)
            y += height + 12

    arrows = [
        ((left_x2, 190), (right_x1, 345), "H1"),
        ((left_x2, 410), (right_x1, 410), "H2"),
        ((left_x2, 630), (right_x1, 480), "H3"),
    ]
    for start, end, label in arrows:
        draw_arrow(draw, start, end)
        lx = (start[0] + end[0]) // 2 - 24
        ly = (start[1] + end[1]) // 2 - 48
        draw.rounded_rectangle((lx - 8, ly - 5, lx + 68, ly + 38), radius=7, fill="white")
        draw.text((lx, ly), label, fill="#1f4e79", font=small)

    draw.text((90, 35), "Independent variables", fill="#555555", font=regular)
    draw.text((1240, 220), "Dependent variable", fill="#555555", font=regular)
    canvas.save(AUTHOR_FRAMEWORK)


def configure_document(doc: Document) -> None:
    base.configure_document(doc)
    doc.sections[0].header.paragraphs[0].text = ""
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.color.rgb = None
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def title_page(doc: Document) -> None:
    for line in [
        "ASSUMPTION UNIVERSITY OF THAILAND",
        "Graduate School of Business and Advanced Technology Management",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(15)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SUBTITLE)
    run.italic = True
    run.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BY\nMR. MIN THIHA OO")
    run.bold = True
    run.font.size = Pt(12)

    for _ in range(2):
        doc.add_paragraph()
    for line in [
        "ITM7000 INDEPENDENT STUDY",
        "Submitted in Partial Fulfillment of the Requirements for the Degree of",
        "Master of Information Technology Management",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11.5)
        if line == "ITM7000 INDEPENDENT STUDY":
            run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("SEPTEMBER 2026").font.size = Pt(12)
    doc.add_page_break()


def add_hypotheses(doc: Document) -> None:
    for code, statement in HYPOTHESES:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.hanging_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(5)
        p.add_run(f"{code}: ").bold = True
        p.add_run(statement)


def add_questionnaire(doc: Document) -> None:
    heading = doc.add_heading("APPENDIX A: QUESTIONNAIRE", level=1)
    heading.paragraph_format.page_break_before = True
    layout.add_body(
        doc,
        "Questionnaire items are adopted from Hunsinger and Smith (2008, pp. 252-254). The target phrase 'IT certification' is replaced consistently with 'short online IT professional certificate' to match the defined behavior. The constructs, referent groups, response anchors, and 12-month time horizon are otherwise retained.",
    )
    layout.add_source_note(
        doc,
        "Full source: Hunsinger, D. S., and Smith, M. A. (2008). Factors that influence information systems undergraduates to pursue IT certification. Journal of Information Technology Education, 7, 247-265. https://doi.org/10.28945/188.",
    )

    doc.add_heading("Participant Information and Consent", level=2)
    layout.add_body(
        doc,
        "This questionnaire is part of a master's research study about Thai IT students' intentions to pursue short online IT professional certificates. Participation is voluntary. No name or email address is requested. Responses will be reported only in aggregate form. A respondent may leave the questionnaire before submission without penalty.",
    )
    layout.add_table(
        doc,
        [
            ["Question", "Response options"],
            ["I have read the information above and voluntarily agree to participate.", "Yes / No"],
        ],
        [4.6, 1.6],
        caption="Table A.1: Informed consent",
        font_size=9,
    )

    doc.add_heading("Eligibility", level=2)
    layout.add_table(
        doc,
        [
            ["Question", "Response options"],
            ["Are you 18 years of age or older?", "Yes / No"],
            ["Are you currently enrolled as an undergraduate student at a university in Thailand?", "Yes / No"],
            ["Are you studying an IT-related programme?", "Yes / No"],
        ],
        [4.6, 1.6],
        caption="Table A.2: Eligibility questions",
        font_size=9,
    )
    layout.add_body(doc, "A 'No' response to any eligibility question ends the questionnaire.", no_indent=True)

    doc.add_heading("Respondent Profile", level=2)
    layout.add_table(
        doc,
        [
            ["Question", "Response options"],
            ["Age group", "18-20 / 21-23 / 24-26 / 27 or older / Prefer not to say"],
            ["Gender", "Woman / Man / Non-binary or another identity / Prefer not to say"],
            ["Year of undergraduate study", "Year 1 / Year 2 / Year 3 / Year 4 / Year 5 or later"],
            ["IT-related field of study", "Information Technology / Information Systems / Computer Science / Software Engineering / Cybersecurity / Data Science / Computer Engineering / Other closely related field"],
            ["Before today, had you heard of short online IT professional certificates?", "Yes / No / Not sure"],
            ["Have you ever enrolled in one?", "No / Yes, but did not complete it / Yes, and completed at least one"],
        ],
        [3.4, 2.8],
        caption="Table A.3: Respondent profile",
        font_size=8.4,
    )

    doc.add_heading("Credential Definition", level=2)
    layout.add_note(
        doc,
        "In this questionnaire, a short online IT professional certificate means an assessed online programme in an IT-related field that is offered by a company or university through a platform such as Coursera, edX, Udacity, or a similar service. It is shorter than a degree and provides a certificate after successful completion. Please answer with this type of certificate in mind.",
    )

    doc.add_heading("Behavioral Intention", level=2)
    layout.add_body(
        doc,
        "For each statement, select one response from -3 (strongly disagree) to +3 (strongly agree), with 0 as neutral.",
        no_indent=True,
    )
    layout.add_table(
        doc,
        [
            ["Code", "Statement"],
            ["BI1", "I plan to earn a short online IT professional certificate in the next twelve months."],
            ["BI2", "I intend to pursue a short online IT professional certificate in the next twelve months."],
            ["BI3", "To the extent possible, I plan to pursue a short online IT professional certificate in the next twelve months."],
        ],
        [0.7, 5.5],
        caption="Table A.4: Behavioral Intention items",
        font_size=9,
    )

    doc.add_heading("Attitude", level=2)
    layout.add_body(
        doc,
        "Complete the statement 'For me, pursuing a short online IT professional certificate in the next twelve months is a ... idea' by marking one of seven positions between each pair of endpoints. The endpoint order follows the published instrument.",
        no_indent=True,
    )
    layout.add_table(
        doc,
        [
            ["Code", "First endpoint", "Opposite endpoint"],
            ["ATT1", "Very good", "Very bad"],
            ["ATT2", "Very positive", "Very negative"],
            ["ATT3", "Very helpful", "Very unhelpful"],
        ],
        [0.7, 2.75, 2.75],
        caption="Table A.5: Attitude items",
        font_size=9,
    )

    doc.add_heading("Subjective Norm", level=2)
    layout.add_body(
        doc,
        "For each referent, first rate the normative-belief statement from +3 (strongly agree) to -3 (strongly disagree). Then rate the motivation-to-comply statement from 1 (very undesirable) to 7 (very desirable).",
        no_indent=True,
    )
    sn_rows = [["Code", "Referent", "Questionnaire statement", "Scale"]]
    for index, (display, grammar) in enumerate(REFERENTS, start=1):
        verb = "thinks" if grammar == "the general public" else "think"
        sn_rows.append(
            [
                f"SN-NB{index}",
                display,
                f"{display} {verb} I should pursue a short online IT professional certificate within the next twelve months.",
                "+3 to -3",
            ]
        )
        sn_rows.append(
            [
                f"SN-MC{index}",
                display,
                f"Generally speaking, I do what {grammar} think I should do." if grammar != "the general public" else "Generally speaking, I do what the general public thinks I should do.",
                "1 to 7",
            ]
        )
    layout.add_table(
        doc,
        sn_rows,
        [0.75, 1.1, 3.6, 0.75],
        caption="Table A.6: Subjective Norm items",
        font_size=7.8,
    )

    doc.add_heading("Perceived Behavioral Control", level=2)
    layout.add_body(
        doc,
        "For each control factor, first rate the control-belief statement from +3 (strongly agree) to -3 (strongly disagree). Then rate the facilitation statement from 1 (much easier) to 7 (much more difficult).",
        no_indent=True,
    )
    pbc_rows = [["Code", "Control factor", "Questionnaire statement", "Scale"]]
    for index, (noun, possession) in enumerate(CONTROLS, start=1):
        pbc_rows.append(
            [
                f"PBC-CB{index}",
                noun.title(),
                f"I {possession} to earn a short online IT professional certificate within the next twelve months.",
                "+3 to -3",
            ]
        )
        pbc_rows.append(
            [
                f"PBC-PF{index}",
                noun.title(),
                f"For me, having the {noun} to pursue a short online IT professional certificate would make it [much easier---much more difficult] to earn a short online IT professional certificate in the next twelve months.",
                "1 to 7",
            ]
        )
    layout.add_table(
        doc,
        pbc_rows,
        [0.75, 1.1, 3.6, 0.75],
        caption="Table A.7: Perceived Behavioral Control items",
        font_size=7.8,
    )

    doc.add_heading("Questionnaire Source and Use", level=2)
    layout.add_body(
        doc,
        "Full source citation: Hunsinger, D. S., and Smith, M. A. (2008). Factors that influence information systems undergraduates to pursue IT certification. Journal of Information Technology Education, 7, 247-265. https://doi.org/10.28945/188. The source instrument appears on pp. 252-254. The field questionnaire above retains its four core TPB constructs, referents, control factors, response anchors, and 12-month time horizon.",
    )
    layout.add_table(
        doc,
        [
            ["Construct", "Source location", "Published evidence"],
            ["Behavioral Intention", "p. 252", "Three items; alpha = .94; test-retest r = .64"],
            ["Attitude", "p. 253", "Three semantic evaluations; alpha = .92; test-retest r = .66"],
            ["Subjective Norm", "p. 253", "Six paired beliefs; test-retest r = .56"],
            ["Perceived Behavioral Control", "pp. 253-254", "Four paired control beliefs; test-retest r = .54"],
        ],
        [1.8, 1.0, 3.4],
        caption="Table A.8: Published questionnaire evidence",
        font_size=8.3,
    )
    layout.add_source_note(
        doc,
        "Questionnaire adapted from Hunsinger and Smith (2008), with 'IT certification' replaced by 'short online IT professional certificate' and the published blanks completed using the published referents and control factors. The publisher licenses the journal's articles under CC BY-NC 4.0 (https://creativecommons.org/licenses/by-nc/4.0/; https://www.informingscience.org/Journals/JITEResearch/Overview). No endorsement by the original authors is implied.",
    )


def build_master() -> None:
    OUTPUTS.mkdir(parents=True, exist_ok=True)
    for source_path in (HUNSINGER_FRAMEWORK, WANG_FRAMEWORK):
        if not source_path.exists():
            raise FileNotFoundError(source_path)
    make_author_framework()

    doc = Document()
    configure_document(doc)
    title_page(doc)

    doc.add_heading("Abstract", level=1)
    layout.add_body(
        doc,
        "Short online professional certificates allow learners to complete assessed, career-related study in less time than a degree. This proposed study examines the factors associated with Thai IT students' intention to pursue such certificates when they are offered online by companies or universities through platforms such as Coursera, edX, or Udacity. The study adopts the core Theory of Planned Behavior relationships tested by Hunsinger and Smith (2008) among information-systems undergraduates and independently supported by Wang (2023) in an undergraduate online-learning context. Attitude, Subjective Norm, and Perceived Behavioral Control are the independent variables, and Behavioral Intention to pursue a certificate within the next 12 months is the dependent variable. A quantitative cross-sectional questionnaire survey will be conducted among undergraduate students aged 18 or older in IT-related programmes at universities in Thailand. The questionnaire is adopted from Hunsinger and Smith (2008), with only the named target behavior changed from 'IT certification' to 'short online IT professional certificate.' A pilot study with 40 eligible students will precede the main survey. The three hypotheses will be tested using multiple linear regression. The study measures intention only and will not claim actual certificate completion, learning outcomes, or employment effects.",
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.add_run("Keywords: ").bold = True
    p.add_run("short online professional certificates; IT students; Theory of Planned Behavior; behavioral intention; Thailand")
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    for label, page in [
        ("Abstract", 2),
        ("CHAPTER 1: INTRODUCTION", 4),
        ("CHAPTER 2: LITERATURE REVIEW", 8),
        ("CHAPTER 3: RESEARCH FRAMEWORK", 12),
        ("CHAPTER 4: RESEARCH METHODOLOGY", 16),
        ("APPENDIX A: QUESTIONNAIRE", 19),
        ("References", 23),
    ]:
        base.add_contents_entry(doc, label, page, bold=label != "Abstract")
    doc.add_page_break()

    doc.add_heading("CHAPTER 1: INTRODUCTION", level=1)
    doc.add_heading("1.1 Background of the Study", level=2)
    layout.add_body(
        doc,
        "Micro-credentials are records of assessed learning outcomes that are narrower in scope than conventional degrees and may be combined with further learning. International definitions emphasize assessment, transparency, quality, and portability rather than a particular platform or badge format (Council of the European Union, 2022; UNESCO, 2022). Within information technology, companies and universities offer short online certificate programmes through large learning platforms. The platform delivers the learning environment, while the company or university normally provides the course and issues the certificate.",
    )
    layout.add_body(
        doc,
        "This study focuses on a practical subset of micro-credentials: assessed short online IT professional certificates. Examples include programmes developed by technology companies such as Google, IBM, Microsoft, and Meta, as well as programmes developed by universities and delivered through platforms such as Coursera, edX, or Udacity. These programmes differ from full degrees, attendance-only badges, in-person-only training, and regulated licences. They also differ from traditional exam-only vendor certifications, although both forms are intended to signal focused knowledge or skill.",
    )
    layout.add_body(
        doc,
        "Research on student decisions about credentials and online learning shows that participation is not explained by availability alone. Hunsinger and Smith (2008) found that Attitude, Subjective Norm, and Perceived Behavioral Control each contributed to information-systems undergraduates' intention to pursue IT certification. Wang (2023) found the same three Theory of Planned Behavior relationships among undergraduates considering MOOCs. Research on MOOC use and micro-credential programmes further shows that students' expectations, resources, and learning context matter (Fianu et al., 2018; Miao et al., 2024; Mohan et al., 2020). These findings support an investigation of how Thai IT students evaluate the decision to pursue a short online IT professional certificate.",
    )
    layout.add_body(
        doc,
        "Thailand has participated in regional discussion and development concerning micro-credentials in higher education (Ministry of Higher Education, Science, Research and Innovation, 2022). However, the reviewed studies do not directly examine adult undergraduate IT students enrolled at Thai universities and their intention to pursue company- or university-provided online IT certificates. Evidence from that population may help universities and certificate providers understand whether students' own evaluations, social expectations, or perceived capacity are most closely associated with intention.",
    )

    doc.add_heading("1.2 Statement of the Problem", level=2)
    layout.add_body(
        doc,
        "Existing research addresses related but different behaviors. Hunsinger and Smith (2008) studied predominantly traditional IT certifications in the United States. Wang (2023), Fianu et al. (2018), and Mohan et al. (2020) examined MOOC use rather than the decision to earn an assessed certificate. Miao et al. (2024) examined acceptance of technology used for micro-credential programmes, while Steenkamp et al. (2024) examined using university-issued digital badges in job applications. These studies cannot by themselves explain why Thai IT students would intend to pursue short online IT professional certificates offered by companies or universities through external learning platforms.",
    )
    layout.add_body(
        doc,
        "A focused study is therefore needed in which the target credential, target population, variables, hypotheses, and questionnaire refer to one defined behavior. The present study applies the established core of the Theory of Planned Behavior to intention to pursue a short online IT professional certificate within the next 12 months. It measures student intention rather than assuming that certificate pursuit leads to actual completion, improved learning, employer recognition, or employment outcomes.",
    )

    doc.add_heading("1.3 Research Objectives", level=2)
    layout.add_numbered(
        doc,
        [
            "To examine the relationship between Attitude and Behavioral Intention to pursue a short online IT professional certificate among Thai IT students.",
            "To examine the relationship between Subjective Norm and Behavioral Intention to pursue a short online IT professional certificate among Thai IT students.",
            "To examine the relationship between Perceived Behavioral Control and Behavioral Intention to pursue a short online IT professional certificate among Thai IT students.",
        ],
    )

    doc.add_heading("1.4 Research Questions", level=2)
    layout.add_numbered(
        doc,
        [
            "What is the relationship between Attitude and Behavioral Intention to pursue a short online IT professional certificate among Thai IT students?",
            "What is the relationship between Subjective Norm and Behavioral Intention to pursue a short online IT professional certificate among Thai IT students?",
            "What is the relationship between Perceived Behavioral Control and Behavioral Intention to pursue a short online IT professional certificate among Thai IT students?",
        ],
    )

    doc.add_heading("1.5 Scope of the Research", level=2)
    layout.add_body(
        doc,
        "The population is adults currently enrolled in undergraduate IT-related programmes at universities in Thailand. Relevant fields include Information Technology, Information Systems, Computer Science, Software Engineering, Cybersecurity, Data Science, Computer Engineering, and closely related programmes. The study includes students with and without previous online-certificate experience. It uses an English-language online questionnaire and a quantitative cross-sectional design.",
    )
    layout.add_body(
        doc,
        "The credential is defined as an assessed online programme in an IT-related field, offered by a company or university through a platform such as Coursera, edX, Udacity, or a similar service, completed in substantially less time than a degree, and resulting in a certificate after successful completion. The dependent variable is Behavioral Intention to pursue such a certificate within the next 12 months. The independent variables are Attitude, Subjective Norm, and Perceived Behavioral Control.",
    )

    doc.add_heading("1.6 Significance of the Study", level=2)
    layout.add_body(
        doc,
        "The study contributes a direct application of the Theory of Planned Behavior to a defined online professional-certificate decision among Thai IT students. It also distinguishes certificate pursuit from general platform use and from displaying a digital badge after completion. Practically, the findings may help universities and credential providers identify whether clearer benefits, social encouragement, or practical support are most strongly associated with students' intentions. Any recommendations will be limited to the measured perceptions and intentions.",
    )

    doc.add_heading("1.7 Limitations", level=2)
    layout.add_body(
        doc,
        "Convenience and snowball sampling will limit statistical generalization to all Thai IT students. The cross-sectional design will identify associations but cannot establish causality. Self-reported intention may not result in later enrollment or completion. The English questionnaire may be understood differently across students. Finally, the source instrument was developed for traditional IT certification; the consistent contextual substitution to short online IT professional certificates will therefore be checked in the pilot and reported as a limitation.",
    )

    doc.add_heading("1.8 Definition of Key Terms", level=2)
    definition_rows = [
        ["Term", "Operational definition"],
        ["Short online IT professional certificate", "An assessed online programme in an IT-related field, offered by a company or university through a large online-learning platform, shorter than a degree, and resulting in a certificate after successful completion."],
        ["Attitude", "The student's favorable or unfavorable evaluation of pursuing the defined certificate within the next 12 months (Ajzen, 1991; Hunsinger & Smith, 2008)."],
        ["Subjective Norm", "The student's perception of whether important referents think the student should pursue the defined certificate, weighted by motivation to comply with those referents (Ajzen, 1991; Hunsinger & Smith, 2008)."],
        ["Perceived Behavioral Control", "The student's perceived ability and resources to pursue the defined certificate, represented by control beliefs and the degree to which each factor facilitates or impedes pursuit (Ajzen, 1991; Hunsinger & Smith, 2008)."],
        ["Behavioral Intention", "The student's stated plan or intention to pursue the defined certificate within the next 12 months (Hunsinger & Smith, 2008)."],
    ]
    layout.add_table(doc, definition_rows, [1.7, 4.5], caption="Table 1.1: Operational definitions", font_size=8.4)

    doc.add_page_break()
    doc.add_heading("CHAPTER 2: LITERATURE REVIEW", level=1)
    doc.add_heading("2.1 Short Online Professional Certificates and Micro-Credentials", level=2)
    layout.add_body(
        doc,
        "Micro-credential terminology is not fully uniform. UNESCO (2022) describes micro-credentials through assessed learning outcomes, a limited volume of learning, quality assurance, and potential stand-alone or combined use. The Council of the European Union (2022) similarly emphasizes documented learning outcomes, assessment, quality, transparency, and portability. These definitions support treating short online professional certificates as a bounded form of micro-credential when they require successful assessment and provide an identifiable record of completion.",
    )
    layout.add_body(
        doc,
        "The distinction between provider and platform is important. A company or university may design the curriculum and issue the certificate, while an online platform supplies course hosting, enrollment, assessment delivery, and the learner interface. The present research therefore does not treat 'Coursera certificate' as a single issuer category. It addresses certificates available through Coursera and similar platforms from multiple providers.",
    )
    layout.add_body(
        doc,
        "MOOC research remains relevant because many short online certificates are assembled from online courses. Fianu et al. (2018) found that computer self-efficacy, performance expectancy, and system quality predicted MOOC-use intention in selected Ghanaian universities. Mohan et al. (2020) examined intention and barriers among 412 postgraduate students in India. Miao et al. (2024) studied technology acceptance for micro-credential programmes among 474 university students. These studies explain online-learning participation, but their dependent variables differ from a student's intention to pursue and earn a certificate.",
    )

    doc.add_heading("2.2 Theory of Planned Behavior", level=2)
    layout.add_body(
        doc,
        "The Theory of Planned Behavior proposes that Behavioral Intention is predicted by Attitude toward the behavior, Subjective Norm, and Perceived Behavioral Control (Ajzen, 1991). Attitude captures favorable or unfavorable evaluation. Subjective Norm captures perceived social pressure. Perceived Behavioral Control captures the perceived ease or difficulty of performing the behavior in light of abilities, opportunities, and resources. The relative strength of these predictors may vary by behavior and population.",
    )
    layout.add_body(
        doc,
        "TPB fits this study because pursuit of a short online professional certificate is a specific planned action. It can be evaluated positively or negatively, may be encouraged by relevant people, and depends on resources and capabilities. The theory also permits a precise correspondence among target, action, context, and time: the student intends to pursue the defined certificate in the online context within the next 12 months.",
    )

    doc.add_heading("2.3 Study Variables", level=2)
    doc.add_heading("2.3.1 Attitude", level=3)
    layout.add_body(
        doc,
        "Attitude is the degree to which a person evaluates performing the behavior favorably or unfavorably (Ajzen, 1991). Hunsinger and Smith (2008) measured attitude toward pursuing IT certification using three semantic evaluations: good or bad, positive or negative, and helpful or unhelpful. This direct measure is retained because it concerns the student's evaluation of certificate pursuit rather than a presumed employment outcome.",
    )
    doc.add_heading("2.3.2 Subjective Norm", level=3)
    layout.add_body(
        doc,
        "Subjective Norm is the perceived social pressure to perform or not perform the behavior (Ajzen, 1991). In the IT-certification study, Hunsinger and Smith (2008) measured beliefs about six referents: IT managers, professors, hiring managers, advisors, parents, and the general public. Each normative belief was paired with the student's motivation to comply with that referent.",
    )
    doc.add_heading("2.3.3 Perceived Behavioral Control", level=3)
    layout.add_body(
        doc,
        "Perceived Behavioral Control represents the perceived ease or difficulty of performing the behavior and reflects relevant capacities and resources (Ajzen, 1991). Hunsinger and Smith (2008) used four control factors identified in their student population: learning ability, knowledge, skills, and money and resources. Each control belief was paired with a rating of the degree to which it would make certification easier or more difficult.",
    )
    doc.add_heading("2.3.4 Behavioral Intention", level=3)
    layout.add_body(
        doc,
        "Behavioral Intention indicates how strongly a person plans or intends to perform the specified behavior. Hunsinger and Smith (2008) measured plans and intentions to pursue IT certification within 12 months. The present study retains this time horizon and changes only the credential phrase so that all respondents answer about the same short online IT professional certificate category.",
    )

    doc.add_heading("2.4 Relationships Between the Variables", level=2)
    doc.add_heading("2.4.1 Attitude and Behavioral Intention", level=3)
    layout.add_body(
        doc,
        "TPB predicts stronger intention when attitude toward the behavior is more favorable (Ajzen, 1991). Hunsinger and Smith (2008) found Attitude to be the strongest of the three TPB predictors of IS undergraduates' intention to pursue IT certification. Wang (2023) also found a positive relationship between Attitude and Behavioral Intention in the MOOC context. Ilias et al. (2022) did not find a significant effect of Attitude in their Malaysian internal-auditing certification sample, indicating that the relationship should be tested rather than assumed in the Thai online IT-certificate context.",
    )
    doc.add_heading("2.4.2 Subjective Norm and Behavioral Intention", level=3)
    layout.add_body(
        doc,
        "TPB proposes that favorable perceived social pressure is associated with stronger intention (Ajzen, 1991). Hunsinger and Smith (2008) found a significant positive contribution from Subjective Norm to certification intention. Wang (2023) reported a positive relationship in undergraduate MOOC use, and Ilias et al. (2022) found that social norms positively influenced Malaysian students' intention to pursue internal-auditing certification. These studies support testing the same construct-to-intention path in the present population.",
    )
    doc.add_heading("2.4.3 Perceived Behavioral Control and Behavioral Intention", level=3)
    layout.add_body(
        doc,
        "When a behavior requires ability, time, knowledge, money, or access, perceived control may contribute to intention (Ajzen, 1991). Hunsinger and Smith (2008) found a significant positive contribution from Perceived Behavioral Control to intention to pursue IT certification. Wang (2023) reported a positive relationship with MOOC intention, and Ilias et al. (2022) found a positive effect on intention to pursue internal-auditing certification. The price, workload, and prerequisite knowledge associated with online certificate programmes make this relationship relevant to Thai IT students.",
    )

    doc.add_heading("2.5 Summary of Previous Studies", level=2)
    prior_rows = [
        ["Study", "Population and method", "Framework or main result", "Use in this study"],
        ["Hunsinger & Smith (2008)", "120 US undergraduate IS students; online survey; hierarchical regression.", "Attitude, Subjective Norm, and PBC each positively predicted intention to pursue IT certification.", "Direct framework, hypotheses, variables, questionnaire, scoring, and 12-month timeframe."],
        ["Wang (2023)", "525 undergraduates in China; survey; PLS-SEM.", "Attitude, Subjective Norm, and PBC positively affected MOOC Behavioral Intention.", "Independent online-learning confirmation of the three TPB paths and a published framework figure."],
        ["Ilias et al. (2022)", "243 final-semester accounting students in Malaysia; questionnaire; regression.", "Social norms and PBC affected intention to pursue internal-auditing certification; Attitude was not significant.", "Regional professional-certification comparison."],
        ["Fianu et al. (2018)", "204 students from selected Ghanaian universities; extended UTAUT survey.", "MOOC intention was associated with self-efficacy, performance expectancy, and system quality.", "Online-learning context and evidence that resources and expectations matter."],
        ["Mohan et al. (2020)", "412 postgraduate students at an Indian private university; UTAUT2; PLS-SEM.", "Examined intention and barriers to MOOC use where free courses and certificates were available.", "Developing-country online-learning context."],
        ["Miao et al. (2024)", "474 private-university students in Karachi; TAM and self-determination model; PLS-SEM.", "Examined attitude and intention toward technology used for micro-credential programmes.", "Direct micro-credential context, but not the source of the current questionnaire."],
        ["Steenkamp et al. (2024)", "57 accounting students in New Zealand; extended TAM3; PLS-SEM.", "Examined intention to display university-issued digital badges in job applications.", "Digital-credential comparison; not the target behavior in this study."],
    ]
    layout.add_table(doc, prior_rows, [1.1, 1.65, 1.75, 1.7], caption="Table 2.1: Summary of previous studies", font_size=7.2)

    doc.add_heading("2.6 Research Gap", level=2)
    layout.add_body(
        doc,
        "The literature provides established explanations for traditional IT certification, professional certification, MOOC use, micro-credential technology acceptance, and post-completion badge use. The reviewed studies do not directly combine the target population of Thai undergraduate IT students with the target behavior of pursuing assessed short online IT professional certificates from company or university providers. The present study addresses this bounded population-and-behavior gap using the common TPB paths tested in the two closest framework studies.",
    )

    doc.add_page_break()
    doc.add_heading("CHAPTER 3: RESEARCH FRAMEWORK", level=1)
    doc.add_heading("3.1 Theoretical Frameworks from Previous Studies", level=2)
    layout.add_body(
        doc,
        "This section presents the published frameworks that supply the variables and relationships before presenting the conceptual framework constructed for the present study. The direct IT-certification source is Hunsinger and Smith (2008), and Wang (2023) provides an independent undergraduate online-learning framework containing the same three TPB paths.",
    )

    doc.add_heading("3.1.1 IT-Certification Framework", level=3)
    layout.add_body(
        doc,
        "Hunsinger and Smith (2008) applied the Theory of Planned Behavior to IS students' intention to earn IT certification. Their model connects Attitude, Subjective Norm, and Perceived Behavioral Control to Behavioral Intention. It also shows Behavioral Intention leading to Behavior and a possible direct path from Perceived Behavioral Control to Behavior. The present proposal borrows the three paths ending in Behavioral Intention because actual certificate pursuit or completion will not be measured.",
    )
    layout.add_figure(
        doc,
        HUNSINGER_FRAMEWORK,
        "Figure 3.1: Model of IS students' intention to earn IT certification",
        width=5.9,
        alt_text="Published model showing Attitude, Subjective Norm, and Perceived Behavioral Control leading to Behavioral Intention, and intention and control leading to Behavior.",
    )
    layout.add_source_note(doc, "Source: Hunsinger and Smith (2008, Figure 2, p. 251). Direct crop of the published diagram. CC BY-NC 4.0 (https://creativecommons.org/licenses/by-nc/4.0/).")

    doc.add_heading("3.1.2 Undergraduate MOOC Framework", level=3)
    layout.add_body(
        doc,
        "Wang (2023) integrated the Technology Acceptance Model and the Theory of Planned Behavior to examine undergraduate MOOC use. The published framework includes the same TPB relationships as the present study: Attitude to Behavioral Intention (source H3), Subjective Norm to Behavioral Intention (source H4), and Perceived Behavioral Control to Behavioral Intention (source H5). All three relationships were supported. Perceived Usefulness and Perceived Ease of Use are not included in the present framework because the direct IT-certification study and the core TPB model do not require them.",
    )
    layout.add_figure(
        doc,
        WANG_FRAMEWORK,
        "Figure 3.2: Proposed research model for undergraduate MOOC intention",
        width=4.6,
        alt_text="Published integrated TAM and TPB model showing Perceived Usefulness and Perceived Ease of Use leading to Attitude, and Attitude, Subjective Norm, and Perceived Behavioral Control leading to Behavioral Intention.",
    )
    layout.add_source_note(doc, "Source: Wang (2023, Figure 1, p. 26). Direct crop of the published diagram. CC BY 4.0 (https://creativecommons.org/licenses/by/4.0/).")

    doc.add_heading("3.2 Conceptual Framework of the Study", level=2)
    layout.add_body(
        doc,
        "Figure 3.3 is constructed for this study from the three common TPB paths tested by Hunsinger and Smith (2008) and Wang (2023). The construct names and arrow directions are unchanged. Only the object of intention is specified as a short online IT professional certificate, and actual Behavior is omitted because the cross-sectional questionnaire does not observe later enrollment or completion.",
    )
    layout.add_figure(
        doc,
        AUTHOR_FRAMEWORK,
        "Figure 3.3: Conceptual framework of the study",
        width=6.0,
        alt_text="Author-constructed framework with arrows labeled H1, H2, and H3 from Attitude, Subjective Norm, and Perceived Behavioral Control to Behavioral Intention to Pursue.",
    )
    layout.add_source_note(doc, "Source: Constructed by the author based on Hunsinger and Smith (2008) and Wang (2023).")

    framework_rows = [
        ["Current path", "Hunsinger & Smith (2008)", "Wang (2023)", "Status in current study"],
        ["Attitude -> Behavioral Intention", "H1; tested and supported", "H3; tested and supported", "H1; retained without changing constructs or direction"],
        ["Subjective Norm -> Behavioral Intention", "H2; tested and supported", "H4; tested and supported", "H2; retained without changing constructs or direction"],
        ["PBC -> Behavioral Intention", "H3; tested and supported", "H5; tested and supported", "H3; retained without changing constructs or direction"],
    ]
    layout.add_table(doc, framework_rows, [1.8, 1.45, 1.35, 1.6], caption="Table 3.1: Direct framework and hypothesis lineage", font_size=7.7)

    doc.add_heading("3.3 Research Hypotheses", level=2)
    layout.add_body(
        doc,
        "The hypotheses retain the three construct relationships used by the source studies. The named behavior is changed consistently from traditional IT certification or MOOC use to the defined short online IT professional certificate.",
    )
    add_hypotheses(doc)

    doc.add_heading("3.4 Operationalization of the Variables", level=2)
    op_rows = [
        ["Variable", "Role", "Operational measure", "Direct questionnaire source"],
        ["Attitude", "Independent", "Three seven-point semantic evaluations: good/bad, positive/negative, and helpful/unhelpful.", "Hunsinger & Smith (2008), p. 253; alpha = .92"],
        ["Subjective Norm", "Independent", "Six normative beliefs multiplied by six matching motivation-to-comply ratings and summed.", "Hunsinger & Smith (2008), p. 253; test-retest r = .56"],
        ["Perceived Behavioral Control", "Independent", "Four control beliefs multiplied by four matching facilitation ratings and summed.", "Hunsinger & Smith (2008), pp. 253-254; test-retest r = .54"],
        ["Behavioral Intention", "Dependent", "Three seven-point bipolar agreement items concerning pursuit within the next 12 months.", "Hunsinger & Smith (2008), p. 252; alpha = .94"],
    ]
    layout.add_table(doc, op_rows, [1.25, 0.8, 2.65, 1.5], caption="Table 3.2: Operationalization and direct questionnaire evidence", font_size=7.6)

    adaptation_rows = [
        ["Element", "Published source form", "Form used in this study", "Change"],
        ["Target behavior", "Pursue or earn 'IT certification' within the next 12 months.", "Pursue or earn 'a short online IT professional certificate' within the next 12 months.", "One consistent contextual substitution."],
        ["Constructs", "Attitude, Subjective Norm, Perceived Behavioral Control, Behavioral Intention.", "Same four constructs.", "None."],
        ["Referents", "IT managers, professors, hiring managers, advisors, parents, general public.", "Same six referents.", "None."],
        ["Control factors", "Learning ability, knowledge, skills, money and resources.", "Same four control factors.", "None."],
        ["Blank prompts", "The source prints sentence stems with blanks and separately lists the exact referents and control factors.", "The listed referents and control factors are inserted into their matching source stems.", "Published prompts completed with the published labels; no additional belief content."],
        ["Response formats", "Seven-point bipolar and semantic differential scales.", "Same response formats and anchors.", "None."],
        ["Time horizon", "Next twelve months.", "Next twelve months.", "None."],
    ]
    layout.add_table(doc, adaptation_rows, [1.05, 2.05, 2.05, 1.05], caption="Table 3.3: Questionnaire adoption and contextual substitution", font_size=7.4)

    doc.add_page_break()
    doc.add_heading("CHAPTER 4: RESEARCH METHODOLOGY", level=1)
    doc.add_heading("4.1 Research Design", level=2)
    layout.add_body(
        doc,
        "The study will use a quantitative, explanatory, cross-sectional questionnaire survey. This design permits estimation of the relationships between the three TPB predictors and Behavioral Intention at one point in time. It does not establish causal effects or observe actual certificate completion. The analysis follows the direct source study by treating Behavioral Intention as the dependent variable in a multiple linear regression with Attitude, Subjective Norm, and Perceived Behavioral Control as predictors (Hunsinger & Smith, 2008).",
    )

    doc.add_heading("4.2 Sampling Procedure", level=2)
    doc.add_heading("4.2.1 Target Population", level=3)
    layout.add_body(
        doc,
        "The target population consists of students aged 18 or older who are currently enrolled as undergraduates in IT-related programmes at universities in Thailand. Restricting the population to undergraduates maintains the closest correspondence with Hunsinger and Smith's (2008) IS-student population.",
    )
    doc.add_heading("4.2.2 Sampling Method and Sample Size", level=3)
    layout.add_body(
        doc,
        "Convenience and snowball sampling will be used because a complete national sampling frame of Thai undergraduate IT students is not available. The survey will be distributed through university programmes, student groups, and academic networks. Hunsinger and Smith (2008) identified 90 responses as the minimum for a three-predictor regression and analyzed 120 completed surveys. To provide a larger sample for stable estimation and subgroup description, the present study will seek 400 responses and retain at least 384 eligible and complete responses. The Krejcie and Morgan (1970) figure is used as a conservative recruitment benchmark; because sampling is non-probability based, it is not presented as guaranteeing a 5% margin of error.",
    )
    doc.add_heading("4.2.3 Inclusion and Exclusion Criteria", level=3)
    layout.add_body(
        doc,
        "A response will be included only when the participant provides consent, is at least 18 years old, is currently enrolled as an undergraduate at a university in Thailand, is studying an IT-related programme, and completes all construct questions. Pilot participants will be excluded from the main-study dataset. No exclusion will be based on whether the student has previously pursued a certificate.",
    )

    doc.add_heading("4.3 Research Instrument and Questionnaire Design", level=2)
    layout.add_body(
        doc,
        "The online questionnaire contains participant information and consent, three separate eligibility questions, a short respondent profile, the common credential definition, and the 26 measurement responses adopted from Hunsinger and Smith (2008, pp. 252-254). Separating the eligibility questions allows respondents to consider each requirement individually. A respondent who selects 'No' to consent or any eligibility question will exit the form.",
    )
    instrument_rows = [
        ["Section", "Content", "Purpose"],
        ["A", "Participant information and consent", "Confirm voluntary participation"],
        ["B", "Age, Thai-university enrollment, and IT-programme eligibility", "Apply inclusion criteria"],
        ["C", "Age group, gender, study year, field, awareness, and prior experience", "Describe the sample"],
        ["D", "Neutral definition of a short online IT professional certificate", "Give all respondents the same target"],
        ["E", "3 BI, 3 Attitude, 12 Subjective Norm, and 8 PBC responses", "Measure the four TPB constructs"],
    ]
    layout.add_table(doc, instrument_rows, [0.6, 3.7, 1.9], caption="Table 4.1: Questionnaire structure", font_size=8.3)
    layout.add_body(
        doc,
        "The questionnaire will be administered in English. The measurement response formats are retained from the source instead of converting every item to a five-point scale. Behavioral Intention and belief strength use seven-point bipolar agreement scales; Attitude uses three seven-point semantic differential scales; motivation to comply and facilitation use seven-point ratings. The complete field questionnaire appears in Appendix A.",
    )

    doc.add_heading("4.4 Pilot Study", level=2)
    layout.add_body(
        doc,
        "Before the main survey, the questionnaire will be pilot tested with 40 eligible Thai IT students who will not participate in the main study. This size is within published guidance for initial pilot work and is comparable to the 27-student pilot reported by the direct questionnaire source (Hertzog, 2008; Hunsinger & Smith, 2008; Johanson & Brooks, 2010). The pilot will check whether the definition, instructions, response scales, and contextual substitution are clear and whether the form routing works as intended. Preliminary internal consistency will be examined for Attitude and Behavioral Intention. Consistent with the source study, Cronbach's alpha will not be applied to the Subjective Norm and Perceived Behavioral Control belief composites because their separate beliefs are not assumed to be interchangeable indicators (Hunsinger & Smith, 2008). Any wording change after the pilot will be documented and approved before main data collection.",
    )

    doc.add_heading("4.5 Data Collection", level=2)
    layout.add_body(
        doc,
        "Following academic and ethical approval, the questionnaire will be distributed online through the selected university and student channels. The first page will explain the study purpose, voluntary participation, eligibility, confidentiality, and withdrawal before submission. No name or email address will be required. Data will be stored securely and used only for the stated academic purpose.",
    )

    doc.add_heading("4.6 Scoring and Statistical Treatment", level=2)
    scoring_rows = [
        ["Construct", "Scoring procedure"],
        ["Behavioral Intention", "Mean of BI1-BI3 after coding -3 to +3; a higher score indicates stronger intention."],
        ["Attitude", "Code the three seven-position responses so that higher scores represent more favorable evaluations, then calculate their mean."],
        ["Subjective Norm", "For each referent, multiply the normative-belief score (+3 to -3) by the matching motivation-to-comply score (1 to 7), then sum the six products."],
        ["Perceived Behavioral Control", "For each control factor, multiply the control-belief score (+3 to -3) by the matching facilitation score (1 = much easier to 7 = much more difficult), then sum the four products, following Hunsinger and Smith (2008)."],
    ]
    layout.add_table(doc, scoring_rows, [1.55, 4.65], caption="Table 4.2: Construct scoring", font_size=8.1)
    layout.add_body(
        doc,
        "Frequencies and percentages will describe eligibility and respondent characteristics. Means and standard deviations will summarize questionnaire responses and construct scores. Cronbach's alpha will be reported for the direct multi-item Attitude and Behavioral Intention scales in both pilot and main samples (Cronbach, 1951). The separate belief composites for Subjective Norm and Perceived Behavioral Control will be calculated according to the published source procedure rather than evaluated as reflective alpha scales.",
    )
    layout.add_body(
        doc,
        "One multiple linear regression will test H1-H3, with Behavioral Intention as the dependent variable and Attitude, Subjective Norm, and Perceived Behavioral Control as predictors. The analysis will report R-squared, adjusted R-squared, the overall F-test, coefficients, standard errors, t-values, p-values, and 95% confidence intervals. Variance inflation factors and residual diagnostics will be examined (Hair et al., 2019). A hypothesis will be supported when its coefficient is positive and p is below .05.",
    )

    doc.add_heading("4.7 Ethical Considerations", level=2)
    layout.add_body(
        doc,
        "The study will follow the university's research-ethics requirements. Participation will be voluntary and limited to adults, who may stop before submitting. No names or email addresses will be collected, and results will be reported in aggregate. Provider names are included only as neutral examples.",
    )

    add_questionnaire(doc)

    doc.add_page_break()
    doc.add_heading("References", level=1)
    layout.add_references(doc, references=sorted(REFERENCES, key=str.casefold), font_size=10)
    doc.save(MASTER_OUT)
    print(MASTER_OUT)


if __name__ == "__main__":
    build_master()
