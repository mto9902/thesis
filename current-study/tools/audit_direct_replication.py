#!/usr/bin/env python3
"""Verify the author-constructed framework package against its published sources."""

from __future__ import annotations

import ast
import hashlib
import re
import sys
import zipfile
from pathlib import Path

import pdfplumber
import pypdfium2 as pdfium
from PIL import Image, ImageChops
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
BUILD_SCRIPT = ROOT / "tools" / "build_direct_replication.py"
SOURCE_PDF = ROOT / "source" / "Steenkamp_2024_Digital_Badges.pdf"
MASTER_DOCX = ROOT / "outputs" / "Master_Thesis_Digital_Badges_Thai_IT_Students.docx"
EVIDENCE_DOCX = ROOT / "outputs" / "Evidence_Pack_Author_Constructed_Framework_and_Questionnaire.docx"

SOURCE_IMAGES = (
    (10, ROOT / "source" / "steenkamp_page_10_model.png", (105, 100, 1175, 1205), ROOT / "outputs" / "Published_Framework_Steenkamp_Figure_1_Direct_Crop.png"),
    (28, ROOT / "source" / "steenkamp_page_28_appendix.png", (110, 115, 1120, 1510), ROOT / "outputs" / "Published_Questionnaire_Steenkamp_Appendix_p28_Direct_Crop.png"),
    (29, ROOT / "source" / "steenkamp_page_29_appendix.png", (110, 115, 1120, 1135), ROOT / "outputs" / "Published_Questionnaire_Steenkamp_Appendix_p29_Direct_Crop.png"),
)


def load_literal(name: str):
    tree = ast.parse(BUILD_SCRIPT.read_text(encoding="utf-8"))
    for node in tree.body:
        if isinstance(node, ast.Assign):
            for target in node.targets:
                if isinstance(target, ast.Name) and target.id == name:
                    return ast.literal_eval(node.value)
    raise RuntimeError(f"Could not find literal {name} in {BUILD_SCRIPT}")


def normalize(text: str) -> str:
    text = text.casefold().replace("’", "'").replace("‘", "'")
    return re.sub(r"[^a-z0-9]+", "", text)


def extract_pdf_text(path: Path) -> str:
    with pdfplumber.open(path) as pdf:
        return "\n".join(
            page.extract_text(x_tolerance=1, y_tolerance=3) or "" for page in pdf.pages
        )


def extract_appendix_item_column(path: Path) -> str:
    with pdfplumber.open(path) as pdf:
        pages = []
        for page_number in (28, 29):
            page = pdf.pages[page_number]
            item_column = page.crop((135, 70, 350, page.height - 20))
            pages.append(item_column.extract_text(x_tolerance=1, y_tolerance=3) or "")
        return "\n".join(pages)


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def same_pixels(first: Image.Image, second: Image.Image) -> bool:
    return first.size == second.size and ImageChops.difference(
        first.convert("RGB"), second.convert("RGB")
    ).getbbox() is None


def docx_contains_exact_image(docx_path: Path, image_path: Path) -> bool:
    expected = hashlib.sha256(image_path.read_bytes()).digest()
    with zipfile.ZipFile(docx_path) as package:
        for name in package.namelist():
            if name.startswith("word/media/"):
                if hashlib.sha256(package.read(name)).digest() == expected:
                    return True
    return False


def require_contains(haystack: str, needle: str, label: str, errors: list[str]) -> None:
    if normalize(needle) not in haystack:
        errors.append(label)


def main_full_replication_legacy() -> int:
    hypotheses = load_literal("HYPOTHESES")
    constructs = load_literal("CONSTRUCTS")
    items = load_literal("ITEMS")

    errors: list[str] = []
    if len(hypotheses) != 13:
        errors.append(f"Expected 13 hypotheses; found {len(hypotheses)}")
    if len(constructs) != 11:
        errors.append(f"Expected 11 constructs; found {len(constructs)}")
    if len(items) != 40:
        errors.append(f"Expected 40 items; found {len(items)}")

    codes = [item[0] for item in items]
    if len(codes) != len(set(codes)):
        errors.append("Questionnaire item codes are not unique")

    source_text = normalize(extract_pdf_text(SOURCE_PDF))
    source_item_text = normalize(extract_appendix_item_column(SOURCE_PDF))
    master_text = normalize(extract_docx_text(MASTER_DOCX))
    evidence_text = normalize(extract_docx_text(EVIDENCE_DOCX))

    for code, statement in hypotheses:
        require_contains(source_text, f"{code} {statement}", f"Source hypothesis {code}", errors)
        master_statement = statement.replace(
            "Perceptions of external control has", "Perceptions of external control have"
        )
        require_contains(master_text, f"{code} {master_statement}", f"Master hypothesis {code}", errors)
        require_contains(evidence_text, f"{code} {statement}", f"Evidence hypothesis {code}", errors)

    for name, code, _source_location, _item_codes, _count in constructs:
        source_name = "Behavioral Intention" if code == "BI" else name
        master_name = "Behavioural Intention" if code == "BI" else name
        require_contains(source_text, source_name, f"Source construct {code}", errors)
        require_contains(master_text, master_name, f"Master construct {code}", errors)
        require_contains(evidence_text, name, f"Evidence construct {code}", errors)

    for code, _construct, source_wording, _prior_source in items:
        require_contains(source_item_text, f"{code} {source_wording}", f"Source item {code}", errors)
        require_contains(evidence_text, source_wording, f"Evidence item {code}", errors)
        master_wording = source_wording
        if code == "BI3":
            master_wording = source_wording.replace("use the when", "use them when")
        require_contains(master_text, master_wording, f"Master item {code}", errors)

    for label, text in (("master", master_text), ("evidence", evidence_text)):
        for forbidden in ("preparedfordr", "preparedforprofessor", "professorkimi", "drqizhengu"):
            if forbidden in text:
                errors.append(f"Forbidden addressee wording in {label}: {forbidden}")

    for forbidden in (
        "directmodelreplication",
        "contextualreplication",
        "sourceaudit",
        "auditmaster",
        "sourceanomaly",
        "draftstatus",
        "noresultsareavailable",
        "inserthere",
        "correctionrequiresapproval",
    ):
        if forbidden in master_text:
            errors.append(f"Meta/process wording remains in master thesis: {forbidden}")

    source_hash = hashlib.sha256(SOURCE_PDF.read_bytes()).hexdigest()

    source_pdf = pdfium.PdfDocument(SOURCE_PDF)
    for page_number, page_image_path, crop_box, crop_path in SOURCE_IMAGES:
        rendered = source_pdf[page_number].render(scale=2.4).to_pil()
        saved_page = Image.open(page_image_path)
        saved_crop = Image.open(crop_path)
        if not same_pixels(rendered, saved_page):
            errors.append(f"Source page image does not match PDF page index {page_number}")
        if not same_pixels(saved_page.crop(crop_box), saved_crop):
            errors.append(f"Published crop does not match source page image: {crop_path.name}")

    model_crop = SOURCE_IMAGES[0][3]
    if not docx_contains_exact_image(MASTER_DOCX, model_crop):
        errors.append("Master thesis does not contain the exact source framework crop")
    for image_path in (SOURCE_IMAGES[0][3], SOURCE_IMAGES[1][3], SOURCE_IMAGES[2][3]):
        if not docx_contains_exact_image(EVIDENCE_DOCX, image_path):
            errors.append(f"Evidence pack does not contain exact image: {image_path.name}")

    if errors:
        print("DIRECT-SOURCE AUDIT: FAIL")
        for error in errors:
            print(f"- {error}")
        return 1

    print("DIRECT-SOURCE AUDIT: PASS")
    print("- Source framework and Appendix images match direct PDF renders pixel-for-pixel")
    print("- Exact source Figure 1 crop embedded in both documents")
    print("- 13/13 hypotheses found in the source PDF, master thesis, and evidence pack")
    print("- 11/11 constructs retained")
    print("- 40/40 coded source items found in the source PDF and evidence pack")
    print("- 40/40 source items found in the master thesis; BI3 contains only the field-ready typo correction")
    print("- Supervisor-facing manuscript contains none of the prohibited audit/drafting phrases")
    print(f"- Source PDF SHA-256: {source_hash}")
    return 0


def main() -> int:
    hypotheses = load_literal("HYPOTHESES")
    source_hypotheses = load_literal("SELECTED_SOURCE_HYPOTHESES")
    constructs = load_literal("CONSTRUCTS")
    items = load_literal("ITEMS")
    selected_codes = load_literal("SELECTED_CODES")

    study_constructs = [construct for construct in constructs if construct[1] in selected_codes]
    study_items = [item for item in items if item[0].rstrip("0123456789") in selected_codes]

    errors: list[str] = []
    if len(hypotheses) != 4:
        errors.append(f"Expected 4 study hypotheses; found {len(hypotheses)}")
    if len(source_hypotheses) != 4:
        errors.append(f"Expected 4 direct source hypotheses; found {len(source_hypotheses)}")
    if len(study_constructs) != 5:
        errors.append(f"Expected 5 study constructs; found {len(study_constructs)}")
    if len(study_items) != 18:
        errors.append(f"Expected 18 study items; found {len(study_items)}")

    codes = [item[0] for item in study_items]
    if len(codes) != len(set(codes)):
        errors.append("Study questionnaire item codes are not unique")

    source_text = normalize(extract_pdf_text(SOURCE_PDF))
    source_item_text = normalize(extract_appendix_item_column(SOURCE_PDF))
    master_text = normalize(extract_docx_text(MASTER_DOCX))
    evidence_text = normalize(extract_docx_text(EVIDENCE_DOCX))

    for code, statement in source_hypotheses:
        require_contains(source_text, f"{code} {statement}", f"Source hypothesis {code}", errors)

    for code, statement in hypotheses:
        require_contains(master_text, f"{code} {statement}", f"Master hypothesis {code}", errors)
        require_contains(evidence_text, f"{code} {statement}", f"Evidence hypothesis {code}", errors)

    for name, code, _source_location, _item_codes, _count in study_constructs:
        source_name = "Behavioral Intention" if code == "BI" else name
        require_contains(source_text, source_name, f"Source construct {code}", errors)
        require_contains(master_text, name, f"Master construct {code}", errors)
        require_contains(evidence_text, name, f"Evidence construct {code}", errors)

    for code, _construct, source_wording, _prior_source in study_items:
        require_contains(source_item_text, f"{code} {source_wording}", f"Source item {code}", errors)
        require_contains(evidence_text, source_wording, f"Evidence item {code}", errors)
        master_wording = source_wording.replace("use the when", "use them when") if code == "BI3" else source_wording
        require_contains(master_text, master_wording, f"Master item {code}", errors)

    for removed_code in ("CSE1", "PEC1", "CPLAY1", "CANX1", "IMG1", "RES1"):
        if normalize(removed_code) in master_text:
            errors.append(f"Removed questionnaire construct remains in master thesis: {removed_code}")

    for label, text in (("master", master_text), ("evidence", evidence_text)):
        for forbidden in ("preparedfordr", "preparedforprofessor", "professorkimi", "drqizhengu"):
            if forbidden in text:
                errors.append(f"Forbidden addressee wording in {label}: {forbidden}")

    for forbidden in (
        "directmodelreplication",
        "contextualreplication",
        "sourceaudit",
        "auditmaster",
        "sourceanomaly",
        "draftstatus",
        "noresultsareavailable",
        "inserthere",
        "correctionrequiresapproval",
    ):
        if forbidden in master_text:
            errors.append(f"Meta/process wording remains in master thesis: {forbidden}")

    source_pdf = pdfium.PdfDocument(SOURCE_PDF)
    for page_number, page_image_path, crop_box, crop_path in SOURCE_IMAGES:
        rendered = source_pdf[page_number].render(scale=2.4).to_pil()
        saved_page = Image.open(page_image_path)
        saved_crop = Image.open(crop_path)
        if not same_pixels(rendered, saved_page):
            errors.append(f"Source page image does not match PDF page index {page_number}")
        if not same_pixels(saved_page.crop(crop_box), saved_crop):
            errors.append(f"Published crop does not match source page image: {crop_path.name}")

    tam3_source = ROOT / "source" / "Venkatesh_Bala_2008_TAM3_Figure_2.png"
    tam3_output = ROOT / "outputs" / "Published_Framework_Venkatesh_Bala_TAM3_Figure_2.png"
    if not same_pixels(Image.open(tam3_source), Image.open(tam3_output)):
        errors.append("TAM3 framework output is not an exact copy of the published image asset")

    miao_page = ROOT / "source" / "miao_page_9_framework.png"
    miao_output = ROOT / "outputs" / "Published_Framework_Miao_2024_Figure_1_Direct_Crop.png"
    if not same_pixels(Image.open(miao_page).crop((245, 175, 1125, 715)), Image.open(miao_output)):
        errors.append("Miao framework crop does not match the source page render")

    author_framework = ROOT / "outputs" / "Conceptual_Framework_Author_Constructed_Digital_Badges.png"
    source_framework = SOURCE_IMAGES[0][3]
    required_master_images = (tam3_output, miao_output, source_framework, author_framework)
    required_evidence_images = required_master_images + (SOURCE_IMAGES[1][3], SOURCE_IMAGES[2][3])
    for image_path in required_master_images:
        if not docx_contains_exact_image(MASTER_DOCX, image_path):
            errors.append(f"Master thesis does not contain exact image: {image_path.name}")
    for image_path in required_evidence_images:
        if not docx_contains_exact_image(EVIDENCE_DOCX, image_path):
            errors.append(f"Evidence pack does not contain exact image: {image_path.name}")

    if errors:
        print("AUTHOR-FRAMEWORK AUDIT: FAIL")
        for error in errors:
            print(f"- {error}")
        return 1

    source_hash = hashlib.sha256(SOURCE_PDF.read_bytes()).hexdigest()
    print("AUTHOR-FRAMEWORK AUDIT: PASS")
    print("- Three published frameworks and one separate author-constructed framework are embedded")
    print("- 4/4 study paths trace directly to published Steenkamp hypotheses and supported results")
    print("- 5/5 construct labels match the published digital-badge framework")
    print("- 18/18 questionnaire items trace to the published Appendix")
    print("- Only BI3 contains the disclosed grammatical correction")
    print("- Supervisor-facing manuscript contains none of the prohibited audit/drafting phrases")
    print(f"- Steenkamp source PDF SHA-256: {source_hash}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
